# Copyright 2023-2025 daiyixr
# # SPDX-License-Identifier: Apache-2.0
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#     http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import sys
import re
import io
import json
import time
import random
import string
import tempfile
import requests
import webbrowser
import fitz
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QLabel, QTextEdit, QPushButton, QFileDialog, QMessageBox, 
    QGroupBox, QScrollArea, QInputDialog, QDialog, 
    QLineEdit, QDialogButtonBox, QComboBox, QProgressBar,
    QFormLayout, QSpinBox, QMenu, QAction, QTabWidget, QTableWidget, QTableWidgetItem, QCheckBox,
    QProgressDialog
)
from PyQt5.QtGui import (QIcon, QColor, QPalette, QLinearGradient, 
                         QBrush, QFont, QPixmap, QPainter, QImage, QPen)
from PyQt5.QtCore import Qt, QTimer, QRectF


def _draw_certificate_pixmap(
    size: int,
    base_color: str = "#3E5F53",
    text_color: str = "#FFFFFF",
    border_color: str = "#5A8A7A",
) -> QPixmap:
    """绘制带有证书徽章风格的单个尺寸图标。"""

    image = QImage(size, size, QImage.Format_ARGB32)
    image.fill(QColor(0, 0, 0, 0))

    painter = QPainter(image)
    painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)

    margin = size * 0.1
    rect = QRectF(margin, margin, size - 2 * margin, size - 2 * margin)
    radius = size * 0.22

    painter.setPen(Qt.PenStyle.NoPen)
    painter.setBrush(QColor(base_color))
    painter.drawRoundedRect(rect, radius, radius)

    inner_margin = size * 0.04
    inner_rect = rect.adjusted(inner_margin, inner_margin, -inner_margin, -inner_margin)
    painter.setBrush(Qt.BrushStyle.NoBrush)
    painter.setPen(QPen(QColor(border_color), max(1.0, size * 0.06)))
    painter.drawRoundedRect(inner_rect, radius * 0.85, radius * 0.85)

    painter.setPen(QPen(QColor(text_color)))
    font = QFont("Microsoft YaHei", max(10, int(size * 0.38)))
    font.setBold(True)
    painter.setFont(font)
    painter.drawText(QRectF(0, 0, size, size), Qt.AlignmentFlag.AlignCenter, "敏")

    painter.end()
    return QPixmap.fromImage(image)


def create_pen_icon() -> QIcon:
    """生成用于窗口的多尺寸图标。"""

    icon = QIcon()
    for size in (16, 24, 32, 48, 64, 96, 128):
        icon.addPixmap(_draw_certificate_pixmap(size))
    return icon
from dataclasses import dataclass
from typing import List, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from openpyxl.worksheet.worksheet import Worksheet
    from openpyxl.cell.cell import Cell
    from openpyxl.workbook.workbook import Workbook

@dataclass
class RedactionRule:
    """脱敏规则数据类"""
    rule_id: str
    name: str
    pattern: str  # 正则表达式模式
    example: str
    is_active: bool = True
    regex: str = ""  # 存储编译后的正则表达式
    marker_char: str = ""  # 标记字符，用于全局搜索替换

class RuleEngine:
    """规则引擎核心类"""
    
    def __init__(self):
        self.rules: List[RedactionRule] = []
        
    def add_rule(self, rule: RedactionRule) -> None:
        """添加新规则"""
        self.rules.append(rule)
        
    def get_active_rules(self) -> List[RedactionRule]:
        """获取所有激活的规则"""
        return [r for r in self.rules if r.is_active]
        
    def load_default_rules(self) -> None:
        """加载默认规则集"""
        # 使用内置的完整规则集，默认只激活姓名规则
        self.rules = [
            # 身份信息类（高风险）
            RedactionRule(
                rule_id="name_rule",
                name="姓名",
                pattern="[\u4e00-\u9fa5]{2,4}",
                example="张三 → 张*",
                regex="[\u4e00-\u9fa5]{2,4}",
                marker_char="*",
                is_active=True  # 默认激活姓名规则
            ),
            RedactionRule(
                rule_id="id_card_rule",
                name="身份证号",
                pattern="[1-9]\\d{5}(18|19|20)\\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\\d|3[01])\\d{3}[0-9Xx]",
                example="123456199001011234 → 123***********1234",
                regex="[1-9]\\d{5}(18|19|20)\\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\\d|3[01])\\d{3}[0-9Xx]",
                marker_char="*",
                is_active=False  # 默认关闭其他规则
            ),
            RedactionRule(
                rule_id="passport_rule",
                name="护照号码",
                pattern="[EG]\\d{8}|[A-Z]\\d{7,8}",
                example="E12345678 → E1****678",
                regex="[EG]\\d{8}|[A-Z]\\d{7,8}",
                marker_char="*",
                is_active=False  # 默认关闭
            ),
            
            # 联系方式类（高风险）
            RedactionRule(
                rule_id="phone_rule",
                name="手机号码",
                pattern="1[3-9]\\d{9}",
                example="13812345678 → 138****5678",
                regex="1[3-9]\\d{9}",
                marker_char="*",
                is_active=False  # 默认关闭
            ),
            RedactionRule(
                rule_id="landline_rule",
                name="座机号码",
                pattern="0\\d{2,3}-?\\d{7,8}",
                example="010-12345678 → 010-****5678",
                regex="0\\d{2,3}-?\\d{7,8}",
                marker_char="*",
                is_active=False  # 默认关闭
            ),
            RedactionRule(
                rule_id="email_rule",
                name="邮箱地址",
                pattern="[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,}",
                example="test@example.com → t***@example.com",
                regex="[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-zA-Z]{2,}",
                marker_char="*",
                is_active=False  # 默认关闭
            ),
            
            # 金融信息类（高风险）
            RedactionRule(
                rule_id="bank_card_rule",
                name="银行卡号",
                pattern="\\d{13,19}",
                example="6228480402564890018 → 6228****0018",
                regex="\\d{13,19}",
                marker_char="*",
                is_active=False
            ),
            
            # 地址信息类（中风险）
            RedactionRule(
                rule_id="address_rule",
                name="详细地址",
                pattern="[\u4e00-\u9fa5]{2,}(省|市|区|县|镇|街道|路|号|室|楼|层)[\u4e00-\u9fa5\\d\\-#]*",
                example="北京市朝阳区建国路1号 → 北京市朝阳区建******",
                regex="[\u4e00-\u9fa5]{2,}(省|市|区|县|镇|街道|路|号|室|楼|层)[\u4e00-\u9fa5\\d\\-#]*",
                marker_char="*",
                is_active=False
            ),
            
            # 车辆信息类（中风险）
            RedactionRule(
                rule_id="license_plate_rule",
                name="车牌号码",
                pattern="[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领][A-Z][A-Z0-9]{4}[A-Z0-9挂学警港澳]",
                example="京A12345 → 京A***45",
                regex="[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领][A-Z][A-Z0-9]{4}[A-Z0-9挂学警港澳]",
                marker_char="*",
                is_active=False
            ),
            
            # 企业信息类（中风险）
            RedactionRule(
                rule_id="organization_code_rule",
                name="组织机构代码",
                pattern="[A-Z0-9]{8}-[A-Z0-9]|[A-Z0-9]{9}",
                example="12345678-9 → 123***8-9",
                regex="[A-Z0-9]{8}-[A-Z0-9]|[A-Z0-9]{9}",
                marker_char="*",
                is_active=False
            ),
            RedactionRule(
                rule_id="tax_id_rule",
                name="纳税人识别号",
                pattern="\\d{15}|\\d{17}[0-9X]|\\d{18}|\\d{20}",
                example="123456789012345 → 1234*******2345",
                regex="\\d{15}|\\d{17}[0-9X]|\\d{18}|\\d{20}",
                marker_char="*",
                is_active=False
            ),
            
            # 内部标识类（低风险）
            RedactionRule(
                rule_id="employee_id_rule",
                name="员工工号",
                pattern="[A-Z0-9]{4,10}",
                example="EMP001234 → EMP***234",
                regex="[A-Z0-9]{4,10}",
                marker_char="*",
                is_active=False
            ),
            
            # 自定义字段类（用户可配置）
            RedactionRule(
                rule_id="custom_field_rule",
                name="自定义字段",
                pattern=".*",
                example="任意字段 → 任意*",
                regex=".*",
                marker_char="*",
                is_active=False
            )
        ]
            
    def verify_rule_examples(self):
        """验证所有规则的example和实际脱敏效果是否一致"""
        print("开始验证规则示例...")
        for rule in self.rules:
            if " → " in rule.example:
                parts = rule.example.split(" → ")
                if len(parts) == 2:
                    input_text = parts[0]
                    expected_output = parts[1]
                    actual_output = self.apply_redaction_rule(rule, input_text, None)
                    
                    if actual_output != expected_output:
                        print(f"规则 {rule.name} 不一致:")
                        print(f"  输入: {input_text}")
                        print(f"  期望: {expected_output}")
                        print(f"  实际: {actual_output}")
                    else:
                        print(f"规则 {rule.name} ✓")
        print("验证完成")

    def apply_redaction_rule(self, rule: RedactionRule, text: str, custom_list=None) -> str:
        """应用脱敏规则到文本"""
        if not rule.regex or not text:
            return text
            
        try:
            import re
            
            # 姓名规则特殊处理：只替换自定义名单中的姓名
            if rule.rule_id == "name_rule":
                result = text
                
                # 检查是否有自定义姓名列表
                if custom_list and isinstance(custom_list, list):
                    # 只对自定义名单中的姓名进行替换
                    for name in custom_list:
                        if name and name in result:
                            # 姓名脱敏：保留第一个字，其余用*替换
                            redacted_name = name[0] + "*" * (len(name) - 1)
                            result = result.replace(name, redacted_name)
                    return result
                else:
                    # 没有自定义名单时，不进行任何替换
                    return text
            
            # 自定义字段规则特殊处理：完全复用姓名规则的逻辑
            if rule.rule_id == "custom_field_rule":
                result = text
                
                # 检查是否有自定义字段列表
                if custom_list and isinstance(custom_list, list):
                    # 只对自定义字段列表中的内容进行替换
                    for field in custom_list:
                        if field and field in result:
                            # 自定义字段脱敏：保留第一个字符，其余用*替换（与姓名规则完全相同）
                            if len(field) > 0:
                                redacted_field = field[0] + "*" * (len(field) - 1)
                                result = result.replace(field, redacted_field)
                    return result
                else:
                    # 没有自定义字段列表时，不进行任何替换
                    return text
            
            # 其他规则保持原有逻辑
            # 查找所有匹配的内容
            matches = re.findall(rule.regex, text)
            if not matches:
                return text
                
            result = text
            
            # 处理复杂匹配（可能包含组的情况）
            processed_matches = []
            for match in matches:
                if isinstance(match, tuple):
                    # 如果是元组（来自分组），取非空的组
                    actual_match = next((group for group in match if group), "")
                else:
                    actual_match = match
                
                if actual_match:
                    processed_matches.append(actual_match)
            
            # 如果没有有效匹配，使用原始匹配逻辑
            if not processed_matches:
                processed_matches = [m for m in matches if m]
            
            for match in processed_matches:
                # 根据规则ID或名称应用不同的脱敏策略
                if rule.rule_id == "id_card_rule" and len(match) == 18:
                    # 身份证脱敏：保留前3位和后4位
                    redacted = match[:3] + "*" * 11 + match[-4:]
                elif rule.rule_id == "phone_rule" and len(match) == 11:
                    # 手机号脱敏：保留前3位和后4位
                    redacted = match[:3] + "****" + match[-4:]
                elif rule.rule_id == "landline_rule":
                    # 座机号脱敏：保留区号和后4位，中间用*替换
                    if "-" in match:
                        parts = match.split("-")
                        if len(parts[1]) > 4:
                            redacted = parts[0] + "-" + "*" * (len(parts[1]) - 4) + parts[1][-4:]
                        else:
                            redacted = parts[0] + "-" + "*" * len(parts[1])
                    else:
                        if len(match) > 8:
                            redacted = match[:4] + "*" * (len(match) - 8) + match[-4:]
                        else:
                            redacted = match[:4] + "*" * (len(match) - 4)
                elif rule.rule_id == "email_rule":
                    # 邮箱脱敏：保留第一个字符和@后的内容
                    at_index = match.find('@')
                    if at_index > 0:
                        redacted = match[0] + "*" * (at_index - 1) + match[at_index:]
                    else:
                        redacted = match
                elif rule.rule_id == "address_rule":
                    # 地址脱敏：保留前6个字符，其余用*替换
                    if len(match) > 6:
                        redacted = match[:6] + "*" * (len(match) - 6)
                    else:
                        redacted = "*" * len(match)
                elif rule.rule_id == "bank_card_rule":
                    # 银行卡号脱敏：保留前4位和后4位
                    if len(match) > 8:
                        redacted = match[:4] + "*" * (len(match) - 8) + match[-4:]
                    else:
                        redacted = "*" * len(match)
                elif rule.rule_id == "license_plate_rule":
                    # 车牌号脱敏：保留省份+字母和最后2位，中间用*替换
                    if len(match) >= 7:
                        redacted = match[:2] + "*" * (len(match) - 4) + match[-2:]
                    else:
                        redacted = "*" * len(match)
                elif rule.rule_id == "passport_rule":
                    # 护照号脱敏：保留前2位和后3位
                    if len(match) > 5:
                        redacted = match[:2] + "*" * (len(match) - 5) + match[-3:]
                    else:
                        redacted = "*" * len(match)
                elif rule.rule_id == "organization_code_rule":
                    # 组织机构代码脱敏：保留前3位和后2位（包括-后的部分）
                    if "-" in match:
                        parts = match.split("-")
                        if len(parts[0]) > 3:
                            redacted = parts[0][:3] + "*" * (len(parts[0]) - 4) + parts[0][-1] + "-" + parts[1]
                        else:
                            redacted = match
                    else:
                        if len(match) > 5:
                            redacted = match[:3] + "*" * (len(match) - 5) + match[-2:]
                        else:
                            redacted = "*" * len(match)
                elif rule.rule_id == "tax_id_rule":
                    # 纳税人识别号脱敏：保留前4位和后4位
                    if len(match) > 8:
                        redacted = match[:4] + "*" * (len(match) - 8) + match[-4:]
                    else:
                        redacted = "*" * len(match)
                elif rule.rule_id == "employee_id_rule":
                    # 员工工号脱敏：保留前3位和后3位，中间用*替换
                    if len(match) > 6:
                        redacted = match[:3] + "*" * (len(match) - 6) + match[-3:]
                    elif len(match) > 3:
                        redacted = match[:3] + "*" * (len(match) - 3)
                    else:
                        redacted = "*" * len(match)
                else:
                    # 改进的默认脱敏方式：使用内置算法
                    # 调用内置算法函数，避免全星号替换
                    redacted = self.smart_redact_for_rule_engine(match)
                
                result = result.replace(match, redacted)
            return result
        except Exception:
            return text
    
    def smart_redact_for_rule_engine(self, text):
        """规则引擎专用的内置算法脱敏函数"""
        import re
        
        # 检测文本类型并应用相应脱敏规则
        text = text.strip()
        
        # 身份证号（18位）
        if re.match(r'^\d{18}$', text):
            return text[:3] + "*" * 11 + text[-4:]
        
        # 手机号（11位数字）
        elif re.match(r'^1[3-9]\d{9}$', text):
            return text[:3] + "****" + text[-4:]
        
        # 邮箱地址
        elif '@' in text and re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', text):
            at_index = text.find('@')
            return text[0] + "*" * (at_index - 1) + text[at_index:]
        
        # 中文姓名（2-4个汉字）
        elif re.match(r'^[\u4e00-\u9fa5]{2,4}$', text):
            return text[0] + "*" * (len(text) - 1)
        
        # 银行卡号（16-19位数字）
        elif re.match(r'^\d{16,19}$', text):
            return text[:4] + "*" * (len(text) - 8) + text[-4:]
        
        # 车牌号
        elif re.match(r'^[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领][A-Z]\d{5}$', text):
            return text[:2] + "***" + text[-2:]
        
        # 护照号
        elif re.match(r'^[A-Z]\d{8}$', text):
            return text[:2] + "****" + text[-3:]
        
        # 如果不匹配任何模式，采用保守的脱敏策略
        else:
            # 对于短文本（小于等于3个字符），保留第一个字符
            if len(text) <= 3:
                return text[0] + "*" * (len(text) - 1)
            # 对于长文本，保留前后各1个字符
            elif len(text) <= 10:
                return text[0] + "*" * (len(text) - 2) + text[-1]
            # 对于很长的文本，保留前后各2个字符
            else:
                return text[:2] + "*" * (len(text) - 4) + text[-2:]

class UniversalRedactionTool(QMainWindow):
    def check_update(self):
        """检查软件更新"""
        import webbrowser
        url = "https://gitee.com/daiyixr/universal-desensitization-tool/raw/master/latest_version.json"
        try:
            info = requests.get(url, timeout=10).json()
            latest = info.get("version", "")
            if latest and latest != self.version:
                # 检测到新版本，直接打开百度网盘下载页面
                baidu_url = "https://pan.baidu.com/s/1_eiYyKkYYMZa3ExVkLt3rg?pwd=muxz"  # 替换为你的真实链接
                webbrowser.open(baidu_url)
                # 同时提示提取码
                QMessageBox.information(
                    self, 
                    "发现新版本", 
                    f"当前版本：{self.version}\n最新版本：{latest}\n\n更新内容：{info.get('desc', '')}\n\n已自动打开下载页面\n提取码：muxz（永久有效）"  # 替换为你的真实提取码
                )
            else:
                QMessageBox.information(self, "检查更新", "当前已是最新版本！")
        except Exception as e:
            QMessageBox.warning(self, "检查更新失败", f"检查更新时发生错误：{str(e)}")

    def show_name_redact_dialog(self):
        """弹窗：用户自定义输入需脱敏的姓名列表，支持文本/表格粘贴，自动识别并反馈未识别项"""
        dialog = QDialog(self)
        dialog.setWindowTitle("自定义姓名脱敏")
        dialog.setWindowIcon(self.windowIcon())
        dialog.resize(500, 400)
        dialog.setModal(True)
        layout = QVBoxLayout()
        
        # 说明标签
        instruction_label = QLabel("请粘贴所有需要脱敏的姓名（每行一个，支持表格/文本粘贴）：")
        instruction_label.setStyleSheet("font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(instruction_label)
        
        # 姓名输入框
        name_edit = QTextEdit()
        name_edit.setPlaceholderText("如：\n张三\n李四\n王五\n\n支持批量粘贴，一行一个姓名")
        name_edit.setMinimumHeight(150)
        layout.addWidget(name_edit)
        
        # 结果显示标签
        result_label = QLabel("")
        result_label.setStyleSheet("color: #27ae60; font-size: 10pt; background-color: #f8f9fa; padding: 8px; border-radius: 4px; margin: 10px 0;")
        result_label.setWordWrap(True)
        layout.addWidget(result_label)
        
        # 存储有效姓名的变量
        valid_names = []
        
        def on_confirm():
            nonlocal valid_names
            raw_text = name_edit.toPlainText().strip()
            if not raw_text:
                QMessageBox.warning(dialog, "提示", "请输入姓名列表！")
                return
            
            # 处理姓名列表
            names = [n.strip() for n in raw_text.splitlines() if n.strip()]
            # 简单中文姓名识别（2-4位汉字）
            valid = [n for n in names if re.match(r"^[\u4e00-\u9fa5]{2,4}$", n)]
            invalid = [n for n in names if n not in valid]
            
            valid_names = valid  # 保存有效姓名
            
            msg = f"✅ 识别成功 {len(valid)} 个姓名"
            if valid:
                msg += f"：\n{', '.join(valid[:10])}"
                if len(valid) > 10:
                    msg += f" ... 等{len(valid)}个"
            
            if invalid:
                msg += f"\n\n❌ 未识别（请检查格式）{len(invalid)}个：\n{', '.join(invalid[:5])}"
                if len(invalid) > 5:
                    msg += f" ... 等{len(invalid)}个"
            
            result_label.setText(msg)
            
            if valid:
                # 将有效姓名添加到姓名规则中
                self.update_name_rule_with_custom_names(valid)
                save_btn.setText("已识别")
                self.set_hollow_button(save_btn, "#27ae60", font_size="14px", padding="10px 20px")
        
        def on_save_and_close():
            if valid_names:
                QMessageBox.information(dialog, "保存成功", f"已成功保存 {len(valid_names)} 个自定义姓名到脱敏规则中。\n\n在进行脱敏处理时，这些姓名将被自动识别和脱敏。")
                dialog.accept()
            else:
                # 如果没有已识别的姓名，先尝试识别
                on_confirm()
                if valid_names:
                    QMessageBox.information(dialog, "保存成功", f"已成功保存 {len(valid_names)} 个自定义姓名到脱敏规则中。\n\n在进行脱敏处理时，这些姓名将被自动识别和脱敏。")
                    dialog.accept()
                else:
                    QMessageBox.warning(dialog, "提示", "请先输入并识别姓名后再保存！")
        
        def on_cancel_name():
            # 取消操作：仅关闭窗口
            dialog.reject()
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        
        save_btn = QPushButton("确定并识别")
        self.set_hollow_button(save_btn, "#3498db", font_size="14px", padding="10px 20px")
        save_btn.clicked.connect(on_confirm)
        
        close_btn = QPushButton("取消")
        self.set_hollow_button(close_btn, "#95a5a6", font_size="14px", padding="10px 20px")
        close_btn.clicked.connect(on_cancel_name)
        
        save_and_close_btn = QPushButton("保存并关闭")
        self.set_hollow_button(save_and_close_btn, "#e74c3c", font_size="14px", padding="10px 20px")
        save_and_close_btn.clicked.connect(on_save_and_close)
        
        btn_layout.addWidget(save_btn)
        btn_layout.addWidget(save_and_close_btn)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def show_custom_field_redact_dialog(self):
        """弹窗：用户自定义输入需脱敏的任意字段列表，支持文本/表格粘贴，无格式限制"""
        dialog = QDialog(self)
        dialog.setWindowTitle("自定义字段脱敏")
        dialog.setWindowIcon(self.windowIcon())
        dialog.resize(500, 400)
        dialog.setModal(True)
        layout = QVBoxLayout()
        
        # 说明标签
        instruction_label = QLabel("请粘贴所有需要脱敏的字段内容（每行一个，支持任意字符）：")
        instruction_label.setStyleSheet("font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(instruction_label)
        
        # 字段输入框
        field_edit = QTextEdit()
        field_edit.setPlaceholderText("如：\n公司A\n部门01\nABC123\n项目-X\n\n支持批量粘贴，一行一个字段\n支持中文、英文、数字、符号等任意字符")
        field_edit.setMinimumHeight(150)
        layout.addWidget(field_edit)
        
        # 结果显示标签
        result_label = QLabel("")
        result_label.setStyleSheet("color: #27ae60; font-size: 10pt; background-color: #f8f9fa; padding: 8px; border-radius: 4px; margin: 10px 0;")
        result_label.setWordWrap(True)
        layout.addWidget(result_label)
        
        # 存储有效字段的变量
        valid_fields = []
        
        def on_confirm():
            nonlocal valid_fields
            raw_text = field_edit.toPlainText().strip()
            if not raw_text:
                QMessageBox.warning(dialog, "提示", "请输入字段列表！")
                return
            
            # 处理字段列表
            fields = [f.strip() for f in raw_text.splitlines() if f.strip()]
            # 对于自定义字段，接受任何非空内容
            valid = [f for f in fields if f]
            
            valid_fields = valid  # 保存有效字段
            
            msg = f"✅ 识别成功 {len(valid)} 个字段"
            if valid:
                msg += f"：\n{', '.join(valid[:10])}"
                if len(valid) > 10:
                    msg += f" ... 等{len(valid)}个"
            
            result_label.setText(msg)
            
            if valid:
                # 将有效字段添加到自定义字段规则中
                self.update_custom_field_rule_with_fields(valid)
                save_btn.setText("已识别")
                self.set_hollow_button(save_btn, "#27ae60", font_size="14px", padding="10px 20px")
        
        def on_save_and_close():
            if valid_fields:
                QMessageBox.information(dialog, "保存成功", f"已成功保存 {len(valid_fields)} 个自定义字段到脱敏规则中。\n\n在进行脱敏处理时，这些字段将被自动识别和脱敏。")
                dialog.accept()
            else:
                # 如果没有已识别的字段，先尝试识别
                on_confirm()
                if valid_fields:
                    QMessageBox.information(dialog, "保存成功", f"已成功保存 {len(valid_fields)} 个自定义字段到脱敏规则中。\n\n在进行脱敏处理时，这些字段将被自动识别和脱敏。")
                    dialog.accept()
                else:
                    QMessageBox.warning(dialog, "提示", "请先输入并识别字段后再保存！")
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        
        save_btn = QPushButton("确定并识别")
        self.set_hollow_button(save_btn, "#3498db", font_size="14px", padding="10px 20px")
        save_btn.clicked.connect(on_confirm)
        
        def on_cancel():
            # 取消操作：仅关闭窗口
            dialog.reject()
        
        close_btn = QPushButton("取消")
        self.set_hollow_button(close_btn, "#95a5a6", font_size="14px", padding="10px 20px")
        close_btn.clicked.connect(on_cancel)
        
        save_and_close_btn = QPushButton("保存并关闭")
        self.set_hollow_button(save_and_close_btn, "#e74c3c", font_size="14px", padding="10px 20px")
        save_and_close_btn.clicked.connect(on_save_and_close)
        
        btn_layout.addWidget(save_btn)
        btn_layout.addWidget(save_and_close_btn)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def update_name_rule_with_custom_names(self, custom_names, save_to_file=True):
        """将自定义姓名更新到姓名脱敏规则中"""
        if not custom_names:
            return
        
        # 找到姓名规则
        name_rule = None
        for rule in self.rule_engine.rules:
            if rule.rule_id == "name_rule" or rule.name == "姓名":
                name_rule = rule
                break
        
        if name_rule:
            # 创建自定义姓名的正则表达式
            # 转义特殊字符，然后用 | 连接
            escaped_names = [re.escape(name) for name in custom_names]
            custom_pattern = "|".join(escaped_names)
            
            # 更新规则的匹配模式，同时保留原有的中文姓名通用匹配
            original_pattern = "[\u4e00-\u9fa5]{2,4}"  # 原有的通用中文姓名匹配
            combined_pattern = f"({custom_pattern})|({original_pattern})"
            
            name_rule.pattern = combined_pattern
            name_rule.regex = combined_pattern
            name_rule.is_active = True  # 确保规则被激活
            
            # 更新示例文本
            example_names = custom_names[:3]  # 取前3个作为示例
            if len(custom_names) > 3:
                example_text = f"{example_names[0]} → {example_names[0][0]}*，{example_names[1]} → {example_names[1][0]}*... 等{len(custom_names)}个"
            else:
                example_text = "，".join([f"{name} → {name[0]}*" for name in example_names])
            name_rule.example = example_text
            
            # 保存自定义姓名列表（用于后续显示）
            self.custom_names = custom_names

            # 只在需要时保存到文件（避免加载时重复保存）
            if save_to_file:
                self.save_unified_custom_rules(custom_names=custom_names)
            print(f"已更新姓名规则，包含 {len(custom_names)} 个自定义姓名")
        else:
            QMessageBox.warning(self, "错误", "未找到姓名脱敏规则，无法添加自定义姓名")

    def save_unified_custom_rules(self, custom_names=None, custom_fields=None):
        """统一保存所有自定义规则到一个JSON文件中（每天一个文件）"""
        import os, json, datetime
        rules_dir = os.path.join(os.path.dirname(__file__), "user_custom_rules")
        if not os.path.exists(rules_dir):
            os.makedirs(rules_dir)
        
        date_str = datetime.datetime.now().strftime("%Y%m%d")
        file_path = os.path.join(rules_dir, f"自定义规则{date_str}.json")
        
        # 加载现有数据（如果文件存在）
        existing_data = {}
        if os.path.exists(file_path):
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    existing_data = json.load(f)
            except Exception as e:
                print(f"读取现有规则文件失败: {e}")
        
        # 更新数据
        if custom_names is not None:
            existing_data["custom_names"] = custom_names
        if custom_fields is not None:
            existing_data["custom_fields"] = custom_fields
        
        # 保存更新后的数据
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(existing_data, f, ensure_ascii=False, indent=2)
            print(f"已保存自定义规则到: {file_path}")
        except Exception as e:
            print(f"保存自定义规则失败: {e}")

    def load_unified_custom_rules(self):
        """从统一的自定义规则文件中加载所有规则"""
        import os, json, glob
        rules_dir = os.path.join(os.path.dirname(__file__), "user_custom_rules")
        if not os.path.exists(rules_dir):
            return
        
        files = glob.glob(os.path.join(rules_dir, "自定义规则*.json"))
        if not files:
            return
            
        # 获取最新的文件
        latest_file = max(files, key=os.path.getmtime)
        try:
            with open(latest_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # 加载自定义姓名
            custom_names = data.get("custom_names", [])
            if custom_names:
                self.update_name_rule_with_custom_names(custom_names, save_to_file=False)
                print(f"已加载 {len(custom_names)} 个自定义姓名")
            
            # 加载自定义字段
            custom_fields = data.get("custom_fields", [])
            if custom_fields:
                self.update_custom_field_rule_with_fields(custom_fields, save_to_file=False)
                print(f"已加载 {len(custom_fields)} 个自定义字段")
                
            if custom_names or custom_fields:
                print(f"已自动加载最新自定义规则：{latest_file}")
        except Exception as e:
            print(f"加载自定义规则失败: {e}")

    def update_custom_field_rule_with_fields(self, custom_fields, save_to_file=True):
        """将自定义字段更新到自定义字段脱敏规则中"""
        if not custom_fields:
            return
        
        # 找到自定义字段规则
        custom_field_rule = None
        for rule in self.rule_engine.rules:
            if rule.rule_id == "custom_field_rule" or rule.name == "自定义字段":
                custom_field_rule = rule
                break
        
        if custom_field_rule:
            # 创建自定义字段的正则表达式
            # 转义特殊字符，然后用 | 连接
            escaped_fields = [re.escape(field) for field in custom_fields]
            custom_pattern = "|".join(escaped_fields)
            
            # 更新规则的匹配模式
            custom_field_rule.pattern = custom_pattern
            custom_field_rule.regex = custom_pattern
            custom_field_rule.is_active = True  # 确保规则被激活
            
            # 更新示例文本
            example_fields = custom_fields[:3]  # 取前3个作为示例
            if len(custom_fields) > 3:
                example_text = f"{example_fields[0]} → {example_fields[0][0]}*，{example_fields[1]} → {example_fields[1][0]}*... 等{len(custom_fields)}个"
            else:
                example_text = "，".join([f"{field} → {field[0]}*" for field in example_fields if field])
            custom_field_rule.example = example_text
            
            # 保存自定义字段列表（用于后续显示）
            self.custom_fields = custom_fields

            # 只在需要时保存到文件（避免加载时重复保存）
            if save_to_file:
                self.save_unified_custom_rules(custom_fields=custom_fields)
            print(f"已更新自定义字段规则，包含 {len(custom_fields)} 个自定义字段")
        else:
            QMessageBox.warning(self, "错误", "未找到自定义字段脱敏规则，无法添加自定义字段")

    def load_latest_custom_names(self):
        """自动加载最新的自定义规则JSON文件并应用到规则引擎（保持兼容性）"""
        # 调用统一的加载方法
        self.load_unified_custom_rules()

    # ======== PDF 处理核心方法（PyMuPDF版） ========
    def reset_pdf_state(self):
        """清空当前的PDF解析状态"""
        self.pdf_doc = None
        self.pdf_char_map = []
        self.pdf_font_cache = {}
        self.pdf_display_text = ""
        self.pdf_pending_redactions = []
        self.pdf_fallback_font_alias = None
        self.pdf_font_counter = 0

    def _allocate_pdf_font_name(self, prefix="font_alias"):
        """生成唯一的PDF字体别名，避免重复注册"""
        self.pdf_font_counter += 1
        return f"{prefix}_{self.pdf_font_counter}"

    def normalize_pdf_font_name(self, font_name):
        """清理PDF字体名称中的随机前缀/修饰"""
        if not font_name:
            return ""
        try:
            name = str(font_name)
            if "+" in name:
                name = name.split("+")[-1]
            if "," in name:
                name = name.split(",")[0]
            return name.strip()
        except Exception:
            return str(font_name) if font_name else ""

    def register_pdf_fallback_font(self):
        """为PDF写入注册中文兼容的后备字体"""
        if getattr(self, 'pdf_fallback_font_alias', None):
            return self.pdf_fallback_font_alias

        if not self.pdf_doc:
            return None

        font_candidates = [
            (r"C:/Windows/Fonts/msyh.ttc", 0),
            (r"C:/Windows/Fonts/msyh.ttf", None),
            (r"C:/Windows/Fonts/simsun.ttc", 0),
            (r"C:/Windows/Fonts/simhei.ttf", None),
            (r"C:/Windows/Fonts/simfang.ttf", None),
        ]

        for path, ttc_index in font_candidates:
            try:
                if not os.path.exists(path):
                    continue
                alias_name = self._allocate_pdf_font_name("font_fallback")
                insert_kwargs = {
                    "fontname": alias_name,
                    "subset": False,
                }
                if path.lower().endswith('.ttc'):
                    insert_kwargs["fontfile"] = path
                    insert_kwargs["ttc_index"] = ttc_index or 0
                else:
                    insert_kwargs["fontfile"] = path
                self.pdf_doc.insert_font(**insert_kwargs)
                self.pdf_fallback_font_alias = alias_name
                return alias_name
            except Exception as font_err:
                print(f"注册PDF后备字体失败: {path} -> {font_err}")
                continue

        self.pdf_fallback_font_alias = None
        return None

    def pdf_text_requires_ext_font(self, text):
        """判断文本中是否包含需要CJK/全宽支持的字符"""
        if not text:
            return False
        for ch in text:
            code = ord(ch)
            if code > 127 and not (0x2000 <= code <= 0x206F):  # 排除常见的空格/标点
                return True
        return False

    def normalize_pdf_color(self, color_value):
        """将PyMuPDF颜色值统一转换为RGB元组"""
        try:
            if isinstance(color_value, (tuple, list)) and len(color_value) >= 3:
                comps = [float(c) for c in color_value[:3]]
                max_comp = max(comps) if comps else 1.0
                if max_comp > 1.0:
                    comps = [c / 255.0 for c in comps]
                return tuple(comps)
            if isinstance(color_value, int):
                r, g, b = fitz.utils.int_to_rgb(color_value)
                return (r / 255.0, g / 255.0, b / 255.0)
        except Exception:
            pass
        return (0.0, 0.0, 0.0)

    def estimate_char_bbox(self, span_bbox, char_index, total_chars):
        """在缺少逐字符位置信息时，按平均宽度估算字符边界"""
        if not span_bbox or total_chars <= 0:
            return [0, 0, 0, 0]
        x0, y0, x1, y1 = span_bbox
        width = max((x1 - x0) / max(total_chars, 1), 0.5)
        start_x = x0 + width * char_index
        end_x = start_x + width
        return [start_x, y0, end_x, y1]

    def load_pdf_with_pymupdf(self, pdf_path):
        """使用PyMuPDF解析PDF并构建字符映射"""
        try:
            doc = fitz.open(pdf_path)
        except Exception as e:
            QMessageBox.warning(self, "警告", f"无法打开PDF文件: {str(e)}")
            return None

        self.pdf_doc = doc
        self.pdf_char_map = []
        self.pdf_display_text = ""
        display_chars = []
        char_index = 0

        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            try:
                raw_dict = page.get_text("rawdict")
            except Exception:
                raw_dict = None

            if not raw_dict:
                continue

            blocks = raw_dict.get("blocks", [])
            last_font = "helv"
            last_size = 12.0
            last_color = (0.0, 0.0, 0.0)

            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        font = span.get("font") or last_font
                        size = float(span.get("size") or last_size)
                        color_value = span.get("color", 0)
                        rgb_color = self.normalize_pdf_color(color_value)
                        chars = span.get("chars")
                        text = span.get("text", "")
                        span_bbox = span.get("bbox")

                        if chars:
                            iterable = chars
                        else:
                            iterable = []
                            for idx, ch in enumerate(text):
                                if not ch:
                                    continue
                                estimated_bbox = self.estimate_char_bbox(span_bbox, idx, len(text))
                                iterable.append({"c": ch, "bbox": estimated_bbox})

                        for char_info in iterable:
                            char_text = char_info.get("c", "")
                            if not char_text:
                                continue
                            bbox = char_info.get("bbox")
                            display_chars.append(char_text)
                            self.pdf_char_map.append({
                                "index": char_index,
                                "char": char_text,
                                "page": page_index,
                                "bbox": bbox,
                                "font": font,
                                "size": size,
                                "color": rgb_color,
                            })
                            char_index += 1

                        last_font = font
                        last_size = size
                        last_color = rgb_color

                    # 行末追加换行符，保持展示结构
                    display_chars.append("\n")
                    self.pdf_char_map.append({
                        "index": char_index,
                        "char": "\n",
                        "page": page_index,
                        "bbox": None,
                        "font": last_font,
                        "size": last_size,
                        "color": last_color,
                    })
                    char_index += 1

            # 页面末尾再补充一个换行，分隔页面
            if display_chars and display_chars[-1] != "\n":
                display_chars.append("\n")
                self.pdf_char_map.append({
                    "index": char_index,
                    "char": "\n",
                    "page": page_index,
                    "bbox": None,
                    "font": last_font,
                    "size": last_size,
                    "color": last_color,
                })
                char_index += 1

        self.pdf_display_text = ''.join(display_chars)
        return self.pdf_display_text

    def build_pdf_font_cache(self):
        """缓存PDF中使用的字体，方便后续复用原字体"""
        if not self.pdf_doc:
            return

        font_map = {}
        for page_index in range(self.pdf_doc.page_count):
            for font_info in self.pdf_doc.get_page_fonts(page_index):
                xref = font_info[0]
                base_name = font_info[3]
                if base_name and base_name not in font_map:
                    font_map[base_name] = xref

        alias_cache = {}
        for base_name, xref in font_map.items():
            try:
                font_tuple = self.pdf_doc.extract_font(xref)
                if not font_tuple:
                    continue
                font_data = None
                if isinstance(font_tuple, dict):
                    font_data = font_tuple.get("fontfile") or font_tuple.get("stream")
                elif isinstance(font_tuple, (tuple, list)):
                    for item in font_tuple:
                        if isinstance(item, (bytes, bytearray)) and item:
                            font_data = item
                            break
                if not font_data:
                    continue
                alias_name = self._allocate_pdf_font_name("font_alias")
                self.pdf_doc.insert_font(fontname=alias_name, fontbuffer=font_data, subset=False)
                alias_cache[base_name] = alias_name
                normalized = self.normalize_pdf_font_name(base_name)
                if normalized and normalized not in alias_cache:
                    alias_cache[normalized] = alias_name
            except Exception:
                continue
        self.pdf_font_cache = alias_cache
        self.register_pdf_fallback_font()

    def ensure_pdf_font_context(self):
        """确保在交互式脱敏时具备可用的PDF字体上下文"""
        # 尝试在需要时重新打开PDF，避免pdf_doc为None
        if not getattr(self, 'pdf_doc', None):
            pdf_path = getattr(self, 'input_file_path', None)
            if pdf_path and os.path.exists(pdf_path):
                try:
                    self.pdf_doc = fitz.open(pdf_path)
                except Exception as reopen_err:
                    print(f"重新打开PDF失败: {reopen_err}")
                    return False
            else:
                return False

        if not getattr(self, 'pdf_font_cache', None):
            self.pdf_font_cache = {}

        if not self.pdf_font_cache:
            self.build_pdf_font_cache()

        if not self.pdf_fallback_font_alias:
            self.register_pdf_fallback_font()

        return True

    def get_pdf_font_alias(self, font_name):
        """根据原字体名称获取可用于写入的字体别名"""
        cache = getattr(self, 'pdf_font_cache', {}) or {}
        candidates = []
        if font_name:
            candidates.append(font_name)
            normalized = self.normalize_pdf_font_name(font_name)
            if normalized and normalized not in candidates:
                candidates.append(normalized)

        for name in candidates:
            if name in cache:
                alias = cache[name]
                if font_name and font_name not in cache:
                    cache[font_name] = alias
                return alias

        fallback_alias = getattr(self, 'pdf_fallback_font_alias', None)
        if fallback_alias:
            return fallback_alias

        registered_alias = self.register_pdf_fallback_font()
        if registered_alias:
            return registered_alias

        return "helv"

    def prepare_pdf_redaction_segments(self, start_index, redacted_text):
        """根据字符索引生成PDF脱敏片段和撤销快照"""
        if not self.pdf_char_map:
            return [], []

        end_index = min(start_index + len(redacted_text), len(self.pdf_char_map))
        segments = []
        backup_chars = []
        current_segment = None

        for offset, char_pos in enumerate(range(start_index, end_index)):
            if char_pos >= len(self.pdf_char_map):
                break
            entry = self.pdf_char_map[char_pos]
            original_char = entry.get('char', '')
            replacement_char = redacted_text[offset] if offset < len(redacted_text) else original_char

            backup_chars.append({'index': char_pos, 'char': original_char})
            self.pdf_char_map[char_pos]['char'] = replacement_char

            bbox = entry.get('bbox')
            if not bbox:
                continue

            page = entry.get('page', 0)
            font = entry.get('font', 'helv')
            size = entry.get('size', 12.0)
            color = entry.get('color', (0.0, 0.0, 0.0))

            if (not current_segment) or current_segment['page'] != page:
                if current_segment:
                    segments.append(current_segment)
                current_segment = {
                    'page': page,
                    'min_x': bbox[0],
                    'min_y': bbox[1],
                    'max_x': bbox[2],
                    'max_y': bbox[3],
                    'font': font,
                    'size': size,
                    'color': color,
                    'original_chars': [original_char],
                    'redacted_chars': [replacement_char],
                    'indices': [char_pos]
                }
            else:
                current_segment['min_x'] = min(current_segment['min_x'], bbox[0])
                current_segment['min_y'] = min(current_segment['min_y'], bbox[1])
                current_segment['max_x'] = max(current_segment['max_x'], bbox[2])
                current_segment['max_y'] = max(current_segment['max_y'], bbox[3])
                current_segment['original_chars'].append(original_char)
                current_segment['redacted_chars'].append(replacement_char)
                current_segment['indices'].append(char_pos)

        if current_segment:
            segments.append(current_segment)

        formatted_segments = []
        for seg in segments:
            formatted_segments.append({
                'page': seg['page'],
                'rect': [seg['min_x'], seg['min_y'], seg['max_x'], seg['max_y']],
                'font': seg['font'],
                'size': seg['size'],
                'color': seg['color'],
                'original': ''.join(seg['original_chars']),
                'redacted': ''.join(seg['redacted_chars']),
                'indices': seg['indices']
            })

        return formatted_segments, backup_chars

    def restore_pdf_characters(self, backups):
        """根据快照恢复PDF字符映射"""
        if not backups:
            return

        for info in backups:
            index = info.get('index')
            char_val = info.get('char')
            if index is None:
                continue
            if 0 <= index < len(self.pdf_char_map):
                self.pdf_char_map[index]['char'] = char_val

    def build_pdf_operations_from_text(self, original_text, updated_text, base_context=None, context_callback=None):
        """根据原始/更新文本差异构建PDF脱敏操作列表"""
        if original_text is None or updated_text is None:
            return []

        if isinstance(base_context, dict):
            base_context = dict(base_context)
        else:
            base_context = {}
        diff_ranges = self.calculate_text_diff_ranges(original_text, updated_text)
        operations = []

        for start, end in diff_ranges:
            original_segment = original_text[start:end]
            replacement = updated_text[start:end]

            segments, backup = self.prepare_pdf_redaction_segments(start, replacement)
            if not segments:
                self.restore_pdf_characters(backup)
                continue

            end_index = start + len(original_segment)
            operation = {
                'start': start,
                'end': end_index,
                'original': original_segment,
                'redacted': replacement,
                'segments': segments,
                'char_backup': backup,
                'timestamp': self.get_current_timestamp()
            }

            if callable(context_callback):
                extra_context = context_callback(start, end, original_segment, replacement)
                if isinstance(extra_context, dict):
                    operation.update(extra_context)

            if base_context:
                operation.update(base_context)
            operations.append(operation)

        if operations:
            self.pdf_display_text = ''.join(entry.get('char', '') for entry in self.pdf_char_map)

        return operations

    def build_default_mask(self, text):
        """为给定文本生成与长度一致的默认掩码"""
        if not text:
            return text

        length = len(text)
        if length == 1:
            return "*"
        if length == 2:
            return text[0] + "*"
        return text[0] + ("*" * (length - 2)) + text[-1]

    def generate_redacted_text(self, original_text):
        """保留原始空白字符结构的脱敏文本"""
        if not original_text:
            return original_text

        match = re.match(r'^(\s*)(.*?)(\s*)$', original_text, re.DOTALL)
        if not match:
            return self.smart_redact_text(original_text)

        leading, core, trailing = match.groups()
        if not core:
            return original_text

        redacted_core = self.smart_redact_text(core)
        if (not redacted_core) or len(redacted_core) != len(core):
            redacted_core = self.build_default_mask(core)

        return f"{leading}{redacted_core}{trailing}"

    def ensure_pdf_text_color(self, color):
        """确保写入PDF的文本颜色具有足够对比度"""
        try:
            if isinstance(color, (list, tuple)) and len(color) >= 3:
                normalized = tuple(max(0.0, min(1.0, float(c))) for c in color[:3])
            else:
                normalized = (0.0, 0.0, 0.0)
        except Exception:
            normalized = (0.0, 0.0, 0.0)

        brightness = sum(normalized) / 3.0 if normalized else 0.0
        if brightness >= 0.85:
            return (0.0, 0.0, 0.0)
        return normalized

    def apply_pdf_segment(self, page, segment):
        """在指定页面应用单个脱敏片段（使用红线脱敏）"""
        try:
            rect = fitz.Rect(segment.get('rect', [0, 0, 0, 0]))
            if rect.is_empty or rect.width == 0 or rect.height == 0:
                return False

            fontsize = float(segment.get('size', 12.0)) or 12.0
            text_color = self.ensure_pdf_text_color(segment.get('color', (0.0, 0.0, 0.0)))
            text = segment.get('redacted', '')

            if not text:
                original_text = segment.get('original', '')
                if original_text:
                    text = self.build_default_mask(original_text)
                else:
                    text = "***"

            font_alias = self.get_pdf_font_alias(segment.get('font', 'helv'))
            if self.pdf_text_requires_ext_font(text) and font_alias == "helv":
                fallback_alias = self.register_pdf_fallback_font()
                if fallback_alias:
                    font_alias = fallback_alias

            page.add_redact_annot(
                rect,
                text=text,
                fill=(1, 1, 1),
                fontname=font_alias,
                fontsize=fontsize,
                text_color=text_color,
                align=fitz.TEXT_ALIGN_LEFT
            )
            return True
        except Exception as e:
            print(f"应用PDF片段失败: {e}")
            return False

    def calculate_text_diff_ranges(self, original_text, updated_text):
        """计算原文本与更新后文本的差异区间"""
        ranges = []
        start = None
        length = min(len(original_text), len(updated_text))

        for idx in range(length):
            if original_text[idx] != updated_text[idx]:
                if start is None:
                    start = idx
            else:
                if start is not None:
                    ranges.append((start, idx))
                    start = None

        if start is not None:
            ranges.append((start, length))

        return ranges

    def auto_redact_pdf(self):
        """根据当前激活规则自动对PDF文本进行脱敏"""
        if not self.pdf_char_map:
            return [], ''

        pdf_text = ''.join(entry.get('char', '') for entry in self.pdf_char_map)
        operations = []

        for rule in self.rule_engine.get_active_rules():
            custom_list = None
            if rule.rule_id == "name_rule":
                custom_list = getattr(self, 'custom_names', None)
            elif rule.rule_id == "custom_field_rule":
                custom_list = getattr(self, 'custom_fields', None)

            processed_text = self.rule_engine.apply_redaction_rule(rule, pdf_text, custom_list)
            if processed_text == pdf_text:
                continue

            base_context = {
                'type': 'auto',
                'rule_name': rule.name,
                'mode': '自动规则脱敏',
                'rule_type': '规则引擎'
            }
            rule_operations = self.build_pdf_operations_from_text(pdf_text, processed_text, base_context)
            if rule_operations:
                operations.extend(rule_operations)

            pdf_text = processed_text

        self.pdf_display_text = pdf_text
        return operations, pdf_text

    def is_pdf_image_based(self, pdf_path):
        """检测PDF是否为图片型（扫描件）"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(pdf_path)
            text_pages = 0
            total_pages = len(doc)
            
            for page_num in range(min(3, total_pages)):  # 检查前3页或全部页面
                page = doc[page_num]
                text_content = page.get_text().strip()
                if text_content:
                    text_pages += 1
            
            doc.close()
            # 如果前3页都没有文本，认为是图片型PDF
            return text_pages == 0
        except Exception as e:
            print(f"PDF类型检测失败: {e}")
            return True  # 出错时保守处理，认为是图片型

    def closeEvent(self, a0):
        """程序退出时的处理"""
        # 检查是否有自定义规则需要清除
        has_custom_names = hasattr(self, 'custom_names') and self.custom_names
        has_custom_fields = hasattr(self, 'custom_fields') and self.custom_fields
        
        if has_custom_names or has_custom_fields:
            reply = QMessageBox.question(self, '确认', 
                                       '是否清除本次使用的自定义规则？\n'
                                       '选择"是"将清除内存中的自定义规则并删除JSON文件（不可恢复）\n'
                                       '选择"否"将保留规则到下次启动',
                                       QMessageBox.StandardButton.Yes,
                                       QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                # 清除自定义规则
                if has_custom_names:
                    self.custom_names = []
                    # 重置姓名规则为默认状态
                    for rule in self.rule_engine.rules:
                        if rule.rule_id == "name_rule" or rule.name == "姓名":
                            rule.pattern = "[\u4e00-\u9fa5]{2,4}"
                            rule.regex = "[\u4e00-\u9fa5]{2,4}"
                            rule.example = "张三 → 张*，李四 → 李*"
                            break
                
                if has_custom_fields:
                    self.custom_fields = []
                    # 重置自定义字段规则为默认状态
                    for rule in self.rule_engine.rules:
                        if rule.rule_id == "custom_field_rule" or rule.name == "自定义字段":
                            rule.pattern = ""
                            rule.regex = ""
                            rule.example = "请先配置自定义字段"
                            rule.is_active = False
                            break
                
                # 删除自定义JSON文件
                import os, glob
                try:
                    rules_dir = os.path.join(os.path.dirname(__file__), "user_custom_rules")
                    if os.path.exists(rules_dir):
                        files = glob.glob(os.path.join(rules_dir, "自定义规则*.json"))
                        deleted_count = 0
                        for file_path in files:
                            try:
                                os.remove(file_path)
                                deleted_count += 1
                            except Exception as e:
                                print(f"删除文件失败 {file_path}: {e}")
                        
                        if deleted_count > 0:
                            print(f"已清除自定义规则及 {deleted_count} 个JSON文件")
                        else:
                            print("已清除自定义规则")
                except Exception as e:
                    print(f"清除JSON文件时出错: {e}")
        
        super().closeEvent(a0)

    def __init__(self):
        super().__init__()
        self.version = "2.4.1"  # 添加版本号属性
        self.setWindowTitle("FileMasker")
        self.setWindowIcon(self.get_app_icon())
        self.setGeometry(200, 120, 800, 650)
        self.setup_ui()
        self.setup_styles()
        
        # 初始化文档对象
        self.current_word_doc = None
        
        # 初始化撤销历史记录
        self.text_redaction_history = []  # 文本脱敏历史记录
        self.word_redaction_history = []  # Word文档脱敏历史记录
        self.pdf_redaction_history = []   # PDF文档脱敏历史记录
        self.excel_redaction_history = []  # Excel脱敏历史记录
        
        # 初始化Excel格式存储
        self.excel_cell_formats = {}  # 存储每个单元格的原始格式信息
        self.original_excel_path = None  # 存储原始Excel文件路径
        
        # 初始化日志导出相关变量
        self.current_redaction_log = []  # 当前操作的脱敏日志
        
        # 自动加载最新的自定义规则
        self.load_latest_custom_names()
        
        # 初始化PDF相关状态
        self.is_pdf_source = False  # 标记当前文件是否来源于PDF
        self.reset_pdf_state()

    def save_cell_format(self, cell, row, col):
        """保存单元格的格式信息"""
        try:
            # 保存单元格的关键格式属性
            format_info = {
                'font': {
                    'name': cell.font.name if cell.font and cell.font.name else None,
                    'size': cell.font.size if cell.font and cell.font.size else None,
                    'bold': cell.font.bold if cell.font else None,
                    'italic': cell.font.italic if cell.font else None,
                    'color': str(cell.font.color.rgb) if cell.font and cell.font.color and hasattr(cell.font.color, 'rgb') else None,
                },
                'fill': {
                    'fill_type': str(cell.fill.fill_type) if cell.fill else None,
                    'start_color': str(cell.fill.start_color.rgb) if cell.fill and cell.fill.start_color and hasattr(cell.fill.start_color, 'rgb') else None,
                },
                'border': {
                    'left': str(cell.border.left.style) if cell.border and cell.border.left else None,
                    'right': str(cell.border.right.style) if cell.border and cell.border.right else None,
                    'top': str(cell.border.top.style) if cell.border and cell.border.top else None,
                    'bottom': str(cell.border.bottom.style) if cell.border and cell.border.bottom else None,
                },
                'alignment': {
                    'horizontal': str(cell.alignment.horizontal) if cell.alignment and cell.alignment.horizontal else None,
                    'vertical': str(cell.alignment.vertical) if cell.alignment and cell.alignment.vertical else None,
                    'wrap_text': cell.alignment.wrap_text if cell.alignment else None,
                },
                'number_format': cell.number_format if hasattr(cell, 'number_format') else None,
            }
            
            # 使用(row, col)作为键存储格式信息
            self.excel_cell_formats[(row, col)] = format_info
        except Exception as e:
            # 如果保存格式失败，忽略错误但记录日志
            print(f"警告：保存单元格({row}, {col})格式失败: {str(e)}")

    def apply_cell_format(self, cell, row, col):
        """将保存的格式应用到单元格"""
        try:
            format_info = self.excel_cell_formats.get((row, col))
            if not format_info:
                return
            
            from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
            
            # 应用字体格式
            font_info = format_info.get('font', {})
            if any(font_info.values()):
                cell.font = Font(
                    name=font_info.get('name'),
                    size=font_info.get('size'),
                    bold=font_info.get('bold'),
                    italic=font_info.get('italic'),
                    color=font_info.get('color')
                )
            
            # 应用填充格式
            fill_info = format_info.get('fill', {})
            if any(fill_info.values()):
                cell.fill = PatternFill(
                    fill_type=fill_info.get('fill_type'),
                    start_color=fill_info.get('start_color')
                )
            
            # 应用边框格式
            border_info = format_info.get('border', {})
            if any(border_info.values()):
                cell.border = Border(
                    left=Side(style=border_info.get('left')),
                    right=Side(style=border_info.get('right')),
                    top=Side(style=border_info.get('top')),
                    bottom=Side(style=border_info.get('bottom'))
                )
            
            # 应用对齐格式
            alignment_info = format_info.get('alignment', {})
            if any(alignment_info.values()):
                cell.alignment = Alignment(
                    horizontal=alignment_info.get('horizontal'),
                    vertical=alignment_info.get('vertical'),
                    wrap_text=alignment_info.get('wrap_text')
                )
            
            # 应用数字格式
            number_format = format_info.get('number_format')
            if number_format:
                cell.number_format = number_format
                
        except Exception as e:
            print(f"警告：应用单元格({row}, {col})格式失败: {str(e)}")

    def get_app_icon(self):
        # 使用更具标识性的自定义徽章图标
        # create_pen_icon() 使用了 _draw_certificate_pixmap 函数来绘制不同尺寸的设备像素图
        return create_pen_icon()

    @staticmethod
    def _hex_to_rgb(hex_color):
        """Convert hex color string to RGB tuple."""
        if not hex_color:
            return 52, 152, 219
        color = hex_color.strip()
        if color.startswith("#"):
            color = color[1:]
        if len(color) == 3:
            color = ''.join(ch * 2 for ch in color)
        try:
            r = int(color[0:2], 16)
            g = int(color[2:4], 16)
            b = int(color[4:6], 16)
            return r, g, b
        except (ValueError, IndexError):
            return 52, 152, 219

    def set_hollow_button(
        self,
        button,
        color="#3498db",
        *,
        text_color=None,
        padding="8px 16px",
        radius=5,
        font_size=None,
        bold=True,
        min_width=None,
        hover_alpha=0.14,
        pressed_alpha=0.24,
    ):
        """Apply a consistent hollow button style."""
        r, g, b = self._hex_to_rgb(color)
        text_color = text_color or color
        font_size_line = f"font-size: {font_size};" if font_size else ""
        font_weight_line = "font-weight: bold;" if bold else ""
        min_width_line = f"min-width: {min_width};" if min_width else ""
        hover_bg = f"rgba({r}, {g}, {b}, {hover_alpha:.2f})"
        pressed_bg = f"rgba({r}, {g}, {b}, {pressed_alpha:.2f})"
        disabled_color = f"rgba({r}, {g}, {b}, 0.35)"
        button.setStyleSheet(
            f"""
            QPushButton {{
                background-color: transparent;
                color: {text_color};
                border: 2px solid {color};
                border-radius: {radius}px;
                padding: {padding};
                {font_weight_line}
                {font_size_line}
                {min_width_line}
            }}
            QPushButton:hover {{
                background-color: {hover_bg};
            }}
            QPushButton:pressed {{
                background-color: {pressed_bg};
            }}
            QPushButton:disabled {{
                color: {disabled_color};
                border-color: {disabled_color};
                background-color: transparent;
            }}
        """
        )

    def setup_ui(self):
        # 主窗口布局
        main_widget = QWidget()
        main_layout = QVBoxLayout()
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # 标题区域
        title_label = QLabel("通用脱敏工具")
        title_label.setFont(QFont("Arial", 18, QFont.Bold))
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(title_label)
        version_label = QLabel(f"版本:{self.version} | 2025 D&Ai ")
        version_label.setObjectName("version_label")
        self.version_label = version_label  # 保存为实例变量
        version_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        version_label.setStyleSheet("color: #7f8c8d; font-size: 10pt;")
        main_layout.addWidget(version_label)

        # 文件操作区（前置）
        file_group = QGroupBox("📁 文件处理")
        file_layout = QVBoxLayout()
        
        # 处理模式选择
        mode_layout = QHBoxLayout()
        mode_label = QLabel("处理模式:")
        mode_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.mode_combo = QComboBox()
        self.mode_combo.addItems([
            "🎯 交互式脱敏（推荐）",
            "⚙️ 自动脱敏（规则模式）"
        ])
        self.mode_combo.setCurrentIndex(0)  # 默认选择交互式脱敏
        self.mode_combo.currentIndexChanged.connect(self.on_mode_changed)
        self.rule_config_btn = QPushButton("📋 配置脱敏规则")
        self.set_hollow_button(self.rule_config_btn, "#3498db", font_size="16px", padding="10px 20px")
        self.rule_config_btn.clicked.connect(self.show_rule_config_dialog)
        self.rule_config_btn.setVisible(False)  # 初始隐藏
        
        mode_layout.addWidget(mode_label)
        mode_layout.addWidget(self.mode_combo)
        mode_layout.addWidget(self.rule_config_btn)
        mode_layout.addStretch()
        file_layout.addLayout(mode_layout)
        
        # 模式说明标签
        self.mode_tip_label = QLabel("💡 交互式脱敏：选中文本或单元格后右键选择脱敏，精确控制每个内容")
        self.mode_tip_label.setStyleSheet("color: #27ae60; font-size: 9pt; background-color: #d5f4e6; padding: 8px; border-radius: 5px; border-left: 3px solid #27ae60;")
        self.mode_tip_label.setWordWrap(True)
        file_layout.addWidget(self.mode_tip_label)
        
        # 文件选择按钮
        file_btn_layout = QHBoxLayout()
        self.input_btn = QPushButton("📂 选择待脱敏文件")
        self.input_btn.setMinimumHeight(40)
        self.set_hollow_button(self.input_btn, "#3498db", font_size="16px", padding="10px 20px")
        self.input_btn.clicked.connect(self.select_input_file)
        
        self.output_btn = QPushButton("💾 设置输出路径")
        self.output_btn.setMinimumHeight(40)
        self.set_hollow_button(self.output_btn, "#3498db", font_size="16px", padding="10px 20px")
        self.output_btn.clicked.connect(self.select_output_path)
        file_btn_layout.addWidget(self.input_btn)
        file_btn_layout.addWidget(self.output_btn)
        file_layout.addLayout(file_btn_layout)
        # 文件信息显示
        self.file_info_label = QLabel("📄 未选择文件")
        self.file_info_label.setStyleSheet(
            "color: #1E3A8A; "
            "font-size: 10pt; "
            "padding: 5px; "
            "background-color: #f8f9fa; "
            "border-radius: 3px;"
        )
        file_layout.addWidget(self.file_info_label)
        # 内容交互区
        self.content_tabs = QTabWidget()
        self.content_tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #ecf0f1;
                padding: 8px 16px;
                margin-right: 2px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }
            QTabBar::tab:selected {
                background-color: #3498db;
                color: white;
            }
        """)

        self.text_tab = QWidget()
        self.excel_tab = QWidget()
        self.word_tab = QWidget()
        self.pdf_tab = QWidget()

        # 文本内容交互
        text_layout = QVBoxLayout()
        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(False)  # 允许编辑以支持交互式脱敏
        # QTextEdit用独立的预览提示框代替setPlaceholderText，保持与Excel风格统一
        text_placeholder = QLabel("📄 选择文本文件后，内容将在此显示\n💡 使用技巧：\n• 选中需要脱敏的文字后右键选择脱敏方式\n• 支持局部脱敏和全文同内容脱敏")
        text_placeholder.setAlignment(Qt.AlignmentFlag.AlignLeft)
        text_placeholder.setStyleSheet("color: #2563EB; font-size: 9pt; background-color: #f8f9fa; border: 1px solid #e9ecef; border-radius: 4px; padding: 10px;")
        text_layout.addWidget(text_placeholder)

        # 添加文本选择上下文菜单
        self.text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.text_edit.customContextMenuRequested.connect(self.show_text_context_menu)

        text_layout.addWidget(self.text_edit)
        
        # 添加清除按钮
        text_clear_btn = QPushButton("🗑️ 清除内容")
        self.set_hollow_button(text_clear_btn, "#e74c3c", font_size="14px", padding="6px 12px")
        text_clear_btn.clicked.connect(lambda: self.text_edit.setPlainText(""))
        text_layout.addWidget(text_clear_btn, alignment=Qt.AlignmentFlag.AlignRight)
        
        # 初始化文本右键菜单
        self.text_menu = QMenu(self)
        self.redact_action = QAction("🎯 标记脱敏（仅选中部分）", self)
        self.redact_action.triggered.connect(self.mark_text_redaction)
        self.text_menu.addAction(self.redact_action)
        
        self.redact_all_action = QAction("🔄 标记脱敏（全文相同内容）", self)
        self.redact_all_action.triggered.connect(self.mark_text_redaction_all)
        self.text_menu.addAction(self.redact_all_action)
        
        # 添加撤销脱敏功能
        self.text_menu.addSeparator()
        self.text_undo_action = QAction("↩️ 撤销脱敏", self)
        self.text_undo_action.triggered.connect(self.undo_text_redaction)
        self.text_menu.addAction(self.text_undo_action)
        
        # 区域撤销功能已移除，仅保留单步撤销
        
        # Excel内容交互
        excel_layout = QVBoxLayout()
        self.table_widget = QTableWidget()
        # QTableWidget没有setPlaceholderText方法，我们用标签代替
        excel_placeholder = QLabel("📊 选择Excel文件后，内容将在此显示\n💡 使用技巧：\n• 点击单元格选中后右键选择脱敏方式\n• 支持单元格、整行、整列脱敏 ")
        excel_placeholder.setAlignment(Qt.AlignmentFlag.AlignLeft)
        excel_placeholder.setStyleSheet("color: #2563EB; font-size: 9pt; background-color: #f8f9fa; border: 1px solid #e9ecef; border-radius: 4px; padding: 10px;")
        excel_layout.addWidget(excel_placeholder)
        excel_layout.addWidget(self.table_widget)
        
        # 添加清除按钮
        excel_clear_btn = QPushButton("🗑️ 清除内容")
        self.set_hollow_button(excel_clear_btn, "#e74c3c", font_size="14px", padding="6px 12px")
        excel_clear_btn.clicked.connect(self.table_widget.clear)
        excel_layout.addWidget(excel_clear_btn, alignment=Qt.AlignmentFlag.AlignRight)
        
        # Word文档内容交互
        word_layout = QVBoxLayout()
        self.word_edit = QTextEdit()
        self.word_edit.setReadOnly(False)  # 允许编辑以支持交互式脱敏
        # QTextEdit用独立的预览提示框代替setPlaceholderText，保持与Excel风格统一
        word_placeholder = QLabel("📝 选择Word文档后，内容将在此显示\n💡 使用技巧：\n• 选中需要脱敏的文字后右键选择脱敏方式\n• 支持局部脱敏和全文同内容脱敏 ")
        word_placeholder.setAlignment(Qt.AlignmentFlag.AlignLeft)
        word_placeholder.setStyleSheet("color: #2563EB; font-size: 9pt; background-color: #f8f9fa; border: 1px solid #e9ecef; border-radius: 4px; padding: 10px;")
        word_layout.addWidget(word_placeholder)
        
        # 添加Word文档选择上下文菜单
        self.word_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.word_edit.customContextMenuRequested.connect(self.show_word_context_menu)
        
        word_layout.addWidget(self.word_edit)
        
        # 添加清除按钮
        word_clear_btn = QPushButton("🗑️ 清除内容")
        self.set_hollow_button(word_clear_btn, "#e74c3c", font_size="14px", padding="6px 12px")
        word_clear_btn.clicked.connect(lambda: self.word_edit.setPlainText(""))
        word_layout.addWidget(word_clear_btn, alignment=Qt.AlignmentFlag.AlignRight)
        
        # 初始化Word右键菜单
        self.word_menu = QMenu(self)
        self.word_redact_action = QAction("🎯 标记脱敏（仅选中部分）", self)
        self.word_redact_action.triggered.connect(self.mark_word_redaction)
        self.word_menu.addAction(self.word_redact_action)
        
        self.word_redact_all_action = QAction("🔄 标记脱敏（全文相同内容）", self)
        self.word_redact_all_action.triggered.connect(self.mark_word_redaction_all)
        self.word_menu.addAction(self.word_redact_all_action)
        
        # 添加撤销脱敏功能
        self.word_menu.addSeparator()
        self.word_undo_action = QAction("↩️ 撤销脱敏", self)
        self.word_undo_action.triggered.connect(self.undo_word_redaction)
        self.word_menu.addAction(self.word_undo_action)
        
        # 区域撤销功能已移除，仅保留单步撤销
        
        # PDF文档内容交互
        pdf_layout = QVBoxLayout()
        self.pdf_edit = QTextEdit()
        self.pdf_edit.setReadOnly(False)  # 允许编辑以支持交互式脱敏
        # PDF标签页的提示信息
        pdf_placeholder = QLabel("📄 选择PDF文档后，内容将在此显示\n💡 使用技巧：\n• 选中需要脱敏的文字后右键选择脱敏方式\n• 支持局部脱敏和全文同内容脱敏\n• PDF处理")
        pdf_placeholder.setAlignment(Qt.AlignmentFlag.AlignLeft)
        pdf_placeholder.setStyleSheet("color: #2563EB; font-size: 9pt; background-color: #f8f9fa; border: 1px solid #e9ecef; border-radius: 4px; padding: 10px;")
        pdf_layout.addWidget(pdf_placeholder)
        
        # 添加PDF文档选择上下文菜单
        self.pdf_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.pdf_edit.customContextMenuRequested.connect(self.show_pdf_context_menu)
        
        pdf_layout.addWidget(self.pdf_edit)
        
        # 添加清除按钮
        pdf_clear_btn = QPushButton("🗑️ 清除内容")
        self.set_hollow_button(pdf_clear_btn, "#e74c3c", font_size="14px", padding="6px 12px")
        pdf_clear_btn.clicked.connect(lambda: self.pdf_edit.setPlainText(""))
        pdf_layout.addWidget(pdf_clear_btn, alignment=Qt.AlignmentFlag.AlignRight)
        
        # 初始化PDF右键菜单
        self.pdf_menu = QMenu(self)
        self.pdf_redact_action = QAction("🎯 标记脱敏（仅选中部分）", self)
        self.pdf_redact_action.triggered.connect(self.mark_pdf_redaction)
        self.pdf_menu.addAction(self.pdf_redact_action)
        
        self.pdf_redact_all_action = QAction("🔄 标记脱敏（全文相同内容）", self)
        self.pdf_redact_all_action.triggered.connect(self.mark_pdf_redaction_all)
        self.pdf_menu.addAction(self.pdf_redact_all_action)
        
        # 添加撤销脱敏功能
        self.pdf_menu.addSeparator()
        self.pdf_undo_action = QAction("↩️ 撤销脱敏", self)
        self.pdf_undo_action.triggered.connect(self.undo_pdf_redaction)
        self.pdf_menu.addAction(self.pdf_undo_action)
        
        self.text_tab.setLayout(text_layout)
        self.excel_tab.setLayout(excel_layout)
        self.word_tab.setLayout(word_layout)
        self.pdf_tab.setLayout(pdf_layout)
        self.content_tabs.addTab(self.word_tab, "📝 Word文档")
        self.content_tabs.addTab(self.pdf_tab, "📄 PDF文档")
        self.content_tabs.addTab(self.excel_tab, "📊 Excel内容")
        self.content_tabs.addTab(self.text_tab, "📄 文本内容")
        file_layout.addWidget(self.content_tabs)
        
        # 初始化表格右键菜单
        self.setup_table_context_menu()
        
        file_group.setLayout(file_layout)
        main_layout.addWidget(file_group)

        # 操作按钮区
        action_btn_layout = QHBoxLayout()
        self.process_btn = QPushButton("🚀 开始脱敏")
        self.process_btn.setMinimumHeight(50)
        self.set_hollow_button(self.process_btn, "#3498db", font_size="17px", padding="12px 24px", radius=8)
        
        self.batch_btn = QPushButton("📦 批量处理")
        self.batch_btn.setMinimumHeight(50)
        self.set_hollow_button(self.batch_btn, "#3498db", font_size="17px", padding="12px 24px", radius=8)
        self.batch_btn.setVisible(False)  # 初始隐藏，只在自动规则模式下显示
        
        self.help_btn = QPushButton("❓ 帮助")
        self.help_btn.setMinimumHeight(50)
        self.set_hollow_button(self.help_btn, "#3498db", font_size="17px", padding="12px 24px", radius=8)
        
        action_btn_layout.addWidget(self.process_btn)
        action_btn_layout.addWidget(self.batch_btn)
        action_btn_layout.addWidget(self.help_btn)
        main_layout.addLayout(action_btn_layout)
        
        # 连接所有按钮
        self.process_btn.clicked.connect(self.process_file)
        self.batch_btn.clicked.connect(self.batch_process)
        self.help_btn.clicked.connect(self.show_help)

        # 进度条和状态栏
        progress_layout = QHBoxLayout()
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #3498db;
                border-radius: 3px;
            }
        """)
        self.status_label = QLabel("✅ 就绪")
        self.status_label.setStyleSheet("color: #27ae60; font-weight: bold;")
        progress_layout.addWidget(self.progress_bar)
        progress_layout.addWidget(self.status_label)
        main_layout.addLayout(progress_layout)

        main_widget.setLayout(main_layout)
        self.setCentralWidget(main_widget)
        
        # 初始化规则引擎（隐藏在后台）
        self.rule_engine = RuleEngine()
        self.rule_engine.load_default_rules()
        
        # 验证规则示例（仅在开发调试时使用）
        # self.rule_engine.verify_rule_examples()

    def show_help(self):
        # 更新记录（简洁版，每条一句话，保留日期）
        update_records = [
            "2025-11-28 V2.4.1：修复word导出错误；新增预览区清除内容按钮，优化用户体验",
            "2025-11-23 V2.4.0：新增自定义规则生成器，支持快速创建个性化脱敏规则",
            "2025-11-06 V2.3.1：修改按钮样式",
            "2025-10-07 V2.3.0：新增PDF脱敏功能，支持文本型PDF脱敏",
            "2025-08-28 V2.2.1：修复日志导出错误",
            "2025-08-22 V2.2.0：增加日志导出功能",
            "2025-08-11 V2.1.5：增加Excel区域撤销功能，提升一致性和稳定性。",
            "2025-08-11 V2.1.4：Excel批量脱敏和全表查找替换功能增强。",
            "2025-08-11 V2.1.3：新增检查更新功能，支持自动获取最新版本",
            "2025-08-11 V2.1.2：新增自定义字段规则，界面优化",
            "2025-08-11 V2.1.1：修复姓名脱敏逻辑，批量粘贴更智能",
            "2025-08-10 V2.1.0：规则选择弹窗升级，界面统一化",
            "2025-08-09 V2.0.0：界面重构，支持Word/Excel批量处理",
            "2025-08-08 V1.9.3：Excel交互增强，视觉优化",
            "2025-08-04 V1.0：初始版本，基础UI与文件操作"
        ]
        help_text = f"""
        <div style='max-width:900px; margin:auto; font-family:Arial;'>
        <div style='color: #c0392b; font-weight: bold; border: 10px solid #c0392b; padding: 20px; margin-bottom: 25px; font-size:20px;'>
        【免责声明】本软件为免费工具，用户自愿使用。开发者不承诺软件绝对安全，对因使用软件导致的数据丢失、系统损坏等后果不承担责任。禁止将软件用于非法目的。
        </div>

        <h2 style='text-align:center; font-size:28px; margin-bottom:18px;'>通用脱敏工具 V{self.version} 使用说明</h2>        <h3 style='color:#2980b9; font-size:25px;'>基本功能</h3>
        <ul style='font-size:18px;'>
            <li>支持 TXT文本、Excel表格、Word文档、PDF文档四种格式的敏感信息脱敏处理</li>
            <li>交互式脱敏：选中文本或单元格，右键标记，精确控制每个内容</li>
            <li>自动脱敏（规则模式）：配置规则后可一键批量处理文件夹或多文件</li>
        </ul>

        <h3 style='color:#2980b9; font-size:25px;'>核心特色</h3>
        <ul style='font-size:18px;'>
            <li>右键快速标记，支持全文同步脱敏</li>
            <li>Excel支持单元格、整行、整列精确脱敏</li>
            <li>自定义规则每日自动保存，支持批量导入/导出</li>
            <li>内置十余种脱敏规则，涵盖生活工作多方面需求</li>
            <li>提供英文脱敏规则文件，可直接导入使用</li>
        </ul>

        <h3 style='color:#2980b9; font-size:25px;'>操作步骤</h3>
        <ol style='font-size:18px;'>
            <li>选择文件或文件夹：支持单文件、多文件或文件夹批量处理</li>
            <li>设置输出路径：可自定义输出目录，自动生成“文件名（脱敏）”格式</li>
            <li>配置脱敏规则：可自定义姓名、字段等规则，支持实时预览</li>
            <li>开始处理：点击“开始脱敏”或“批量处理”按钮自动完成脱敏</li>
        </ol>

        <div style='display:flex; align-items:center;'>
            <h3 style='color:#2980b9; font-size:25px; margin:0;'>更新记录</h3>
        </div>
        <div style='max-height:160px; overflow-y:auto; background:#f8f9fa; border:1px solid #e9ecef; border-radius:6px; padding:8px; font-size:18px;'>
        <ul style='margin:0;'>
            {''.join([f'<li>{rec}</li>' for rec in update_records])}
        </ul>
        </div>

        <div style='background-color:#f0f8ff; padding:10px; margin:14px 0; border-left:4px solid #4a90e2; font-size:20px;'>
        <b>使用建议</b><br>
        • 新手建议先用交互式模式熟悉功能<br>
        • 批量处理前建议先单个文件测试<br>
        • 重要文件处理前请务必备份
        </div>

        <p style='text-align:center; margin-top:18px; font-size:20px;'>
        <b>版本 V{self.version}</b> 基于 Apache License 2.0 | 2025 D&Ai <br>
        <b>更多功能请在使用中探索发现 😊</b>
        </p>
        </div>
        """

        # 使用 QDialog + QScrollArea + QLabel 实现带滚动条的帮助窗口
        help_window = QDialog(self)
        help_window.setWindowTitle("帮助")
        help_window.setWindowModality(Qt.WindowModality.ApplicationModal)
        help_window.resize(900, 600)

        # 主布局
        main_layout = QVBoxLayout(help_window)
        
        # 顶部按钮区域
        top_layout = QHBoxLayout()
        top_layout.addStretch()  # 左侧弹簧
        
        check_update_btn = QPushButton("检查更新")
        self.set_hollow_button(check_update_btn, "#4a90e2", font_size="14px", padding="8px 20px")
        check_update_btn.clicked.connect(self.check_update)
        top_layout.addWidget(check_update_btn)
        
        main_layout.addLayout(top_layout)

        scroll_area = QScrollArea(help_window)
        scroll_area.setWidgetResizable(True)

        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)

        help_label = QLabel(help_text)
        help_label.setWordWrap(True)
        help_label.setStyleSheet("font-size: 14px; padding: 10px;")
        help_label.setTextFormat(Qt.TextFormat.RichText)
        content_layout.addWidget(help_label)

        scroll_area.setWidget(content_widget)

        main_layout.addWidget(scroll_area)
        help_window.setLayout(main_layout)

        help_window.exec_()

    def update_rule_list(self):
        """更新规则总览（弹窗显示，避免主窗口属性缺失）"""
        total_rules = len(self.rule_engine.rules)
        active_rules_list = self.rule_engine.get_active_rules()
        active_count = len(active_rules_list)
        if active_count == 0:
            msg = f"当前未激活任何规则，共 {total_rules} 条"
        else:
            names = '\n'.join([f"{i+1}. {r.name}" for i, r in enumerate(active_rules_list)])
            msg = f"已激活 {active_count} 条规则，共 {total_rules} 条：\n\n{names}"
        QMessageBox.information(self, "规则设置", msg)

    def add_rule(self):
        """添加新规则"""
        try:
            rule_data = self.build_custom_rule_data()
        except ValueError as err:
            QMessageBox.warning(self, "提示", str(err))
            return
        except Exception as exc:
            QMessageBox.critical(self, "错误", f"生成规则失败: {str(exc)}")
            return

        try:
            new_rule = RedactionRule(
                rule_id=rule_data["rule_id"],
                name=rule_data["name"],
                pattern=rule_data["pattern"],
                example=rule_data["example_display"],
                regex=rule_data["pattern"],
                marker_char="*",
                is_active=True,
            )
            self.rule_engine.add_rule(new_rule)
            self.update_rule_list()
            QMessageBox.information(self, "成功", f"已添加规则：{new_rule.name}")
            self.refresh_custom_rule_preview()
        except Exception as exc:
            QMessageBox.critical(self, "错误", f"添加规则失败: {str(exc)}")

    def clear_rules(self):
        """清空所有规则"""
        reply = QMessageBox.question(
            self, 
            "确认", 
            "确定要清空所有规则吗?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.rule_engine.rules.clear()
            self.update_rule_list()

    def export_rules(self):
        """导出规则到JSON文件"""
        if not self.rule_engine.rules:
            QMessageBox.warning(self, "警告", "没有可导出的规则")
            return
        
        # 默认保存到 user_custom_rules 文件夹
        import os, datetime
        rules_dir = os.path.join(os.path.dirname(__file__), "user_custom_rules")
        if not os.path.exists(rules_dir):
            os.makedirs(rules_dir)
        
        date_str = datetime.datetime.now().strftime("%Y%m%d")
        default_filename = os.path.join(rules_dir, f"导出规则{date_str}.json")
            
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出规则",
            default_filename,
            "JSON文件 (*.json);;所有文件 (*)"
        )
        
        if file_path:
            try:
                import json
                rules_data = [{
                    'rule_id': rule.rule_id,
                    'name': rule.name,
                    'pattern': rule.pattern,
                    'example': rule.example,
                    'is_active': rule.is_active
                } for rule in self.rule_engine.rules]
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(rules_data, f, ensure_ascii=False, indent=2)
                    
                QMessageBox.information(self, "成功", "规则导出成功")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")

    def import_rules(self):
        """从JSON文件导入规则"""
        # 默认打开 user_custom_rules 文件夹
        import os
        rules_dir = os.path.join(os.path.dirname(__file__), "user_custom_rules")
        if not os.path.exists(rules_dir):
            os.makedirs(rules_dir)
        
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "导入规则",
            rules_dir,
            "JSON文件 (*.json);;所有文件 (*)"
        )
        
        if file_path:
            try:
                import json
                with open(file_path, 'r', encoding='utf-8') as f:
                    rules_data = json.load(f)
                
                # 检查是否是数组格式
                if not isinstance(rules_data, list):
                    # 如果是对象格式，检查是否有rules字段
                    if isinstance(rules_data, dict) and 'rules' in rules_data:
                        rules_data = rules_data['rules']
                    else:
                        raise ValueError("JSON格式不正确，应为规则数组或包含rules字段的对象")
                
                # 询问用户是否要替换现有规则还是添加到现有规则
                reply = QMessageBox.question(
                    self, 
                    "导入选项", 
                    f"检测到 {len(rules_data)} 条规则\n\n是否要替换现有的内置规则？\n\n点击 '是' 替换所有规则\n点击 '否' 添加到现有规则中\n点击 '取消' 取消导入",
                    QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
                )
                
                if reply == QMessageBox.Cancel:
                    return
                
                new_rules = []
                for rule_data in rules_data:
                    new_rule = RedactionRule(
                        rule_id=rule_data.get('id', rule_data.get('rule_id', f"imported_{len(new_rules)+1}")),
                        name=rule_data.get('name', '导入规则'),
                        pattern=rule_data.get('pattern', ''),
                        example=rule_data.get('example', ''),
                        regex=rule_data.get('pattern', ''),
                        marker_char="*",
                        is_active=rule_data.get('is_active', True)
                    )
                    new_rules.append(new_rule)
                
                if reply == QMessageBox.Yes:
                    # 替换现有规则
                    self.rule_engine.rules = new_rules
                    action = "替换"
                else:
                    # 添加到现有规则
                    self.rule_engine.rules.extend(new_rules)
                    action = "添加"
                
                self.update_rule_list()
                active_count = len([r for r in new_rules if r.is_active])
                QMessageBox.information(
                    self, 
                    "成功", 
                    f"成功{action} {len(new_rules)} 条规则\n其中 {active_count} 条已激活\n当前总规则数: {len(self.rule_engine.rules)}"
                )
                
            except Exception as e:
                QMessageBox.critical(self, "错误", f"导入失败: {str(e)}")

    def edit_rule(self):
        """编辑现有规则"""
        if not self.rule_engine.rules:
            QMessageBox.warning(self, "警告", "没有可编辑的规则")
            return
            
        # 创建规则选择对话框
        rule_names = [f"{i+1}. {rule.name} ({'✅' if rule.is_active else '❌'})" for i, rule in enumerate(self.rule_engine.rules)]
        rule_name, ok = QInputDialog.getItem(
            self,
            "选择规则",
            "请选择要编辑的规则:",
            rule_names,
            0,
            False
        )
        
        if ok and rule_name:
            try:
                # 获取选中的规则索引
                rule_index = int(rule_name.split('.')[0]) - 1
                selected_rule = self.rule_engine.rules[rule_index]
                
                # 创建编辑对话框
                dialog = QDialog(self)
                dialog.setWindowTitle("编辑规则")
                dialog.setMinimumWidth(500)
                layout = QVBoxLayout()
                
                # 规则名称
                name_label = QLabel("规则名称:")
                name_edit = QLineEdit(selected_rule.name)
                layout.addWidget(name_label)
                layout.addWidget(name_edit)
                
                # 规则模式
                pattern_label = QLabel("匹配模式:")
                pattern_edit = QLineEdit(selected_rule.pattern)
                layout.addWidget(pattern_label)
                layout.addWidget(pattern_edit)
                
                # 示例
                example_label = QLabel("示例:")
                example_edit = QLineEdit(selected_rule.example)
                layout.addWidget(example_label)
                layout.addWidget(example_edit)
                
                # 规则状态区域
                status_layout = QHBoxLayout()
                status_label = QLabel("规则状态:")
                status_layout.addWidget(status_label)
                
                # 激活/禁用按钮
                toggle_btn = QPushButton()
                def apply_toggle_style():
                    if selected_rule.is_active:
                        toggle_btn.setText("禁用规则")
                        self.set_hollow_button(toggle_btn, "#ff6b6b", padding="6px 12px")
                    else:
                        toggle_btn.setText("激活规则")
                        self.set_hollow_button(toggle_btn, "#51cf66", padding="6px 12px")

                apply_toggle_style()

                def toggle_rule_status():
                    selected_rule.is_active = not selected_rule.is_active
                    apply_toggle_style()
                
                toggle_btn.clicked.connect(toggle_rule_status)
                status_layout.addWidget(toggle_btn)
                status_layout.addStretch()
                
                layout.addLayout(status_layout)
                
                # 按钮
                btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
                btn_box.accepted.connect(dialog.accept)
                btn_box.rejected.connect(dialog.reject)
                layout.addWidget(btn_box)
                
                dialog.setLayout(layout)
                
                if dialog.exec_() == QDialog.Accepted:
                    # 更新规则
                    selected_rule.name = name_edit.text()
                    selected_rule.pattern = pattern_edit.text()
                    selected_rule.regex = pattern_edit.text()  # 同时更新regex
                    selected_rule.example = example_edit.text()
                    self.update_rule_list()
                    QMessageBox.information(self, "成功", "规则更新成功")
                    
            except Exception as e:
                QMessageBox.critical(self, "错误", f"编辑规则失败: {str(e)}")

    def preview_rule(self):
        """预览规则效果"""
        if not self.rule_engine.rules:
            QMessageBox.warning(self, "警告", "没有可预览的规则")
            return
            
        # 创建预览对话框
        dialog = QDialog(self)
        dialog.setWindowTitle("规则预览")
        dialog.resize(500, 400)
        layout = QVBoxLayout()
        
        # 规则选择下拉框
        rule_combo = QComboBox()
        rule_combo.addItems([rule.name for rule in self.rule_engine.rules])
        layout.addWidget(QLabel("选择规则:"))
        layout.addWidget(rule_combo)
        
        # 测试文本输入
        test_text_edit = QTextEdit()
        test_text_edit.setPlaceholderText("输入测试文本...")
        layout.addWidget(QLabel("测试文本:"))
        layout.addWidget(test_text_edit)
        
        # 预览结果
        result_label = QLabel("预览结果:")
        result_text = QTextEdit()
        result_text.setReadOnly(True)
        layout.addWidget(result_label)
        layout.addWidget(result_text)
        
        # 预览按钮
        preview_btn = QPushButton("预览")
        def on_preview():
            selected_index = rule_combo.currentIndex()
            if selected_index >= 0:
                rule = self.rule_engine.rules[selected_index]
                test_text = test_text_edit.toPlainText()
                if not test_text:
                    result_text.setPlainText("请输入测试文本")
                    return
                    
                # 使用与交互脱敏相同的内置算法
                try:
                    import re
                    # 使用规则引擎的模式匹配，但结合内置算法
                    # 使用 re.finditer 来获取完整匹配，避免分组问题
                    matches = []
                    for match_obj in re.finditer(rule.pattern, test_text):
                        matches.append(match_obj.group(0))  # group(0) 是完整匹配
                    
                    if not matches:
                        result = test_text
                        matches_found = False
                    else:
                        result = test_text
                        matches_found = True
                        for match in matches:
                            # 使用smart_redact_text的内置算法
                            redacted = self.smart_redact_text(match)
                            result = result.replace(match, redacted)
                    
                    # 显示详细结果
                    status = "✅ 匹配成功" if matches_found else "❌ 未匹配"
                    match_info = f"匹配到的内容: {matches}" if matches_found else "无匹配内容"
                    result_text.setPlainText(f"""应用规则: {rule.name}
规则描述: {rule.example}
匹配状态: {status}
{match_info}

原始文本: {test_text}

脱敏结果: {result}

注：此预览使用与实际脱敏相同的算法""")
                    
                except Exception as e:
                    result_text.setPlainText(f"预览出错: {str(e)}\n\n请检查规则格式是否正确")
                    
        preview_btn.clicked.connect(on_preview)
        layout.addWidget(preview_btn)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def read_file_with_encoding(self, file_path):
        """尝试多种编码格式读取文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig', 'latin1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    self.original_encoding = encoding  # 保存原始编码
                    self.status_label.setText(f"使用 {encoding} 编码成功读取文件")
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.status_label.setText(f"读取文件时发生错误：{str(e)}")
                continue
        
        return None

    def load_word_document(self, file_path):
        """加载Word文档内容"""
        try:
            if file_path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                full_text.append(cell.text)
                    self.current_word_doc = doc
                    return '\n'.join(full_text)
                except ImportError:
                    QMessageBox.warning(self, "警告", "未安装python-docx库，无法处理DOCX文件\n请运行: pip install python-docx")
                    return None
            else:
                QMessageBox.warning(self, "警告", "仅支持DOCX格式文件，请先转换DOC为DOCX后再处理。")
                return None
        except Exception as e:
            QMessageBox.warning(self, "警告", f"加载Word文档失败: {str(e)}")
            return None

    def on_mode_changed(self):
        """处理模式切换"""
        if self.mode_combo.currentIndex() == 0:
            # 交互式脱敏模式
            self.mode_tip_label.setText("💡 交互式脱敏：选中文本或单元格后右键选择脱敏，精确控制每个内容")
            self.mode_tip_label.setStyleSheet("color: #27ae60; font-size: 9pt; background-color: #d5f4e6; padding: 8px; border-radius: 5px; border-left: 3px solid #27ae60;")
            self.rule_config_btn.setVisible(False)  # 隐藏规则配置按钮
            self.batch_btn.setVisible(False)  # 隐藏批量处理按钮
        else:
            # 自动脱敏模式
            self.mode_tip_label.setText("⚙️ 自动脱敏：将对整个文件应用脱敏规则，请点击【配置脱敏规则】按钮设置规则")
            self.mode_tip_label.setStyleSheet("color: #e74c3c; font-size: 9pt; background-color: #fdf2f2; padding: 8px; border-radius: 5px; border-left: 3px solid #e74c3c;")
            self.rule_config_btn.setVisible(True)  # 显示规则配置按钮
            self.batch_btn.setVisible(True)  # 显示批量处理按钮

    def show_rule_config_dialog(self):
        """显示规则配置弹窗（菜单式复选框，每条规则可勾选激活/禁用）"""
        dialog = QDialog(self)
        dialog.setWindowTitle("📋 脱敏规则配置（可视化选择）")
        dialog.setModal(True)
        dialog.resize(700, 600)

        layout = QVBoxLayout()

        # 规则复选框列表区
        rules_group = QGroupBox("内置规则选择（勾选表示激活）")
        rules_layout = QVBoxLayout()
        self.rule_checkboxes = []
        for i, rule in enumerate(self.rule_engine.rules):
            cb = QCheckBox(f"{i+1}. {rule.name}  ——  {rule.example}")
            cb.setChecked(rule.is_active)
            cb.setToolTip(f"匹配规则: {rule.pattern}")
            self.rule_checkboxes.append(cb)
            row_layout = QHBoxLayout()
            row_layout.addWidget(cb)
            # 如果是姓名规则，在后面加自定义按钮
            if rule.name == "姓名":
                name_btn = QPushButton("自定义名单")
                self.set_hollow_button(name_btn, "#3498db", font_size="12px", padding="6px 12px")
                name_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                name_btn.clicked.connect(self.show_name_redact_dialog)
                row_layout.addSpacing(10)
                row_layout.addWidget(name_btn)
            # 如果是自定义字段规则，在后面加自定义按钮
            elif rule.name == "自定义字段":
                field_btn = QPushButton("自定义字段")
                self.set_hollow_button(field_btn, "#e67e22", font_size="12px", padding="6px 12px")
                field_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                field_btn.clicked.connect(self.show_custom_field_redact_dialog)
                row_layout.addSpacing(10)
                row_layout.addWidget(field_btn)
            rules_layout.addLayout(row_layout)
        rules_group.setLayout(rules_layout)
        layout.addWidget(rules_group)

        # 自定义规则生成器（表单+实时预览）
        custom_group = QGroupBox("自定义规则生成器")
        custom_layout = QVBoxLayout()

        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignmentFlag.AlignLeft)

        self.custom_field_name_input = QLineEdit()
        self.custom_field_name_input.setPlaceholderText("例如：客户姓名")
        self.custom_field_name_input.setText("客户姓名")
        self.custom_field_name_input.setToolTip("填写这条自定义规则的名称，便于在规则列表中识别，例如“客户姓名”或“培训时长”。")
        form_layout.addRow("字段名称：", self.custom_field_name_input)

        self.custom_match_type_combo = QComboBox()
        match_items = [
            ("完全自定义", "custom"),
            ("仅字母", "alpha"),
            ("仅数字", "digit"),
            ("字母+数字", "alnum"),
            ("仅汉字", "han"),
            ("字母数字汉字", "mixed"),
        ]
        for label_text, value in match_items:
            self.custom_match_type_combo.addItem(label_text, userData=value)
        self.custom_match_type_combo.setCurrentIndex(4)
        self.custom_match_type_combo.setToolTip("选择待匹配文本的大致字符类型。若要指定更复杂的范围，请选择“完全自定义”。")
        form_layout.addRow("匹配方式：", self.custom_match_type_combo)

        self.custom_charset_label = QLabel("自定义字符集：")
        self.custom_charset_input = QLineEdit()
        self.custom_charset_input.setPlaceholderText("请输入字符集合，例如：A-Za-z0-9")
        self.custom_charset_input.setToolTip("仅在选择“完全自定义”时需要填写。支持连字符区间（如A-Z）或直接列出允许的字符，也可包含\\u4e00-\\u9fa5。")
        form_layout.addRow(self.custom_charset_label, self.custom_charset_input)
        self.custom_charset_label.setVisible(False)
        self.custom_charset_input.setVisible(False)

        self.custom_min_length_spin = QSpinBox()
        self.custom_min_length_spin.setRange(1, 100)
        self.custom_min_length_spin.setValue(2)
        self.custom_min_length_spin.setToolTip("匹配文本的最小长度。数字或字符个数不足该值将不会命中规则。")
        form_layout.addRow("最小长度：", self.custom_min_length_spin)

        self.custom_max_length_spin = QSpinBox()
        self.custom_max_length_spin.setRange(1, 100)
        self.custom_max_length_spin.setValue(10)
        self.custom_max_length_spin.setToolTip("匹配文本的最大长度。可与最小值相同以限定固定长度。")
        form_layout.addRow("最大长度：", self.custom_max_length_spin)

        self.custom_separator_input = QLineEdit()
        self.custom_separator_input.setPlaceholderText("例如：空格 / . / - / \\n")
        self.custom_separator_input.setToolTip("若需要多段内容（如身份证格式的分段），在此指定段与段之间的分隔符。支持输入\\n、\\t代表换行和制表符。")
        form_layout.addRow("分隔符：", self.custom_separator_input)

        self.custom_parts_spin = QSpinBox()
        self.custom_parts_spin.setRange(1, 5)
        self.custom_parts_spin.setValue(1)
        self.custom_parts_spin.setToolTip("设置需要匹配的段数。例如银行卡可分成多段；若仅匹配一段文本，请保持为1。")
        form_layout.addRow("段数：", self.custom_parts_spin)

        self.custom_template_combo = QComboBox()
        template_items = [
            ("首字母+星号", "first_asterisk"),
            ("全部星号", "all_asterisk"),
            ("保留前3位", "keep_3"),
            ("保留头尾", "keep_head_tail"),
        ]
        for label_text, value in template_items:
            self.custom_template_combo.addItem(label_text, userData=value)
        self.custom_template_combo.setToolTip("选择匹配到的数据在脱敏后的展示方式。不同模板会保留不同的关键信息。")
        form_layout.addRow("脱敏模板：", self.custom_template_combo)

        custom_layout.addLayout(form_layout)

        helper_layout = QHBoxLayout()
        helper_layout.addStretch()
        preview_button = QPushButton("刷新预览")
        self.set_hollow_button(preview_button, "#27ae60", font_size="12px", padding="6px 16px")
        helper_layout.addWidget(preview_button)
        copy_button = QPushButton("复制 JSON")
        self.set_hollow_button(copy_button, "#4a90e2", font_size="12px", padding="6px 16px")
        helper_layout.addWidget(copy_button)
        helper_layout.addStretch()
        custom_layout.addLayout(helper_layout)

        preview_group = QGroupBox("实时预览")
        preview_layout = QVBoxLayout()

        self.custom_regex_preview = QLabel("/")
        self.custom_regex_preview.setWordWrap(True)
        self.custom_regex_preview.setStyleSheet("color: #27ae60; font-family: Consolas, 'Courier New', monospace;")
        preview_layout.addWidget(QLabel("正则表达式："))
        preview_layout.addWidget(self.custom_regex_preview)

        self.custom_example_preview = QLabel("示例预览将在此显示")
        self.custom_example_preview.setWordWrap(True)
        preview_layout.addWidget(QLabel("示例："))
        preview_layout.addWidget(self.custom_example_preview)

        self.custom_json_preview = QTextEdit()
        self.custom_json_preview.setReadOnly(True)
        self.custom_json_preview.setMinimumHeight(140)
        self.custom_json_preview.setStyleSheet("font-family: Consolas, 'Courier New', monospace;")
        preview_layout.addWidget(QLabel("JSON："))
        preview_layout.addWidget(self.custom_json_preview)

        preview_group.setLayout(preview_layout)
        custom_layout.addWidget(preview_group)

        custom_group.setLayout(custom_layout)
        layout.addWidget(custom_group)

        def handle_match_type_change(index: int) -> None:
            value = self.custom_match_type_combo.itemData(index)
            is_custom = value == "custom"
            self.custom_charset_label.setVisible(is_custom)
            self.custom_charset_input.setVisible(is_custom)
            if not is_custom:
                self.custom_charset_input.clear()
            self.refresh_custom_rule_preview()

        self.custom_match_type_combo.currentIndexChanged.connect(handle_match_type_change)
        self.custom_field_name_input.textChanged.connect(self.refresh_custom_rule_preview)
        self.custom_charset_input.textChanged.connect(self.refresh_custom_rule_preview)
        self.custom_min_length_spin.valueChanged.connect(self.refresh_custom_rule_preview)
        self.custom_max_length_spin.valueChanged.connect(self.refresh_custom_rule_preview)
        self.custom_separator_input.textChanged.connect(self.refresh_custom_rule_preview)
        self.custom_parts_spin.valueChanged.connect(self.refresh_custom_rule_preview)
        self.custom_template_combo.currentIndexChanged.connect(self.refresh_custom_rule_preview)
        preview_button.clicked.connect(self.refresh_custom_rule_preview)
        copy_button.clicked.connect(self.copy_custom_rule_json)
        handle_match_type_change(self.custom_match_type_combo.currentIndex())

        # 规则操作按钮（保留原有功能）
        rule_btn_layout = QHBoxLayout()
        add_btn = QPushButton("添加规则")
        self.set_hollow_button(add_btn, "#3498db", font_size="14px", padding="8px 18px")
        add_btn.clicked.connect(self.add_rule)
        import_btn = QPushButton("导入规则")
        self.set_hollow_button(import_btn, "#3498db", font_size="14px", padding="8px 18px")
        import_btn.clicked.connect(self.import_rules)
        export_btn = QPushButton("导出规则")
        self.set_hollow_button(export_btn, "#3498db", font_size="14px", padding="8px 18px")
        export_btn.clicked.connect(self.export_rules)
        edit_btn = QPushButton("编辑规则")
        self.set_hollow_button(edit_btn, "#3498db", font_size="14px", padding="8px 18px")
        edit_btn.clicked.connect(self.edit_rule)
        preview_btn = QPushButton("预览规则")
        self.set_hollow_button(preview_btn, "#9b59b6", font_size="14px", padding="8px 18px")
        preview_btn.clicked.connect(self.preview_rule)
        clear_btn = QPushButton("清空规则")
        self.set_hollow_button(clear_btn, "#e74c3c", font_size="14px", padding="8px 18px")
        clear_btn.clicked.connect(self.clear_rules)
        rule_btn_layout.addWidget(add_btn)
        rule_btn_layout.addWidget(import_btn)
        rule_btn_layout.addWidget(export_btn)
        rule_btn_layout.addWidget(edit_btn)
        rule_btn_layout.addWidget(preview_btn)
        rule_btn_layout.addWidget(clear_btn)
        layout.addLayout(rule_btn_layout)

        # 对话框按钮
        dialog_btn_layout = QHBoxLayout()
        dialog_btn_layout.addStretch()
        
        ok_btn = QPushButton("继续")
        self.set_hollow_button(
            ok_btn,
            "#27ae60",
            font_size="16px",
            padding="12px 24px",
            radius=8,
            min_width="100px",
        )
        
        cancel_btn = QPushButton("取消")
        self.set_hollow_button(
            cancel_btn,
            "#e74c3c",
            font_size="16px",
            padding="12px 24px",
            radius=8,
            min_width="100px",
        )
        
        dialog_btn_layout.addWidget(ok_btn)
        dialog_btn_layout.addWidget(cancel_btn)
        dialog_btn_layout.addStretch()
        layout.addLayout(dialog_btn_layout)

        # 按钮事件：保存激活状态
        def on_ok():
            for cb, rule in zip(self.rule_checkboxes, self.rule_engine.rules):
                rule.is_active = cb.isChecked()
            active_count = len([r for r in self.rule_engine.rules if r.is_active])
            QMessageBox.information(self, "规则设置", f"已激活 {active_count} 条规则")
            self.update_rule_list()
            dialog.accept()
        ok_btn.clicked.connect(on_ok)
        cancel_btn.clicked.connect(dialog.reject)

        dialog.setLayout(layout)
        dialog.exec_()

    def build_custom_rule_data(self, preview_only: bool = False) -> dict:
        """根据表单输入生成规则配置与预览数据"""
        if not hasattr(self, "custom_field_name_input"):
            raise ValueError("自定义规则控件尚未初始化")

        field_name = self.custom_field_name_input.text().strip()
        if not field_name:
            raise ValueError("请填写字段名称")

        match_type = self.custom_match_type_combo.currentData() or "custom"
        min_length = self.custom_min_length_spin.value()
        max_length = self.custom_max_length_spin.value()
        if min_length > max_length:
            raise ValueError("最小长度不能大于最大长度")

        parts = self.custom_parts_spin.value()
        separator_raw = self.custom_separator_input.text()
        separator_processed = separator_raw.replace("\\n", "\n").replace("\\t", "\t")

        pattern_map = {
            "alpha": "A-Za-z",
            "digit": "0-9",
            "alnum": "A-Za-z0-9",
            "han": "\\u4e00-\\u9fa5",
            "mixed": "A-Za-z0-9\\u4e00-\\u9fa5",
        }

        if match_type == "custom":
            custom_charset = re.sub(r"\s+", "", self.custom_charset_input.text())
            if not custom_charset:
                raise ValueError("请填写自定义字符集")
            pattern_charset = custom_charset.replace("[", "").replace("]", "")
        else:
            pattern_charset = pattern_map.get(match_type, "A-Za-z0-9")

        pattern_charset = re.sub(r"\s+", "", pattern_charset)
        if not pattern_charset:
            raise ValueError("字符集不能为空")

        segment = f"[{pattern_charset}]{{{min_length},{max_length}}}"
        if separator_processed:
            if separator_raw == "\\n":
                separator_pattern = r"\n"
            elif separator_raw == "\\t":
                separator_pattern = r"\t"
            else:
                separator_pattern = re.escape(separator_processed)
        else:
            separator_pattern = ""

        if parts == 1:
            pattern = segment
        else:
            if separator_pattern:
                pattern = f"{segment}(?:{separator_pattern}{segment}){{{parts-1}}}"
            else:
                pattern = f"{segment}(?:{segment}){{{parts-1}}}"

        pool_source = pattern_charset if match_type != "custom" else self.custom_charset_input.text()
        char_pool = self._build_char_pool(match_type, pool_source)
        example_text = self._generate_custom_example(char_pool, parts, min_length, max_length, separator_processed)
        template_value = self.custom_template_combo.currentData() or "first_asterisk"
        masked_example = self._apply_mask_template(example_text, template_value)

        rule_id = "rule_preview" if preview_only else f"custom_{int(time.time())}_{random.randint(1000, 9999)}"

        return {
            "rule_id": rule_id,
            "name": field_name,
            "pattern": pattern,
            "regex_display": f"/{pattern}/",
            "example_display": f"{example_text} → {masked_example}",
            "payload": {
                "rule_id": rule_id,
                "name": field_name,
                "pattern": pattern,
                "example": f"{example_text} → {masked_example}",
                "regex": pattern,
                "marker_char": "*",
                "template": template_value,
                "is_active": True,
                "metadata": {
                    "match_type": match_type,
                    "min_length": min_length,
                    "max_length": max_length,
                    "parts": parts,
                    "separator": separator_raw,
                },
            },
            "json": json.dumps(
                {
                    "rule_id": rule_id,
                    "name": field_name,
                    "pattern": pattern,
                    "example": f"{example_text} → {masked_example}",
                    "regex": pattern,
                    "marker_char": "*",
                    "template": template_value,
                    "is_active": True,
                    "metadata": {
                        "match_type": match_type,
                        "min_length": min_length,
                        "max_length": max_length,
                        "parts": parts,
                        "separator": separator_raw,
                    },
                },
                ensure_ascii=False,
                indent=2,
            ),
        }

    def _build_char_pool(self, match_type: str, charset: str) -> str:
        """构建示例所需的字符集合"""
        base_map = {
            "alpha": string.ascii_letters,
            "digit": string.digits,
            "alnum": string.ascii_letters + string.digits,
            "han": "张王李赵刘陈杨黄周吴徐孙胡郭林何高马罗梁宋郑谢韩唐冯许曹",
            "mixed": string.ascii_letters + string.digits + "张王李赵刘陈杨黄周吴徐孙胡郭林何高马罗梁宋郑谢韩唐冯许曹",
        }

        if match_type in base_map:
            pool = base_map[match_type]
        else:
            pool = ""
            charset_clean = re.sub(r"\s+", "", charset or "")
            charset_clean = charset_clean.replace("[", "").replace("]", "")

            if "\\u4e00-\\u9fa5" in charset_clean:
                pool += "张王李赵刘陈杨黄周吴徐孙胡郭林何高马罗梁宋郑谢韩唐冯许曹"
                charset_clean = charset_clean.replace("\\u4e00-\\u9fa5", "")

            for start, end in re.findall(r"([A-Za-z0-9])\-([A-Za-z0-9])", charset_clean):
                if ord(start) <= ord(end):
                    for code in range(ord(start), ord(end) + 1):
                        pool += chr(code)
                charset_clean = charset_clean.replace(f"{start}-{end}", "")

            charset_clean = charset_clean.replace("\\", "")
            pool += charset_clean

        unique_pool = "".join(dict.fromkeys(pool))
        return unique_pool or (string.ascii_letters + string.digits)

    def _generate_custom_example(self, pool: str, parts: int, min_length: int, max_length: int, separator: str) -> str:
        """生成示例文本"""
        if not pool:
            raise ValueError("字符集为空，无法生成示例")

        segments = []
        for _ in range(parts):
            length = max(random.randint(min_length, max_length), 1)
            segment = ''.join(random.choice(pool) for _ in range(length))
            segments.append(segment)

        if separator:
            return separator.join(segments)
        return ''.join(segments)

    def _apply_mask_template(self, text: str, template: str) -> str:
        """根据模板生成脱敏示例"""
        if not text:
            return ""

        if template == "all_asterisk":
            return "*" * len(text)
        if template == "keep_3":
            return text[:3] + "*" * max(len(text) - 3, 0) if len(text) > 3 else "*" * len(text)
        if template == "keep_head_tail":
            if len(text) <= 2:
                return "*" * len(text)
            return text[0] + "*" * (len(text) - 2) + text[-1]

        parts = re.split(r"(\s+)", text)
        masked_parts = []
        for part in parts:
            if not part or part.isspace():
                masked_parts.append(part)
            else:
                masked_parts.append(part[0] + "*" * max(len(part) - 1, 0))
        return ''.join(masked_parts)

    def refresh_custom_rule_preview(self) -> None:
        """刷新右侧预览"""
        if not hasattr(self, "custom_regex_preview"):
            return
        try:
            preview = self.build_custom_rule_data(preview_only=True)
        except ValueError as err:
            self.custom_regex_preview.setText(str(err))
            self.custom_example_preview.setText("—")
            self.custom_json_preview.setPlainText("")
            return
        except Exception as exc:
            self.custom_regex_preview.setText(str(exc))
            self.custom_example_preview.setText("—")
            self.custom_json_preview.setPlainText("")
            return

        self.custom_regex_preview.setText(preview["regex_display"])
        self.custom_example_preview.setText(preview["example_display"])
        self.custom_json_preview.setPlainText(preview["json"])

    def copy_custom_rule_json(self) -> None:
        """复制预览中的JSON"""
        if not hasattr(self, "custom_json_preview"):
            return
        json_text = self.custom_json_preview.toPlainText().strip()
        if not json_text:
            QMessageBox.information(self, "提示", "当前没有可复制的JSON内容")
            return
        QApplication.clipboard().setText(json_text)
        QMessageBox.information(self, "成功", "JSON内容已复制到剪贴板")

    def sanitize_excel_value(self, value):
        """清理Excel单元格值，避免特殊字符问题"""
        if value is None:
            return ""
        
        # 转换为字符串
        str_value = str(value)
        
        # 替换可能导致问题的字符
        # Excel不允许某些字符，特别是控制字符
        forbidden_chars = ['\x00', '\x01', '\x02', '\x03', '\x04', '\x05', '\x06', '\x07', '\x08']
        for char in forbidden_chars:
            str_value = str_value.replace(char, '')
        
        # 限制长度（Excel单元格最大32767字符）
        if len(str_value) > 32767:
            str_value = str_value[:32767]
        
        return str_value

    def select_input_file(self):
        """选择输入文件并加载内容"""
        # 根据当前标签页确定文件类型过滤器
        current_tab = self.content_tabs.currentIndex()
        if current_tab == 0:  # Word标签页
            file_filter = "Word文档 (*.docx *.doc);;Word 2007及以上 (*.docx);;Word 97-2003 (*.doc);;所有文件 (*)"
            dialog_title = "选择Word文档"
        elif current_tab == 1:  # PDF标签页
            file_filter = "PDF文档 (*.pdf);;所有文件 (*)"
            dialog_title = "选择PDF文件"
        elif current_tab == 2:  # Excel标签页
            file_filter = "Excel文件 (*.xlsx);;所有文件 (*)"
            dialog_title = "选择Excel文件"
        elif current_tab == 3:  # 文本标签页
            file_filter = "文本文件 (*.txt);;所有文件 (*)"
            dialog_title = "选择文本文件"
        else:
            file_filter = "Word文档 (*.docx *.doc);;PDF文档 (*.pdf);;Excel文件 (*.xlsx);;文本文件 (*.txt);;所有文件 (*)"
            dialog_title = "选择输入文件"
            
        script_dir = os.path.dirname(os.path.abspath(__file__))
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            dialog_title,
            script_dir,
            file_filter
        )
        if file_path:
            self.input_file_path = file_path
            self.file_info_label.setText(f"输入文件: {os.path.basename(file_path)}")
            self.status_label.setText("已选择输入文件")
            
            # 根据文件类型加载内容
            try:
                if file_path.lower().endswith('.txt'):
                    # 尝试多种编码格式
                    content = self.read_file_with_encoding(file_path)
                    if content is not None:
                        self.text_edit.setPlainText(content)
                        self.content_tabs.setCurrentIndex(2)  # 切换到文本标签页
                    else:
                        QMessageBox.warning(self, "警告", "无法读取文件，请检查文件编码格式")
                        return
                elif file_path.lower().endswith('.xlsx'):
                    try:
                        from openpyxl import load_workbook
                        
                        wb = load_workbook(file_path)
                        if wb and wb.active:
                            # 保存工作簿信息以便后续保存
                            self.current_workbook = wb
                            self.current_sheet_name = wb.active.title
                            self.original_excel_path = file_path  # 保存原始Excel文件路径
                            
                            ws = wb.active
                            
                            # 清空表格和格式存储
                            self.table_widget.clear()
                            self.excel_cell_formats = {}  # 清空格式存储
                            
                            # 获取行数和列数
                            max_row = ws.max_row if hasattr(ws, 'max_row') else 0
                            max_col = ws.max_column if hasattr(ws, 'max_column') else 0
                            self.table_widget.setRowCount(max_row)
                            self.table_widget.setColumnCount(max_col)
                            
                            # 加载Excel数据并保存格式信息
                            try:
                                for row_idx, row in enumerate(ws.iter_rows(), 1):
                                    for col_idx, cell in enumerate(row, 1):
                                        if cell and cell.value is not None:
                                            try:
                                                value = str(cell.value) if cell.value is not None else ""
                                                self.table_widget.setItem(
                                                    row_idx - 1,  # QTableWidget从0开始
                                                    col_idx - 1,
                                                    QTableWidgetItem(value))
                                                
                                                # 保存单元格格式信息
                                                self.save_cell_format(cell, row_idx - 1, col_idx - 1)
                                            except Exception:
                                                continue
                            except Exception as e:
                                QMessageBox.warning(self, "警告", f"加载Excel数据失败: {str(e)}")
                                return
                            
                            self.content_tabs.setCurrentIndex(1)
                        else:
                            QMessageBox.warning(self, "警告", "Excel文件中没有活动工作表")
                    except Exception as e:
                        QMessageBox.warning(self, "警告", f"加载Excel文件失败: {str(e)}")
                elif file_path.lower().endswith('.pdf'):
                    progress = QProgressDialog("正在解析PDF文件，请稍候...", "取消", 0, 0, self)
                    progress.setWindowTitle("PDF处理中")
                    progress.setModal(True)
                    progress.show()
                    QApplication.processEvents()

                    try:
                        progress.setLabelText("正在读取PDF文本...")
                        QApplication.processEvents()

                        self.reset_pdf_state()
                        display_text = self.load_pdf_with_pymupdf(file_path)
                        if display_text is None:
                            progress.close()
                            return

                        if not self.pdf_char_map:
                            progress.close()
                            QMessageBox.warning(self, "警告", "未检测到可解析的文本内容，可能是扫描件PDF")
                            self.reset_pdf_state()
                            return

                        progress.setLabelText("正在缓存原PDF字体...")
                        QApplication.processEvents()
                        self.build_pdf_font_cache()

                        self.pdf_edit.setPlainText(display_text)
                        self.content_tabs.setCurrentIndex(1)
                        self.is_pdf_source = True
                        self.pdf_redaction_history.clear()

                        progress.close()
                        QMessageBox.information(self, "PDF加载完成", 
                            f"PDF文件已成功解析，可直接在界面中进行脱敏操作\n"
                            f"文件: {os.path.basename(file_path)}")
                        self.status_label.setText("PDF文本已加载，支持原格式脱敏")
                    except Exception as e:
                        progress.close()
                        QMessageBox.warning(self, "警告", f"处理PDF文件失败: {str(e)}")
                        self.reset_pdf_state()
                        return
                elif file_path.lower().endswith(('.docx', '.doc')):
                    try:
                        # .doc 文件仅提示用户先转换为 .docx，避免直接处理
                        if file_path.lower().endswith('.doc'):
                            QMessageBox.information(
                                self,
                                "格式提示",
                                "很抱歉，由于兼容性原因，暂不支持直接处理DOC格式文件。\n请先使用word将DOC文件另存为DOCX格式后再继续脱敏工作。"
                            )
                            return

                        # 处理DOCX文档
                        self.is_pdf_source = False  # 标记非PDF来源
                        content = self.load_word_document(file_path)
                        if content is not None:
                            self.word_edit.setPlainText(content)
                            self.content_tabs.setCurrentIndex(0)  # 切换到Word标签页
                        else:
                            QMessageBox.warning(self, "警告", "无法读取DOCX文档")
                            return
                    except Exception as e:
                        QMessageBox.warning(self, "警告", f"加载Word文档失败: {str(e)}")
                else:
                    QMessageBox.warning(self, "警告", "不支持的文件格式")
            except Exception as e:
                QMessageBox.warning(self, "警告", f"加载文件失败: {str(e)}")

    def select_output_path(self):
        """设置输出路径"""
        if not hasattr(self, 'input_file_path'):
            QMessageBox.warning(self, "警告", "请先选择输入文件")
            return
            
        # 获取输入文件信息
        input_dir = os.path.dirname(self.input_file_path)
        input_name = os.path.splitext(os.path.basename(self.input_file_path))[0]
        input_ext = os.path.splitext(self.input_file_path)[1]
        
        # 生成默认输出文件名
        default_name = f"{input_name}（脱敏）{input_ext}"
        default_path = os.path.join(input_dir, default_name)
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "设置输出路径",
            default_path,  # 设置默认路径和文件名
            f"相同类型文件 (*{input_ext});;所有文件 (*)"
        )
        if file_path:
            # 确保输出文件扩展名与输入文件一致
            output_ext = os.path.splitext(file_path)[1]
            if output_ext.lower() != input_ext.lower():
                file_path = file_path + input_ext
            self.output_file_path = file_path
            self.file_info_label.setText(f"{self.file_info_label.text()}\n输出路径: {os.path.basename(file_path)}")
            self.status_label.setText("已设置输出路径")

    def process_file(self):
        """执行文件脱敏处理"""
        if not hasattr(self, 'input_file_path'):
            QMessageBox.warning(self, "警告", "请先选择输入文件")
            return
            
        if not hasattr(self, 'output_file_path'):
            QMessageBox.warning(self, "警告", "请先设置输出路径")
            return
        
        # 检查处理模式
        is_interactive_mode = self.mode_combo.currentIndex() == 0
        
        if is_interactive_mode:
            # 交互式脱敏模式 - 只保存当前界面上的修改
            self.save_interactive_changes()
        else:
            # 自动脱敏模式 - 应用规则到整个文件
            if not self.rule_engine.get_active_rules():
                QMessageBox.warning(self, "警告", "没有可用的脱敏规则")
                return
            self.auto_process_file()
    
    def save_interactive_changes(self):
        """保存交互式脱敏的更改"""
        try:
            # 如果没有设置输出路径，自动生成默认路径
            if not hasattr(self, 'output_file_path') or not self.output_file_path:
                input_dir = os.path.dirname(self.input_file_path)
                input_name = os.path.splitext(os.path.basename(self.input_file_path))[0]
                input_ext = os.path.splitext(self.input_file_path)[1]
                default_name = f"{input_name}（脱敏）{input_ext}"
                self.output_file_path = os.path.join(input_dir, default_name)
            
            if self.input_file_path.endswith(('.xlsx', '.xls')):
                self.save_excel_changes()
            elif self.input_file_path.endswith('.txt'):
                self.save_text_changes()
            elif self.input_file_path.endswith('.doc'):
                QMessageBox.information(self, "DOC文件格式提示", "由于兼容性原因，暂不支持直接处理DOC格式文件。\n请先将DOC文件转换为DOCX格式后再进行脱敏处理。")
                return
            elif self.input_file_path.endswith('.docx'):
                self.save_word_changes()
            elif self.input_file_path.endswith('.pdf'):
                # PDF文件：从PDF标签页保存内容
                self.save_pdf_changes()
            else:
                QMessageBox.warning(self, "警告", "当前文件类型不支持交互式脱敏")
                return
            
            QMessageBox.information(self, "完成", f"文件已保存到: {self.output_file_path}")

            # 新增：询问是否导出日志
            log_exported = self.show_export_log_dialog()

            # 询问是否打开文件
            reply = QMessageBox.question(self, "打开文件", "是否立即打开刚保存的文件？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes:
                try:
                    import subprocess
                    import platform
                    file_path = self.output_file_path
                    if platform.system() == "Windows":
                        os.startfile(file_path)
                    elif platform.system() == "Darwin":
                        subprocess.Popen(["open", file_path])
                    else:
                        subprocess.Popen(["xdg-open", file_path])
                except Exception as e_open:
                    QMessageBox.critical(self, "打开失败", f"文件已保存，但打开失败：{str(e_open)}")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存文件失败: {str(e)}")
    
    def save_excel_changes(self):
        """保存Excel文件的交互式修改（保持格式）"""
        import openpyxl
        from openpyxl import load_workbook
        
        try:
            # 如果有原始文件路径，直接加载原始工作簿以保持格式
            if hasattr(self, 'original_excel_path') and self.original_excel_path:
                wb_new = load_workbook(self.original_excel_path)
            else:
                # fallback：创建新的工作簿
                wb_new = openpyxl.Workbook()
                if wb_new.active:
                    wb_new.remove(wb_new.active)  # 删除默认工作表
            
            # 从界面上的表格获取数据并更新到工作簿
            if hasattr(self, 'current_workbook') and self.current_workbook:
                # 处理当前显示的工作表
                if hasattr(self, 'current_sheet_name') and self.current_sheet_name:
                    ws_new = wb_new[self.current_sheet_name] if self.current_sheet_name in wb_new.sheetnames else wb_new.active
                    
                    # 用界面上修改后的数据更新工作表，但保持原有格式
                    for row in self.table_widget.selectedItems():
                        if row:
                            row_idx = row.row()
                            col_idx = row.column()
                            cell = ws_new.cell(row_idx + 1, col_idx + 1)
                            item = self.table_widget.item(row_idx, col_idx)
                            if item:
                                cell.value = item.text()
                                # 如果有保存的格式信息，应用格式
                                self.apply_cell_format(cell, row_idx, col_idx)

                    # 处理未选中但在历史记录中的单元格
                    for row_idx in range(self.table_widget.rowCount()):
                        for col_idx in range(self.table_widget.columnCount()):
                            item = self.table_widget.item(row_idx, col_idx)
                            if item:
                                # 检查该单元格是否在当前选中范围内
                                if not any(row_idx == r.row() and col_idx == r.column() for r in self.table_widget.selectedItems()):
                                    cell = ws_new.cell(row_idx + 1, col_idx + 1)
                                    cell.value = item.text()
                                    # 如果有保存的格式信息，应用格式
                                    self.apply_cell_format(cell, row_idx, col_idx)
            
            try:
                wb_new.save(self.output_file_path)
            except PermissionError:
                QMessageBox.critical(self, "文件权限错误", 
                    f"无法保存文件到：{self.output_file_path}\n\n"
                    "可能的解决方案：\n"
                    "1. 关闭正在使用该文件的Excel程序\n"
                    "2. 检查文件是否设为只读\n"
                    "3. 以管理员身份运行本程序\n"
                    "4. 选择其他保存位置")
                return  # 停止处理
            
        except Exception as e:
            # 如果格式保持失败，使用简化方法
            print(f"格式保持失败，使用简化保存方法: {str(e)}")
            self.save_excel_changes_simple()
    
    def save_excel_changes_simple(self):
        """简化的Excel保存方法（不保持格式）"""
        import openpyxl
        
        # 创建新的工作簿
        wb_new = openpyxl.Workbook()
        ws_new = wb_new.active
        ws_new.title = getattr(self, 'current_sheet_name', 'Sheet1')
        
        # 从界面表格获取数据
        for row in range(self.table_widget.rowCount()):
            for col in range(self.table_widget.columnCount()):
                item = self.table_widget.item(row, col)
                if item and item.text():
                    ws_new.cell(row + 1, col + 1).value = item.text()
        
        try:
            wb_new.save(self.output_file_path)
        except PermissionError:
            QMessageBox.critical(self, "文件权限错误", 
                f"无法保存文件到：{self.output_file_path}\n\n"
                "可能的解决方案：\n"
                "1. 关闭正在使用该文件的Excel程序\n"
                "2. 检查文件是否设为只读\n"
                "3. 以管理员身份运行本程序\n"
                "4. 选择其他保存位置")
            return
    
    def save_text_changes(self):
        """保存文本文件的交互式修改"""
        # 获取文本编辑器中的内容
        content = self.text_edit.toPlainText()
        
        # 使用原始编码保存
        encoding = getattr(self, 'original_encoding', 'utf-8')
        try:
            with open(self.output_file_path, 'w', encoding=encoding) as f:
                f.write(content)
        except PermissionError:
            QMessageBox.critical(self, "文件权限错误", 
                f"无法保存文件到：{self.output_file_path}\n\n"
                "可能的解决方案：\n"
                "1. 关闭正在使用该文件的其他程序\n"
                "2. 检查文件是否设为只读\n"
                "3. 以管理员身份运行本程序\n"
                "4. 选择其他保存位置")
            raise
    
    def save_word_changes(self):
        """保存Word文档的交互式修改（保持原格式）"""
        try:
            if not hasattr(self, 'current_word_doc') or not self.current_word_doc:
                QMessageBox.warning(self, "警告", "没有加载可保存的 Word 文档，请先打开文件。")
                return
            from docx import Document
            
            # 只支持DOCX文件的保存
            if self.input_file_path.lower().endswith('.doc'):
                QMessageBox.warning(self, "不支持的格式", 
                    "DOC格式不支持交互式脱敏保存\n\n请先将DOC文件转换为DOCX格式：\n"
                    "1. 用Word打开DOC文件\n"
                    "2. 另存为DOCX格式\n"
                    "3. 重新选择DOCX文件进行脱敏")
                return
            else:
                # 直接打开DOCX文件
                doc = Document(self.input_file_path)
            
            # 获取原始文本内容（用于比对）
            original_text = self.get_word_text_content(doc)
            
            # 获取编辑器内容作为新内容
            new_content = self.word_edit.toPlainText()
            
            # 计算需要替换的内容
            replacements = self.calculate_text_replacements(original_text, new_content)
            
            # 应用替换到文档（保持格式）
            self.apply_word_replacements(doc, replacements)
            
            # 保存文档
            if self.output_file_path.lower().endswith(('.doc', '.docx')):
                output_path = self.output_file_path
                if not output_path.lower().endswith('.docx'):
                    output_path = os.path.splitext(output_path)[0] + '.docx'
            else:
                output_path = os.path.splitext(self.output_file_path)[0] + '.docx'
            
            try:
                doc.save(output_path)
                self.output_file_path = output_path
            except PermissionError:
                QMessageBox.critical(self, "文件权限错误", 
                    f"无法保存文件到：{output_path}\n\n"
                    "可能的解决方案：\n"
                    "1. 关闭正在使用该文件的Word程序\n"
                    "2. 检查文件是否设为只读\n"
                    "3. 以管理员身份运行本程序\n"
                    "4. 选择其他保存位置")
                raise
            
            # 已保存，主流程统一弹窗，无需此处弹窗
            
        except ImportError:
            QMessageBox.warning(self, "警告", "未安装python-docx库，无法保存DOCX文件")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存Word文档时出错: {str(e)}")

    def save_pdf_changes(self):
        """保存PDF文档的交互式修改（直接在PDF上写入）"""
        progress = QProgressDialog("正在保存PDF文件，请稍候...", "取消", 0, 0, self)
        progress.setWindowTitle("PDF保存中")
        progress.setModal(True)
        progress.show()
        QApplication.processEvents()

        try:
            progress.setLabelText("正在准备PDF文档...")
            QApplication.processEvents()

            doc = fitz.open(self.input_file_path)
            self.pdf_doc = doc
            self.build_pdf_font_cache()

            output_path = self.output_file_path
            if not output_path.lower().endswith('.pdf'):
                output_path = os.path.splitext(output_path)[0] + '.pdf'
                self.output_file_path = output_path

            if not self.pdf_pending_redactions:
                # 没有实际修改，直接另存即可
                doc.save(self.output_file_path)
                progress.close()
                return

            progress.setLabelText("正在应用脱敏内容...")
            QApplication.processEvents()

            applied_segments = 0
            redacted_pages = set()
            for operation in self.pdf_pending_redactions:
                segments = operation.get('segments', [])
                for segment in segments:
                    page_index = segment.get('page', 0)
                    if page_index >= doc.page_count:
                        continue
                    page = doc.load_page(page_index)
                    if self.apply_pdf_segment(page, segment):
                        applied_segments += 1
                        redacted_pages.add(page_index)

            if applied_segments == 0:
                progress.close()
                QMessageBox.warning(self, "提示", "未能在PDF中定位可写入的文本区域，已保留原PDF")
                doc.save(self.output_file_path)
                return

            # 应用所有红线脱敏标记，彻底删除原文
            for page_index in sorted(redacted_pages):
                try:
                    redact_page = doc.load_page(page_index)
                    redact_page.apply_redactions()
                except Exception as apply_err:
                    print(f"应用PDF红线脱敏失败: 页面 {page_index}, 错误: {apply_err}")

            progress.setLabelText("正在写出PDF文件...")
            QApplication.processEvents()
            doc.save(self.output_file_path, garbage=4, deflate=True)
            progress.close()
            self.pdf_pending_redactions = []

        except PermissionError:
            progress.close()
            QMessageBox.critical(self, "文件权限错误", 
                f"无法保存文件到：{self.output_file_path}\n\n"
                "可能的解决方案：\n"
                "1. 关闭正在使用该文件的PDF阅读器\n"
                "2. 检查文件是否设为只读\n"
                "3. 以管理员身份运行本程序\n"
                "4. 选择其他保存位置")
            raise
        except Exception as e:
            progress.close()
            QMessageBox.critical(self, "错误", f"保存PDF文档时出错: {str(e)}")

    def get_word_text_content(self, doc):
        """提取Word文档的纯文本内容"""
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return '\n'.join(text_parts)
    
    def calculate_text_replacements(self, original_text, new_text):
        """计算文本差异，返回需要替换的内容"""
        # 简单的逐行比较替换策略
        original_lines = original_text.split('\n')
        new_lines = new_text.split('\n')
        
        replacements = []
        min_len = min(len(original_lines), len(new_lines))
        
        for i in range(min_len):
            if original_lines[i] != new_lines[i]:
                replacements.append((original_lines[i], new_lines[i]))
        
        return replacements
    
    def apply_word_replacements(self, doc, replacements):
        """将替换应用到Word文档，保持格式"""
        for old_text, new_text in replacements:
            if old_text.strip() and old_text != new_text:
                # 替换段落中的文本
                for para in doc.paragraphs:
                    if old_text in para.text:
                        self.replace_text_in_paragraph(para, old_text, new_text)
                
                # 替换表格中的文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if old_text in cell.text:
                                for para in cell.paragraphs:
                                    if old_text in para.text:
                                        self.replace_text_in_paragraph(para, old_text, new_text)
    
    def replace_text_in_paragraph(self, paragraph, old_text, new_text):
        """在段落中替换文本，保持格式"""
        if old_text in paragraph.text:
            # 获取段落的所有runs
            runs = paragraph.runs
            paragraph_text = paragraph.text
            
            if old_text in paragraph_text:
                # 找到替换位置
                start_pos = paragraph_text.find(old_text)
                end_pos = start_pos + len(old_text)
                
                # 记录格式信息
                current_pos = 0
                start_run_idx = -1
                end_run_idx = -1
                start_run_pos = 0
                end_run_pos = 0
                
                # 找到起始和结束run的位置
                for i, run in enumerate(runs):
                    run_len = len(run.text)
                    if start_run_idx == -1 and current_pos + run_len > start_pos:
                        start_run_idx = i
                        start_run_pos = start_pos - current_pos
                    if current_pos + run_len >= end_pos:
                        end_run_idx = i
                        end_run_pos = end_pos - current_pos
                        break
                    current_pos += run_len
                
                # 执行替换
                if start_run_idx != -1 and end_run_idx != -1:
                    if start_run_idx == end_run_idx:
                        # 在同一个run内替换
                        run = runs[start_run_idx]
                        run.text = run.text[:start_run_pos] + new_text + run.text[end_run_pos:]
                    else:
                        # 跨多个run替换
                        runs[start_run_idx].text = runs[start_run_idx].text[:start_run_pos] + new_text
                        for i in range(start_run_idx + 1, end_run_idx + 1):
                            if i < len(runs):
                                if i == end_run_idx:
                                    runs[i].text = runs[i].text[end_run_pos:]
                                else:
                                    runs[i].text = ""

    def auto_process_file(self):
        """自动脱敏模式 - 应用规则到整个文件"""
            
        try:
            # 确保输出目录存在
            os.makedirs(os.path.dirname(self.output_file_path), exist_ok=True)

            # Excel
            if self.input_file_path.lower().endswith('.xlsx'):
                try:
                    from openpyxl import load_workbook
                    from copy import copy
                    import re

                    wb_original = load_workbook(self.input_file_path)
                    wb_new = load_workbook(self.input_file_path)  # 创建副本保留格式

                    # 记录自动脱敏的历史
                    auto_redaction_history = []

                    for ws_name in wb_original.sheetnames:
                        ws_original = wb_original[ws_name]
                        ws_new = wb_new[ws_name]

                        # 复制合并单元格信息
                        try:
                            for merged_range in list(ws_original.merged_cells.ranges):
                                ws_new.merge_cells(str(merged_range))
                        except:
                            pass  # 如果合并单元格操作失败，继续处理

                        for row in ws_original.iter_rows():
                            for cell in row:
                                if cell.value is not None:
                                    cell_value = str(cell.value)
                                    original_value = cell_value
                                    applied_rules = []  # 记录应用的规则
                                    
                                    for rule in self.rule_engine.get_active_rules():
                                        old_value = cell_value
                                        if rule.name == "自定义字段":
                                            custom_fields = getattr(self, 'custom_fields', None)
                                            cell_value = self.rule_engine.apply_redaction_rule(rule, cell_value, custom_fields)
                                        else:
                                            custom_names = getattr(self, 'custom_names', None)
                                            cell_value = self.rule_engine.apply_redaction_rule(rule, cell_value, custom_names)
                                        
                                        # 如果这个规则产生了变化，记录它
                                        if cell_value != old_value:
                                            applied_rules.append(rule)
                                    
                                    # 更新单元格值并保持格式
                                    target_cell = ws_new[cell.coordinate]
                                    if cell_value != original_value:
                                        target_cell.value = cell_value
                                        
                                        # 记录自动脱敏历史（仅限活动工作表用于界面显示）
                                        if ws_name == wb_original.active.title:
                                            # 使用最后一个应用的规则，或者合并规则名
                                            rule_names = [r.name for r in applied_rules] if applied_rules else ['自动规则脱敏']
                                            primary_rule = applied_rules[-1] if applied_rules else None
                                            
                                            auto_redaction_history.append({
                                                'row': cell.row - 1,  # 转换为0索引
                                                'col': cell.column - 1,
                                                'original_text': original_value,
                                                'redacted_text': cell_value,
                                                'original_background': QColor(),
                                                'original_tooltip': '',
                                                'rule_name': ', '.join(rule_names),  # 合并所有应用的规则名
                                                'mode': '自动规则脱敏',
                                                'timestamp': self.get_current_timestamp(),
                                                'rule': primary_rule  # 保存主要规则对象引用
                                            })
                                    
                                    # 复制所有格式属性
                                    if cell.has_style:
                                        target_cell.font = copy(cell.font)
                                        target_cell.border = copy(cell.border)
                                        target_cell.fill = copy(cell.fill)
                                        target_cell.number_format = cell.number_format
                                        target_cell.protection = copy(cell.protection)
                                        target_cell.alignment = copy(cell.alignment)

                    try:
                        wb_new.save(self.output_file_path)
                    except PermissionError:
                        QMessageBox.critical(self, "文件权限错误", 
                            f"无法保存文件到：{self.output_file_path}\n\n"
                            "可能的解决方案：\n"
                            "1. 关闭正在使用该文件的Excel程序\n"
                            "2. 检查文件是否设为只读\n"
                            "3. 以管理员身份运行本程序\n"
                            "4. 选择其他保存位置")
                        return  # 停止处理
                    
                    # 保存历史记录
                    if auto_redaction_history:
                        self.excel_redaction_history.append({
                            'type': 'auto_rule_redaction',
                            'operations': auto_redaction_history
                        })
                    
                    # 在界面中显示脱敏后的结果（优化输出格式与交互模式一致）
                    try:
                        # 保存工作簿信息以便后续操作
                        self.current_workbook = wb_new
                        self.current_sheet_name = wb_new.active.title
                        
                        ws = wb_new.active
                        
                        # 清空表格
                        self.table_widget.clear()
                        
                        # 获取行数和列数
                        max_row = ws.max_row if hasattr(ws, 'max_row') else 0
                        max_col = ws.max_column if hasattr(ws, 'max_column') else 0
                        self.table_widget.setRowCount(max_row)
                        self.table_widget.setColumnCount(max_col)
                        
                        # 加载脱敏后的Excel数据到界面
                        for row_idx, row in enumerate(ws.iter_rows(), 1):
                            for col_idx, cell in enumerate(row, 1):
                                if cell and cell.value is not None:
                                    try:
                                        value = str(cell.value) if cell.value is not None else ""
                                        table_item = QTableWidgetItem(value)
                                        
                                        # 检查是否为脱敏单元格，如果是则标记
                                        for operation in auto_redaction_history:
                                            if (operation['row'] == row_idx - 1 and 
                                                operation['col'] == col_idx - 1):
                                                table_item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                                                table_item.setToolTip(f"已脱敏 - 原文本: {operation['original_text']}")
                                                break
                                        
                                        self.table_widget.setItem(
                                            row_idx - 1,  # QTableWidget从0开始
                                            col_idx - 1,
                                            table_item)
                                    except Exception:
                                        continue
                        
                        # 切换到Excel标签页显示结果
                        self.content_tabs.setCurrentIndex(1)
                        
                    except Exception as e:
                        QMessageBox.warning(self, "警告", f"在界面显示脱敏结果失败: {str(e)}")

                except Exception as e:
                    QMessageBox.warning(self, "警告", f"处理Excel文件失败: {str(e)}")
                    return

            # Word (.docx 或 .doc)
            elif self.input_file_path.lower().endswith(('.docx', '.doc')):
                try:
                    # 仅提示 .doc，让用户先转换
                    if self.input_file_path.lower().endswith('.doc'):
                        QMessageBox.information(
                            self,
                            "格式提示",
                            "很抱歉，由于兼容性原因，暂不支持直接处理DOC格式文件。\n请先使用word将DOC文件转换为DOCX格式后再继续脱敏工作。"

                        )
                        return

                    # 检查文件是否被占用
                    try:
                        with open(self.input_file_path, 'r+b') as test_file:
                            pass  # 如果能打开，说明文件未被占用
                    except PermissionError:
                        QMessageBox.critical(self, "文件被占用", 
                                           f"无法处理Word文档，文件可能正在被使用：\n\n"
                                           f"请关闭可能正在使用该文件的程序（如Microsoft Word）后重试。\n\n"
                                           f"文件路径：{self.input_file_path}")
                        return

                    from docx import Document
                    doc = Document(self.input_file_path)

                    # 段落
                    for para_idx, para in enumerate(doc.paragraphs):
                        original_text = para.text
                        processed_text = original_text
                        for rule in self.rule_engine.get_active_rules():
                            if rule.name == "自定义字段":
                                custom_fields = getattr(self, 'custom_fields', None)
                                processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_fields)
                            else:
                                custom_names = getattr(self, 'custom_names', None)
                                processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_names)
                        if original_text != processed_text:
                            para.text = processed_text
                            # 记录Word脱敏历史
                            self.word_redaction_history.append({
                                'original': original_text,
                                'redacted': processed_text,
                                'timestamp': self.get_current_timestamp(),
                                'rule_name': '自动脱敏',
                                'mode': '自动脱敏',
                                'rule_type': '规则引擎',
                                'position_desc': f"段落 {para_idx + 1}"
                            })

                    # 表格
                    for table_idx, table in enumerate(doc.tables):
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                original_text = cell.text
                                processed_text = original_text
                                for rule in self.rule_engine.get_active_rules():
                                    if rule.name == "自定义字段":
                                        custom_fields = getattr(self, 'custom_fields', None)
                                        processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_fields)
                                    else:
                                        custom_names = getattr(self, 'custom_names', None)
                                        processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_names)
                                if original_text != processed_text:
                                    cell.text = processed_text
                                    # 记录Word表格脱敏历史
                                    self.word_redaction_history.append({
                                        'original': original_text,
                                        'redacted': processed_text,
                                        'timestamp': self.get_current_timestamp(),
                                        'rule_name': '自动脱敏',
                                        'mode': '自动脱敏',
                                        'rule_type': '规则引擎',
                                        'position_desc': f"表格{table_idx + 1} 行{row_idx + 1} 列{cell_idx + 1}"
                                    })

                    # 保存为 docx
                    try:
                        if self.output_file_path.lower().endswith('.docx'):
                            doc.save(self.output_file_path)
                        else:
                            output_path = os.path.splitext(self.output_file_path)[0] + '.docx'
                            doc.save(output_path)
                    except PermissionError:
                        QMessageBox.critical(self, "文件权限错误", 
                            f"无法保存文件到：{self.output_file_path}\n\n"
                            "可能的解决方案：\n"
                            "1. 关闭正在使用该文件的Word程序\n"
                            "2. 检查文件是否设为只读\n"
                            "3. 以管理员身份运行本程序\n"
                            "4. 选择其他保存位置")
                        return  # 停止处理
                        self.output_file_path = output_path

                except ImportError:
                    QMessageBox.warning(self, "警告", "处理Word文档需要安装python-docx库\n请运行: pip install python-docx")
                    return
                except PermissionError as e:
                    QMessageBox.critical(self, "权限错误", 
                                       f"无法访问Word文档，请检查：\n\n"
                                       f"1. 是否有其他程序（如Microsoft Word）正在使用该文件\n"
                                       f"2. 文件是否被设置为只读\n"
                                       f"3. 是否有足够的文件访问权限\n\n"
                                       f"错误详情: {str(e)}")
                    return
                except Exception as e:
                    error_msg = str(e)
                    if "Permission denied" in error_msg or "Errno 13" in error_msg:
                        QMessageBox.critical(self, "权限错误", 
                                           f"无法访问Word文档，请检查：\n\n"
                                           f"1. 是否有其他程序（如Microsoft Word）正在使用该文件\n"
                                           f"2. 文件是否被设置为只读\n"
                                           f"3. 是否有足够的文件访问权限\n\n"
                                           f"错误详情: {error_msg}")
                    else:
                        QMessageBox.warning(self, "警告", f"处理Word文档失败: {error_msg}")
                    return

            # 文本
            elif self.input_file_path.lower().endswith('.txt'):
                encoding = None
                encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig', 'latin1']
                import re

                lines = []
                for enc in encodings:
                    try:
                        with open(self.input_file_path, 'r', encoding=enc) as f:
                            lines = f.readlines()
                        encoding = enc
                        break
                    except UnicodeDecodeError:
                        continue

                if encoding is None or not lines:
                    QMessageBox.critical(self, "错误", "无法读取文件内容或确定编码格式")
                    return

                processed_lines = []
                for line_idx, line in enumerate(lines):
                    original_line = line
                    processed_line = line
                    for rule in self.rule_engine.get_active_rules():
                        if rule.name == "自定义字段":
                            custom_fields = getattr(self, 'custom_fields', None)
                            processed_line = self.rule_engine.apply_redaction_rule(rule, processed_line, custom_fields)
                        else:
                            custom_names = getattr(self, 'custom_names', None)
                            processed_line = self.rule_engine.apply_redaction_rule(rule, processed_line, custom_names)
                    
                    # 记录文本脱敏历史
                    if original_line != processed_line:
                        self.text_redaction_history.append({
                            'original': original_line.strip(),
                            'redacted': processed_line.strip(),
                            'timestamp': self.get_current_timestamp(),
                            'rule_name': '自动脱敏',
                            'mode': '自动脱敏',
                            'rule_type': '规则引擎',
                            'position_desc': f"第 {line_idx + 1} 行"
                        })
                    
                    processed_lines.append(processed_line)

                try:
                    with open(self.output_file_path, 'w', encoding=encoding) as f:
                        f.writelines(processed_lines)
                except PermissionError:
                    QMessageBox.critical(self, "文件权限错误", 
                        f"无法保存文件到：{self.output_file_path}\n\n"
                        "可能的解决方案：\n"
                        "1. 关闭正在使用该文件的其他程序\n"
                        "2. 检查文件是否设为只读\n"
                        "3. 以管理员身份运行本程序\n"
                        "4. 选择其他保存位置")
                    return  # 停止处理

            # PDF文件处理
            elif self.input_file_path.lower().endswith('.pdf'):
                self.reset_pdf_state()
                display_text = self.load_pdf_with_pymupdf(self.input_file_path)

                if display_text is None or not self.pdf_char_map:
                    QMessageBox.warning(self, "警告", "未能解析PDF文本，可能为扫描件或受保护的PDF")
                    return

                operations, updated_text = self.auto_redact_pdf()

                if not operations:
                    doc = fitz.open(self.input_file_path)
                    doc.save(self.output_file_path)
                    doc.close()
                    QMessageBox.information(self, "提示", "未检测到可脱敏的内容，已复制原PDF文件。")
                    return

                self.pdf_pending_redactions = operations
                self.pdf_redaction_history.extend(operations)
                self.pdf_edit.setPlainText(updated_text)
                self.content_tabs.setCurrentIndex(1)
                self.is_pdf_source = True

                self.save_pdf_changes()
                self.status_label.setText(f"自动脱敏完成 {len(operations)} 处敏感信息")

            else:
                QMessageBox.warning(self, "警告", "不支持的文件格式")

            QMessageBox.information(self, "成功", "文件脱敏处理完成")
            self.status_label.setText("处理完成")

            # 新增：询问是否导出日志
            log_exported = self.show_export_log_dialog()

            # 询问是否打开文件
            reply = QMessageBox.question(self, "打开文件", "是否立即打开刚保存的文件？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes:
                try:
                    import subprocess
                    import platform
                    file_path = self.output_file_path
                    if platform.system() == "Windows":
                        os.startfile(file_path)
                    elif platform.system() == "Darwin":
                        subprocess.Popen(["open", file_path])
                    else:
                        subprocess.Popen(["xdg-open", file_path])
                except Exception as e_open:
                    QMessageBox.critical(self, "打开失败", f"文件已保存，但打开失败：{str(e_open)}")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"处理文件失败: {str(e)}")
            self.status_label.setText("处理失败")

    def batch_process(self):
        """批量处理文件"""
        # 第一步：检查是否在自动规则模式
        if self.mode_combo.currentIndex() != 1:
            QMessageBox.warning(self, "提示", "批量处理功能仅在【自动脱敏（规则模式）】下可用！\n\n请先切换到自动脱敏模式。")
            return
        
        # 第二步：检查是否有激活的规则
        active_rules = self.rule_engine.get_active_rules()
        if not active_rules:
            QMessageBox.warning(self, "提示", "没有激活的脱敏规则！\n\n请先点击【配置脱敏规则】按钮激活需要的规则。")
            return
        
        # 第三步：显示步骤提示和选择处理方式
        help_msg = """
        <h3>📦 批量处理使用说明</h3>
        
        <p><b>✅ 当前激活规则：</b><br>
        """ + "<br>".join([f"• {rule.name}" for rule in active_rules]) + """</p>
        
    <p><b>📁 支持格式：</b>Excel (.xlsx)、Word (.docx)、文本 (.txt)、PDF (.pdf)</p>
        
    <p><b>🛠️ 使用步骤：</b>选择文件或文件夹 → 选择输出目录 → 自动处理完成</p>
        
    <p><b>💡 提示：</b>输出文件将自动添加"（脱敏）"标识</p>
        """
        
        # 选择批量处理方式
        choice_dialog = QMessageBox(self)
        choice_dialog.setWindowTitle("批量处理方式选择")
        choice_dialog.setText(help_msg + "\n\n请选择批量处理方式：")
        choice_dialog.setIcon(QMessageBox.Icon.Question)
        
        folder_btn = choice_dialog.addButton("📁 选择文件夹", QMessageBox.ButtonRole.AcceptRole)
        multi_files_btn = choice_dialog.addButton("多选文件", QMessageBox.ButtonRole.AcceptRole)
        cancel_btn = choice_dialog.addButton("取消", QMessageBox.ButtonRole.RejectRole)
        
        choice_dialog.exec()
        clicked_btn = choice_dialog.clickedButton()
        
        if clicked_btn == cancel_btn:
            return
            
        # 根据选择方式获取文件列表
        input_files = []
        output_dir = None
        
        if clicked_btn == folder_btn:
            # 选择文件夹方式
            default_dir = os.path.dirname(os.path.abspath(__file__))
            input_dir = QFileDialog.getExistingDirectory(
                self,
                "第1步：选择待处理文件夹",
                default_dir
            )
            if not input_dir:
                return
                
            # 选择输出文件夹
            output_dir = QFileDialog.getExistingDirectory(
                self,
                "第2步：选择输出目录",
                default_dir
            )
            if not output_dir:
                return
            
            # 获取文件夹中的所有支持文件
            for filename in os.listdir(input_dir):
                if filename.lower().endswith(('.xlsx', '.docx', '.txt', '.pdf')):
                    input_files.append(os.path.join(input_dir, filename))
                    
        elif clicked_btn == multi_files_btn:
            # 多选文件方式
            script_dir = os.path.dirname(os.path.abspath(__file__))
            file_paths, _ = QFileDialog.getOpenFileNames(
                self,
                "选择需要批量处理的文件（可多选）",
                script_dir,
                "支持的文件 (*.xlsx *.docx *.txt *.pdf);;Excel文件 (*.xlsx);;Word文档 (*.docx);;PDF文档 (*.pdf);;文本文件 (*.txt);;所有文件 (*)"
            )
            if not file_paths:
                return
                
            input_files = file_paths
            
            # 选择输出文件夹
            default_dir = os.path.dirname(os.path.abspath(__file__))
            output_dir = QFileDialog.getExistingDirectory(
                self,
                "选择输出目录",
                default_dir
            )
            if not output_dir:
                return
        
        if not input_files:
            QMessageBox.information(self, "提示", "没有找到支持的文件格式\n\n支持格式：.xlsx、.docx、.txt")
            return
        
        # 确保 output_dir 不为 None
        if output_dir is None:
            QMessageBox.warning(self, "错误", "输出目录未选择")
            return
        
        try:
            processed_count = 0
            failed_files = []
            
            self.progress_bar.setRange(0, len(input_files))
            self.progress_bar.setValue(0)
            
            for i, input_path in enumerate(input_files):
                filename = os.path.basename(input_path)
                
                # 生成带（脱敏）标识的输出文件名
                name_without_ext = os.path.splitext(filename)[0]
                file_ext = os.path.splitext(filename)[1]
                output_filename = f"{name_without_ext}（脱敏）{file_ext}"
                output_path = os.path.join(output_dir, output_filename)
                
                if not os.path.isfile(input_path):
                    continue
                    
                try:
                    # 处理Excel文件
                    if filename.lower().endswith('.xlsx'):
                        self.status_label.setText(f"正在处理Excel: {filename}")
                        QApplication.processEvents()
                        
                        try:
                            from openpyxl import load_workbook
                            from copy import copy
                            
                            wb_original = load_workbook(input_path)
                            wb_new = load_workbook(input_path)  # 创建副本保留格式
                            
                            for ws_name in wb_original.sheetnames:
                                ws_original = wb_original[ws_name]
                                ws_new = wb_new[ws_name]
                                
                                for row in ws_original.iter_rows():
                                    for cell in row:
                                        if cell.value is not None:
                                            cell_value = str(cell.value)
                                            original_value = cell_value  # 保存原始值
                                            
                                            # 使用规则引擎进行脱敏处理
                                            for rule in active_rules:
                                                if rule.name == "姓名":
                                                    custom_names = getattr(self, 'custom_names', None)
                                                    cell_value = self.rule_engine.apply_redaction_rule(rule, cell_value, custom_names)
                                                elif rule.name == "自定义字段":
                                                    custom_fields = getattr(self, 'custom_fields', None)
                                                    cell_value = self.rule_engine.apply_redaction_rule(rule, cell_value, custom_fields)
                                                else:
                                                    cell_value = self.rule_engine.apply_redaction_rule(rule, cell_value)
                                            
                                            if cell_value != original_value:  # 仅在值发生变化时更新
                                                ws_new[cell.coordinate].value = cell_value
                                            if hasattr(cell, 'style'):
                                                ws_new[cell.coordinate].style = copy(cell.style)
                            
                            try:
                                wb_new.save(output_path)
                            except PermissionError:
                                failed_files.append(f"{filename} (文件权限错误: 无法保存到 {output_path})")
                                continue
                            
                        except Exception as e:
                            failed_files.append(f"{filename} (Excel处理失败: {str(e)})")
                            continue
                            
                    elif filename.lower().endswith('.docx'):
                        # 处理Word文档
                        self.status_label.setText(f"正在处理Word: {filename}")
                        QApplication.processEvents()
                        
                        try:
                            from docx import Document
                            
                            doc = Document(input_path)
                            
                            # 处理段落文本
                            for para in doc.paragraphs:
                                original_text = para.text
                                processed_text = original_text
                                
                                # 应用所有激活的脱敏规则
                                for rule in active_rules:
                                    if rule.name == "姓名":
                                        custom_names = getattr(self, 'custom_names', None)
                                        processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_names)
                                    elif rule.name == "自定义字段":
                                        custom_fields = getattr(self, 'custom_fields', None)
                                        processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_fields)
                                    else:
                                        processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text)
                                
                                # 更新段落文本
                                if original_text != processed_text:
                                    para.text = processed_text
                            
                            # 处理表格中的文本
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        original_text = cell.text
                                        processed_text = original_text
                                        
                                        # 应用所有激活的脱敏规则
                                        for rule in active_rules:
                                            if rule.name == "姓名":
                                                custom_names = getattr(self, 'custom_names', None)
                                                processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_names)
                                            elif rule.name == "自定义字段":
                                                custom_fields = getattr(self, 'custom_fields', None)
                                                processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text, custom_fields)
                                            else:
                                                processed_text = self.rule_engine.apply_redaction_rule(rule, processed_text)
                                        
                                        # 更新单元格文本
                                        if original_text != processed_text:
                                            cell.text = processed_text
                            
                            # 保存DOCX文档
                            try:
                                doc.save(output_path)
                            except PermissionError:
                                failed_files.append(f"{filename} (文件权限错误: 无法保存到 {output_path})")
                                continue
                            
                        except ImportError:
                            failed_files.append(f"{filename} (需要python-docx库处理Word文档)")
                            continue
                        except Exception as e:
                            failed_files.append(f"{filename} (Word处理失败: {str(e)})")
                            continue
                            
                    elif filename.lower().endswith('.txt'):
                        # 处理文本文件
                        self.status_label.setText(f"正在处理文本: {filename}")
                        QApplication.processEvents()
                        
                        encoding = None
                        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig', 'latin1']
                        content = ""
                        
                        # 读取文件内容
                        for enc in encodings:
                            try:
                                with open(input_path, 'r', encoding=enc) as f:
                                    content = f.read()
                                encoding = enc
                                break
                            except UnicodeDecodeError:
                                continue
                        
                        if encoding is None:
                            failed_files.append(f"{filename} (编码问题)")
                            continue
                        
                        # 使用规则引擎处理内容
                        processed_content = content
                        for rule in active_rules:
                            if rule.name == "姓名":
                                custom_names = getattr(self, 'custom_names', None)
                                processed_content = self.rule_engine.apply_redaction_rule(rule, processed_content, custom_names)
                            elif rule.name == "自定义字段":
                                custom_fields = getattr(self, 'custom_fields', None)
                                processed_content = self.rule_engine.apply_redaction_rule(rule, processed_content, custom_fields)
                            else:
                                processed_content = self.rule_engine.apply_redaction_rule(rule, processed_content)
                        
                        # 写入文件
                        try:
                            with open(output_path, 'w', encoding=encoding) as f:
                                f.write(processed_content)
                        except PermissionError:
                            failed_files.append(f"{filename} (文件权限错误: 无法保存到 {output_path})")
                            continue
                    
                    elif filename.lower().endswith('.pdf'):
                        # 处理PDF文件 - 直接使用PyMuPDF保持原格式
                        self.status_label.setText(f"正在处理PDF: {filename}")
                        QApplication.processEvents()

                        previous_input = getattr(self, 'input_file_path', None)
                        previous_output = getattr(self, 'output_file_path', None)

                        try:
                            if self.is_pdf_image_based(input_path):
                                failed_files.append(f"{filename} (图片型PDF暂不支持)")
                            else:
                                self.reset_pdf_state()
                                self.input_file_path = input_path
                                self.output_file_path = output_path

                                display_text = self.load_pdf_with_pymupdf(input_path)
                                if display_text is None or not self.pdf_char_map:
                                    failed_files.append(f"{filename} (无法解析PDF文本)")
                                else:
                                    operations, _ = self.auto_redact_pdf()

                                    if not operations:
                                        doc = fitz.open(input_path)
                                        doc.save(output_path)
                                        doc.close()
                                    else:
                                        self.pdf_pending_redactions = operations
                                        self.save_pdf_changes()

                                    processed_count += 1
                        except Exception as e:
                            failed_files.append(f"{filename} (PDF处理失败: {str(e)})")
                        finally:
                            self.reset_pdf_state()
                            if previous_input is not None:
                                self.input_file_path = previous_input
                            else:
                                if hasattr(self, 'input_file_path'):
                                    del self.input_file_path
                            if previous_output is not None:
                                self.output_file_path = previous_output
                            else:
                                if hasattr(self, 'output_file_path'):
                                    del self.output_file_path

                        self.progress_bar.setValue(i + 1)
                        QApplication.processEvents()
                        continue
                    
                    processed_count += 1
                    self.progress_bar.setValue(i + 1)
                    QApplication.processEvents()
                    
                except Exception as e:
                    failed_files.append(f"{filename} ({str(e)})")
            
            # 显示处理结果
            self.progress_bar.setValue(len(input_files))
            result_msg = f"✅ 批量处理完成！\n\n📊 处理统计：\n• 成功处理：{processed_count} 个文件\n• 输出位置：{output_dir}"
            
            if failed_files:
                result_msg += f"\n• 失败文件：{len(failed_files)} 个"
                if len(failed_files) <= 3:
                    result_msg += f"\n\n❌ 失败详情：\n" + "\n".join([f"• {f}" for f in failed_files])
                else:
                    result_msg += f"\n\n❌ 失败详情（前3个）：\n" + "\n".join([f"• {f}" for f in failed_files[:3]])
                    result_msg += f"\n... 还有 {len(failed_files)-3} 个文件失败"
            
            QMessageBox.information(self, "批量处理完成", result_msg)
            self.status_label.setText("✅ 批量处理完成")
            self.progress_bar.setValue(0)
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"批量处理失败: {str(e)}")
            self.status_label.setText("❌ 批量处理失败")
            self.progress_bar.setValue(0)

    def show_text_context_menu(self, position):
        """显示文本编辑器的右键菜单"""
        # 在自动脱敏模式下不显示右键菜单
        if self.mode_combo.currentIndex() == 1:  # 自动脱敏模式
            return
        self.text_menu.exec(self.text_edit.mapToGlobal(position))
        
    def mark_text_redaction(self):
        """标记选中的文本为脱敏内容"""
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            if not selected_text.strip():
                QMessageBox.warning(self, "警告", "选中的文本为空")
                return
                
            # 使用内置算法进行脱敏
            redacted_text = self.smart_redact_text(selected_text)
            
            # 记录撤销历史
            start_pos = cursor.selectionStart()
            end_pos = cursor.selectionEnd()
            self.text_redaction_history.append({
                'start': start_pos,
                'end': start_pos + len(redacted_text),
                'original': selected_text,
                'redacted': redacted_text,
                'timestamp': self.get_current_timestamp(),
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法',
                'position_desc': f"字符位置 {start_pos}-{end_pos}"
            })
            
            cursor.insertText(redacted_text)
            self.text_edit.setTextCursor(cursor)
        else:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的文本")

    def mark_text_redaction_all(self):
        """标记选中文本在文本编辑器中的所有相同内容为脱敏"""
        cursor = self.text_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            if not selected_text.strip():
                QMessageBox.warning(self, "警告", "选中的文本为空")
                return
                
            # 获取全文内容
            full_text = self.text_edit.toPlainText()
            
            # 使用内置算法进行脱敏
            redacted_text = self.smart_redact_text(selected_text)
            
            # 记录批量撤销历史
            count = full_text.count(selected_text)
            self.text_redaction_history.append({
                'type': 'replace_all',
                'original': selected_text,
                'redacted': redacted_text,
                'count': count,
                'full_original': full_text,
                'timestamp': self.get_current_timestamp(),
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法',
                'position_desc': '整个文档'
            })
            
            # 在全文中替换所有相同的内容
            new_text = full_text.replace(selected_text, redacted_text)
            
            # 更新文本内容
            self.text_edit.setPlainText(new_text)
            
            # 显示替换结果
            QMessageBox.information(self, "脱敏完成", f"已替换 {count} 处相同内容：\n原文：{selected_text}\n脱敏后：{redacted_text}")
        else:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的文本")
    
    def smart_redact_text(self, text):
        """内置算法脱敏文本内容 - 用于交互式脱敏"""
        import re
        
        # 检测文本类型并应用相应脱敏规则
        text = text.strip()
        
        # 身份证号（18位）
        if re.match(r'^\d{18}$', text):
            return text[:3] + "*" * 11 + text[-4:]
        
        # 手机号（11位数字）
        elif re.match(r'^1[3-9]\d{9}$', text):
            return text[:3] + "****" + text[-4:]
        
        # 邮箱地址
        elif '@' in text and re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', text):
            at_index = text.find('@')
            return text[0] + "*" * (at_index - 1) + text[at_index:]
        
        # 中文姓名（2-4个汉字）
        elif re.match(r'^[\u4e00-\u9fa5]{2,4}$', text):
            return text[0] + "*" * (len(text) - 1)
        
        # 银行卡号（16-19位数字）
        elif re.match(r'^\d{16,19}$', text):
            return text[:4] + "*" * (len(text) - 8) + text[-4:]
        
        # 默认脱敏：保留首尾，中间用*替换
        else:
            if len(text) <= 2:
                return "*" * len(text)
        
        # 对于其他任意文本，简单保留首尾字符
        if len(text) > 2:
            return text[0] + ("*" * (len(text) - 2)) + text[-1]
        return text

    def show_word_context_menu(self, position):
        """显示Word编辑器的右键菜单"""
        # 在自动脱敏模式下不显示右键菜单
        if self.mode_combo.currentIndex() == 1:  # 自动脱敏模式
            return
        self.word_menu.exec(self.word_edit.mapToGlobal(position))
        
    def mark_word_redaction(self):
        """标记选中的Word文档文本为脱敏内容"""
        cursor = self.word_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            if not selected_text.strip():
                QMessageBox.warning(self, "警告", "选中的文本为空")
                return
                
            # 使用内置算法进行脱敏
            redacted_text = self.smart_redact_text(selected_text)
            
            # 记录撤销历史
            start_pos = cursor.selectionStart()
            end_pos = cursor.selectionEnd()
            self.word_redaction_history.append({
                'start': start_pos,
                'end': start_pos + len(redacted_text),
                'original': selected_text,
                'redacted': redacted_text,
                'timestamp': self.get_current_timestamp(),
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法',
                'position_desc': f"字符位置 {start_pos}-{end_pos}"
            })
            
            cursor.insertText(redacted_text)
            self.word_edit.setTextCursor(cursor)
        else:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的文本")

    def mark_word_redaction_all(self):
        """标记选中文本在Word文档中的所有相同内容为脱敏"""
        cursor = self.word_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            if not selected_text.strip():
                QMessageBox.warning(self, "警告", "选中的文本为空")
                return
                
            # 获取全文内容
            full_text = self.word_edit.toPlainText()
            
            # 使用内置算法进行脱敏
            redacted_text = self.smart_redact_text(selected_text)
            
            # 记录批量撤销历史
            count = full_text.count(selected_text)
            self.word_redaction_history.append({
                'type': 'replace_all',
                'original': selected_text,
                'redacted': redacted_text,
                'count': count,
                'full_original': full_text,
                'timestamp': self.get_current_timestamp(),
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法',
                'position_desc': '整个文档'
            })
            
            # 在全文中替换所有相同的内容
            new_text = full_text.replace(selected_text, redacted_text)
            
            # 更新文本内容
            self.word_edit.setPlainText(new_text)
            
            # 显示替换结果
            QMessageBox.information(self, "脱敏完成", f"已替换 {count} 处相同内容：\n原文：{selected_text}\n脱敏后：{redacted_text}")
        else:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的文本")

    def setup_table_context_menu(self):
        """设置表格右键菜单"""
        self.table_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table_widget.customContextMenuRequested.connect(self.show_table_context_menu)
        
        self.table_menu = QMenu(self)
        # 添加菜单项
        self.redact_cell_action = QAction("脱敏选中单元格", self)
        self.redact_cell_action.triggered.connect(self.mark_cell_redaction)
        self.table_menu.addAction(self.redact_cell_action)
        
        self.redact_row_action = QAction("脱敏整行", self)
        self.redact_row_action.triggered.connect(self.mark_row_redaction)
        self.table_menu.addAction(self.redact_row_action)
        
        self.redact_col_action = QAction("脱敏整列", self)
        self.redact_col_action.triggered.connect(self.mark_column_redaction)
        self.table_menu.addAction(self.redact_col_action)
        
        # 添加全局查找替换脱敏功能
        self.table_redact_all_action = QAction("🔄 全表相同内容脱敏", self)
        self.table_redact_all_action.triggered.connect(self.mark_table_redaction_all)
        self.table_menu.addAction(self.table_redact_all_action)
        
        # 添加分隔符
        self.table_menu.addSeparator()
        
        # 添加撤销脱敏功能
        self.undo_redaction_action = QAction("撤销脱敏", self)
        self.undo_redaction_action.triggered.connect(self.undo_redaction)
        self.table_menu.addAction(self.undo_redaction_action)
        
        # 添加撤销当前区域脱敏功能
        self.excel_undo_current_action = QAction("撤销当前区域脱敏", self)
        self.excel_undo_current_action.triggered.connect(self.undo_current_excel_redaction)
        self.table_menu.addAction(self.excel_undo_current_action)

    def show_table_context_menu(self, position):
        """显示表格的右键菜单"""
        # 在自动脱敏模式下不显示右键菜单
        if self.mode_combo.currentIndex() == 1:  # 自动脱敏模式
            return
        
        # 保存右键点击的位置对应的单元格索引，用于整行/整列脱敏
        item = self.table_widget.itemAt(position)
        if item:
            self.current_right_click_row = item.row()
            self.current_right_click_col = item.column()
        else:
            self.current_right_click_row = -1
            self.current_right_click_col = -1
        self.table_menu.exec(self.table_widget.mapToGlobal(position))

    def mark_cell_redaction(self):
        """标记选中的单元格为脱敏内容"""
        selected_items = self.table_widget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的单元格")
            return
            
        # 获取用户确认
        reply = QMessageBox.question(self, "确认脱敏", 
                                   f"确定要对选中的 {len(selected_items)} 个单元格进行脱敏吗？\n请慎重操作！",
                                   QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
            
        # 记录脱敏的单元格数量和历史记录
        redacted_count = 0
        operation_history = []
        
        for item in selected_items:
            if item and item.text().strip():
                original_text = item.text().strip()
                redacted_text = self.smart_redact_text(original_text)
                if redacted_text != original_text:
                        # 记录历史
                        operation_history.append({
                            'row': item.row(),
                            'col': item.column(),
                            'original_text': original_text,
                            'redacted_text': redacted_text,
                            'original_background': item.background(),
                            'original_tooltip': item.toolTip(),
                            'timestamp': self.get_current_timestamp(),
                            'rule_name': '交互式脱敏',
                            'mode': '交互式脱敏',
                            'rule_type': '内置算法',
                            'position_desc': f"单元格 {self.get_excel_column_letter(item.column() + 1)}{item.row() + 1}"
                        })
                        
                        item.setText(redacted_text)
                        # 标记脱敏的单元格
                        item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                        item.setToolTip(f"已脱敏 - 原文本: {original_text}")
                        redacted_count += 1
        
        # 记录撤销历史
        if operation_history:
            self.excel_redaction_history.append({
                'type': 'cell_redaction',
                'operations': operation_history
            })
        
        if redacted_count > 0:
            QMessageBox.information(self, "脱敏完成", f"已成功脱敏 {redacted_count} 个单元格")
        else:
            QMessageBox.information(self, "提示", "没有找到需要脱敏的内容")

    def mark_row_redaction(self):
        """标记整行为脱敏内容（基于右键点击的单元格位置）"""
        # 如果有右键点击的单元格，则脱敏该行
        if hasattr(self, 'current_right_click_row') and self.current_right_click_row >= 0:
            target_row = self.current_right_click_row
            total_cells = self.table_widget.columnCount()
            
            # 获取用户确认
            effective_cells = (total_cells - 1) if total_cells > 1 else total_cells
            reply = QMessageBox.question(self, "确认脱敏", 
                                       f"确定要对第 {target_row + 1} 行（{effective_cells} 个单元格，跳过第1列）进行脱敏吗？\n请慎重操作！",
                                       QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
            
            if reply != QMessageBox.StandardButton.Yes:
                return
                
            # 记录脱敏的单元格数量和历史记录
            redacted_count = 0
            operation_history = []
            
            # 遍历该行的所有列（跳过第一列，通常是序号列）
            start_col = 1 if self.table_widget.columnCount() > 1 else 0
            for col in range(start_col, self.table_widget.columnCount()):
                item = self.table_widget.item(target_row, col)
                if item and item.text().strip():
                    original_text = item.text().strip()
                    redacted_text = self.smart_redact_text(original_text)
                    if redacted_text != original_text:
                        # 记录历史
                        operation_history.append({
                            'row': target_row,
                            'col': col,
                            'original_text': original_text,
                            'redacted_text': redacted_text,
                            'original_background': item.background(),
                            'original_tooltip': item.toolTip()
                        })
                        
                        item.setText(redacted_text)
                        # 标记脱敏的单元格
                        item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                        item.setToolTip(f"已脱敏 - 原文本: {original_text}")
                        redacted_count += 1
            
            # 记录撤销历史
            if operation_history:
                self.excel_redaction_history.append({
                    'type': 'row_redaction',
                    'operations': operation_history
                })
            
            if redacted_count > 0:
                QMessageBox.information(self, "脱敏完成", f"已成功脱敏第 {target_row + 1} 行的 {redacted_count} 个单元格")
            else:
                QMessageBox.information(self, "提示", f"第 {target_row + 1} 行没有找到需要脱敏的内容")
        else:
            # 回退到原来的逻辑：处理用户选中的范围
            selected_ranges = self.table_widget.selectedRanges()
            if not selected_ranges:
                QMessageBox.warning(self, "警告", "请先点击要脱敏的单元格")
                return
                
            # 计算总单元格数量
            total_cells = 0
            for r in selected_ranges:
                total_cells += (r.bottomRow() - r.topRow() + 1) * self.table_widget.columnCount()
                
            # 获取用户确认
            reply = QMessageBox.question(self, "确认脱敏", 
                                       f"确定要对选中的行（约 {total_cells} 个单元格）进行脱敏吗？\n请慎重操作！",
                                       QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
            
            if reply != QMessageBox.StandardButton.Yes:
                return
                
            # 记录脱敏的单元格数量和历史记录
            redacted_count = 0
            operation_history = []
            
            for r in selected_ranges:
                for row in range(r.topRow(), r.bottomRow() + 1):
                    for col in range(self.table_widget.columnCount()):
                        item = self.table_widget.item(row, col)
                        if item and item.text().strip():
                            original_text = item.text().strip()
                            redacted_text = self.smart_redact_text(original_text)
                            if redacted_text != original_text:
                                # 记录历史
                                operation_history.append({
                                    'row': row,
                                    'col': col,
                                    'original_text': original_text,
                                    'redacted_text': redacted_text,
                                    'original_background': item.background(),
                                    'original_tooltip': item.toolTip()
                                })
                                
                                item.setText(redacted_text)
                                # 标记脱敏的单元格
                                item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                                item.setToolTip(f"已脱敏 - 原文本: {original_text}")
                                redacted_count += 1
            
            # 记录撤销历史
            if operation_history:
                self.excel_redaction_history.append({
                    'type': 'row_range_redaction',
                    'operations': operation_history
                })
            
            if redacted_count > 0:
                QMessageBox.information(self, "脱敏完成", f"已成功脱敏 {redacted_count} 个单元格")
            else:
                QMessageBox.information(self, "提示", "没有找到需要脱敏的内容")

    def mark_column_redaction(self):
        """标记整列为脱敏内容（基于右键点击的单元格位置）"""
        # 如果有右键点击的单元格，则脱敏该列
        if hasattr(self, 'current_right_click_col') and self.current_right_click_col >= 0:
            target_col = self.current_right_click_col
            total_cells = self.table_widget.rowCount()
            
            # 获取用户确认
            effective_cells = (total_cells - 1) if total_cells > 1 else total_cells
            reply = QMessageBox.question(self, "确认脱敏", 
                                       f"确定要对第 {target_col + 1} 列（{effective_cells} 个单元格，跳过第1行）进行脱敏吗？\n请慎重操作！",
                                       QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
            
            if reply != QMessageBox.StandardButton.Yes:
                return
                
            # 记录脱敏的单元格数量和历史记录
            redacted_count = 0
            operation_history = []
            
            # 遍历该列的所有行（跳过第一行，通常是表头）
            start_row = 1 if self.table_widget.rowCount() > 1 else 0
            for row in range(start_row, self.table_widget.rowCount()):
                item = self.table_widget.item(row, target_col)
                if item and item.text().strip():
                    original_text = item.text().strip()
                    redacted_text = self.smart_redact_text(original_text)
                    if redacted_text != original_text:
                        # 记录历史
                        operation_history.append({
                            'row': row,
                            'col': target_col,
                            'original_text': original_text,
                            'redacted_text': redacted_text,
                            'original_background': item.background(),
                            'original_tooltip': item.toolTip()
                        })
                        
                        item.setText(redacted_text)
                        # 标记脱敏的单元格
                        item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                        item.setToolTip(f"已脱敏 - 原文本: {original_text}")
                        redacted_count += 1
            
            # 记录撤销历史
            if operation_history:
                self.excel_redaction_history.append({
                    'type': 'column_redaction',
                    'operations': operation_history
                })
            
            if redacted_count > 0:
                QMessageBox.information(self, "脱敏完成", f"已成功脱敏第 {target_col + 1} 列的 {redacted_count} 个单元格")
            else:
                QMessageBox.information(self, "提示", f"第 {target_col + 1} 列没有找到需要脱敏的内容")
        else:
            # 回退到原来的逻辑：处理用户选中的范围
            selected_ranges = self.table_widget.selectedRanges()
            if not selected_ranges:
                QMessageBox.warning(self, "警告", "请先点击要脱敏的单元格")
                return
                
            # 获取用户确认
            reply = QMessageBox.question(self, "确认脱敏", 
                                       "确定要对选中的单元格进行脱敏吗？\n请慎重操作！",
                                       QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
            
            if reply != QMessageBox.StandardButton.Yes:
                return
                
            # 记录脱敏的单元格数量和历史记录
            redacted_count = 0
            operation_history = []
            
            # 只处理选中的单元格范围
            for r in selected_ranges:
                for row in range(r.topRow(), r.bottomRow() + 1):
                    for col in range(r.leftColumn(), r.rightColumn() + 1):
                        item = self.table_widget.item(row, col)
                        if item and item.text().strip():
                            original_text = item.text().strip()
                            redacted_text = self.smart_redact_text(original_text)
                            if redacted_text != original_text:
                                # 记录历史
                                operation_history.append({
                                    'row': row,
                                    'col': col,
                                    'original_text': original_text,
                                    'redacted_text': redacted_text,
                                    'original_background': item.background(),
                                    'original_tooltip': item.toolTip()
                                })
                                
                                item.setText(redacted_text)
                                # 标记脱敏的单元格
                                item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                                item.setToolTip(f"已脱敏 - 原文本: {original_text}")
                                redacted_count += 1
            
            # 记录撤销历史
            if operation_history:
                self.excel_redaction_history.append({
                    'type': 'column_range_redaction',
                    'operations': operation_history
                })
            
            if redacted_count > 0:
                QMessageBox.information(self, "脱敏完成", f"已成功脱敏 {redacted_count} 个单元格")
            else:
                QMessageBox.information(self, "提示", "没有找到需要脱敏的内容")

    def mark_table_redaction_all(self):
        """Excel表格全局查找替换脱敏（类似文本和Word的功能）"""
        # 获取当前选中的单元格内容作为查找目标
        selected_items = self.table_widget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选择包含要脱敏内容的单元格")
            return
        
        # 使用第一个选中单元格的内容作为查找目标
        target_text = selected_items[0].text().strip()
        if not target_text:
            QMessageBox.warning(self, "警告", "选中的单元格内容为空")
            return
        
        # 询问用户确认
        reply = QMessageBox.question(self, "确认全表脱敏", 
                                   f"确定要对表格中所有包含「{target_text}」的单元格进行脱敏吗？\n请慎重操作！",
                                   QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # 记录脱敏的单元格数量和历史记录
        redacted_count = 0
        operation_history = []
        
        # 遍历整个表格查找匹配的内容
        for row in range(self.table_widget.rowCount()):
            for col in range(self.table_widget.columnCount()):
                item = self.table_widget.item(row, col)
                if item and target_text in item.text():
                    original_text = item.text()
                    # 使用内置算法替换目标文本
                    redacted_target = self.smart_redact_text(target_text)
                    redacted_text = original_text.replace(target_text, redacted_target)
                    
                    if redacted_text != original_text:
                        # 记录历史
                        operation_history.append({
                            'row': row,
                            'col': col,
                            'original_text': original_text,
                            'redacted_text': redacted_text,
                            'original_background': item.background(),
                            'original_tooltip': item.toolTip()
                        })
                        
                        item.setText(redacted_text)
                        # 标记脱敏的单元格
                        item.setBackground(QColor(255, 235, 235))  # 浅红色背景
                        item.setToolTip(f"已脱敏 - 原文本: {original_text}")
                        redacted_count += 1
        
        # 记录撤销历史
        if operation_history:
            self.excel_redaction_history.append({
                'type': 'table_find_replace_redaction',
                'operations': operation_history,
                'target_text': target_text
            })
        
        if redacted_count > 0:
            QMessageBox.information(self, "脱敏完成", f"已成功脱敏表格中包含「{target_text}」的 {redacted_count} 个单元格")
        else:
            QMessageBox.information(self, "提示", f"表格中没有找到包含「{target_text}」的其他单元格")

    def setup_styles(self):
        # 保持原有配色方案
        palette = self.palette()
        gradient = QLinearGradient(0, 0, 0, self.height())
        gradient.setColorAt(0, QColor(230, 245, 255))
        gradient.setColorAt(1, QColor(180, 220, 255))
        palette.setBrush(QPalette.Window, QBrush(gradient))
        self.setPalette(palette)

        # 统一控件样式
        self.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 1px solid #3498db;
                border-radius: 5px;
                margin-top: 1ex;
                background-color: rgba(255, 255, 255, 180);
            }
            QGroupBox::title {
                color: #2980b9;
                subcontrol-origin: margin;
                subcontrol-position: top center;
                padding: 0 5px;
            }
            QPushButton {
                background-color: transparent;
                color: #3498db;
                border: 2px solid #3498db;
                padding: 8px 16px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: rgba(52, 152, 219, 0.14);
            }
            QPushButton:pressed {
                background-color: rgba(52, 152, 219, 0.24);
            }
            QPushButton:disabled {
                color: rgba(52, 152, 219, 0.35);
                border-color: rgba(52, 152, 219, 0.35);
                background-color: transparent;
            }
            QTextEdit, QLabel {
                font-size: 11pt;
            }
        """)

    def undo_redaction(self):
        """撤销Excel表格的脱敏操作（基于历史记录的逐步撤销）"""
        if not self.excel_redaction_history:
            QMessageBox.information(self, "提示", "没有可撤销的操作")
            return
        
        # 获取最后一次操作
        last_operation = self.excel_redaction_history.pop()
        
        # 根据操作类型显示确认信息
        op_type = last_operation['type']
        op_operations = last_operation['operations']
        
        type_names = {
            'cell_redaction': '单元格脱敏',
            'row_redaction': '行脱敏',
            'column_redaction': '列脱敏'
        }
        
        op_name = type_names.get(op_type, '脱敏操作')
        
        # 获取用户确认
        reply = QMessageBox.question(self, "确认撤销", 
                                   f"确定要撤销最后一次{op_name}操作吗？\n将影响 {len(op_operations)} 个单元格",
                                   QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            # 用户取消，重新添加到历史记录
            self.excel_redaction_history.append(last_operation)
            return
        
        # 执行撤销
        restored_count = 0
        for operation in op_operations:
            item = self.table_widget.item(operation['row'], operation['col'])
            if item:
                # 恢复原始文本
                item.setText(operation['original_text'])
                # 恢复原始背景色
                item.setBackground(operation['original_background'])
                # 恢复原始工具提示
                item.setToolTip(operation['original_tooltip'])
                restored_count += 1
        
        if restored_count > 0:
            QMessageBox.information(self, "撤销完成", f"已成功撤销{op_name}，恢复了 {restored_count} 个单元格")
        else:
            QMessageBox.information(self, "提示", "撤销操作未能恢复任何单元格")
    
    def undo_text_redaction(self):
        """撤销文本编辑器的脱敏操作"""
        if not self.text_redaction_history:
            QMessageBox.information(self, "提示", "没有可撤销的脱敏操作")
            return
        
        # 获取用户确认
        reply = QMessageBox.question(self, "确认撤销", 
                                   f"确定要撤销最后 {len(self.text_redaction_history)} 个脱敏操作吗？",
                                   QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # 逆序撤销操作（后进先出）
        restored_count = 0
        while self.text_redaction_history:
            operation = self.text_redaction_history.pop()
            
            if operation.get('type') == 'replace_all':
                # 批量替换的撤销
                self.text_edit.setPlainText(operation['full_original'])
                restored_count += operation['count']
            else:
                # 单个选择替换的撤销
                cursor = self.text_edit.textCursor()
                cursor.setPosition(operation['start'])
                cursor.setPosition(operation['end'], cursor.KeepAnchor)
                cursor.insertText(operation['original'])
                restored_count += 1
        
        if restored_count > 0:
            QMessageBox.information(self, "撤销完成", f"已成功撤销 {restored_count} 处脱敏内容")
    
    def undo_word_redaction(self):
        """撤销Word文档编辑器的脱敏操作"""
        if not self.word_redaction_history:
            QMessageBox.information(self, "提示", "没有可撤销的脱敏操作")
            return
        
        # 获取用户确认
        reply = QMessageBox.question(self, "确认撤销", 
                                   f"确定要撤销最后 {len(self.word_redaction_history)} 个脱敏操作吗？",
                                   QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # 逆序撤销操作（后进先出）
        restored_count = 0
        while self.word_redaction_history:
            operation = self.word_redaction_history.pop()
            
            if operation.get('type') == 'replace_all':
                # 批量替换的撤销
                self.word_edit.setPlainText(operation['full_original'])
                restored_count += operation['count']
            else:
                # 单个选择替换的撤销
                cursor = self.word_edit.textCursor()
                cursor.setPosition(operation['start'])
                cursor.setPosition(operation['end'], cursor.KeepAnchor)
                cursor.insertText(operation['original'])
                restored_count += 1
        
        if restored_count > 0:
            QMessageBox.information(self, "撤销完成", f"已成功撤销 {restored_count} 处脱敏内容")

    # PDF文档相关方法
    def show_pdf_context_menu(self, position):
        """显示PDF编辑器的右键菜单"""
        # 在自动脱敏模式下不显示右键菜单
        if self.mode_combo.currentIndex() == 1:  # 自动脱敏模式
            return
        self.pdf_menu.exec(self.pdf_edit.mapToGlobal(position))
        
    def mark_pdf_redaction(self):
        """标记选中的PDF文档文本为脱敏内容"""
        cursor = self.pdf_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            normalized_text = selected_text.replace('\u2029', '\n')
            if not normalized_text.strip():
                QMessageBox.warning(self, "警告", "选中的文本为空")
                return
            
            # 使用内置算法进行脱敏，同时保持原有长度
            redacted_text = self.generate_redacted_text(normalized_text)

            start_pos = cursor.selectionStart()
            end_pos = cursor.selectionEnd()
            if end_pos > len(self.pdf_char_map):
                QMessageBox.warning(self, "警告", "选中范围超出PDF解析范围")
                return

            if not self.ensure_pdf_font_context():
                QMessageBox.warning(self, "提示", "未能准备PDF字体上下文，请重新载入PDF后再试")
                return

            if 0 <= start_pos < len(self.pdf_char_map):
                selected_font = self.pdf_char_map[start_pos].get('font')
                alias = self.get_pdf_font_alias(selected_font)
                print(f"PDF交互脱敏字体: 原始={selected_font} -> 别名={alias}")

            current_text = ''.join(entry.get('char', '') for entry in self.pdf_char_map)
            updated_text = current_text[:start_pos] + redacted_text + current_text[end_pos:]

            base_context = {
                'type': 'single',
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法'
            }

            def _context_builder(s, e, original_segment, replacement_segment):
                original_end = s + len(original_segment)
                return {'position_desc': f"字符位置 {s}-{original_end}"}

            operations = self.build_pdf_operations_from_text(
                current_text,
                updated_text,
                base_context,
                context_callback=_context_builder
            )

            if not operations:
                QMessageBox.warning(self, "警告", "未能定位所选文本的坐标，无法完成脱敏")
                return

            operation = operations[0]
            self.pdf_redaction_history.append(operation)
            self.pdf_pending_redactions.append(operation)

            final_text = ''.join(entry.get('char', '') for entry in self.pdf_char_map)
            self.pdf_edit.blockSignals(True)
            self.pdf_edit.setPlainText(final_text.replace('\n', '\u2029'))
            self.pdf_edit.blockSignals(False)

            new_cursor = self.pdf_edit.textCursor()
            new_cursor.setPosition(operation.get('start', start_pos))
            new_cursor.setPosition(operation.get('end', start_pos + len(redacted_text)), new_cursor.KeepAnchor)
            self.pdf_edit.setTextCursor(new_cursor)
        else:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的文本")

    def mark_pdf_redaction_all(self):
        """标记选中文本在PDF文档中的所有相同内容为脱敏"""
        cursor = self.pdf_edit.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            normalized_text = selected_text.replace('\u2029', '\n')
            if not normalized_text.strip():
                QMessageBox.warning(self, "警告", "选中的文本为空")
                return
            
            full_text = ''.join(entry.get('char', '') for entry in self.pdf_char_map)
            occurrences = []
            search_pos = 0
            target_len = len(normalized_text)
            while True:
                idx = full_text.find(normalized_text, search_pos)
                if idx == -1:
                    break
                occurrences.append(idx)
                search_pos = idx + target_len

            if len(occurrences) <= 1:
                QMessageBox.information(self, "提示", "该文本在文档中只出现一次，建议使用单独脱敏")
                return

            preview_text = normalized_text[:20] + ('...' if len(normalized_text) > 20 else '')
            reply = QMessageBox.question(
                self,
                "确认脱敏",
                f"在文档中找到 {len(occurrences)} 处相同文本 \"{preview_text}\"\n确定要全部脱敏吗？",
                QMessageBox.StandardButton.Yes,
                QMessageBox.StandardButton.No
            )

            if reply != QMessageBox.StandardButton.Yes:
                return

            redacted_text = self.generate_redacted_text(normalized_text)

            if not self.ensure_pdf_font_context():
                QMessageBox.warning(self, "提示", "未能准备PDF字体上下文，请重新载入PDF后再试")
                return

            if occurrences:
                probe_index = occurrences[0]
                if 0 <= probe_index < len(self.pdf_char_map):
                    probe_font = self.pdf_char_map[probe_index].get('font')
                    alias = self.get_pdf_font_alias(probe_font)
                    print(f"PDF批量脱敏字体: 原始={probe_font} -> 别名={alias}")

            new_full_text = full_text.replace(normalized_text, redacted_text)

            base_context = {
                'type': 'batch',
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法'
            }

            operations = self.build_pdf_operations_from_text(full_text, new_full_text, base_context)
            if not operations:
                QMessageBox.warning(self, "警告", "未能定位文本坐标，批量脱敏已取消")
                return

            if len(operations) != len(occurrences):
                for op in operations:
                    self.restore_pdf_characters(op.get('char_backup'))
                self.pdf_display_text = full_text
                QMessageBox.warning(self, "警告", "部分文本未能定位，批量脱敏已取消")
                return

            segments_all = []
            backup_all = []
            position_list = []
            for op in operations:
                position_list.append(op.get('start', 0))
                segments_all.extend(op.get('segments', []))
                backup_all.extend(list(op.get('char_backup', [])))

            final_text = ''.join(entry.get('char', '') for entry in self.pdf_char_map)

            aggregated_operation = {
                'type': 'replace_all',
                'target': normalized_text,
                'replacement': redacted_text,
                'original': normalized_text,
                'redacted': redacted_text,
                'count': len(operations),
                'positions': position_list,
                'segments': segments_all,
                'char_backup': backup_all,
                'full_original': full_text,
                'full_new': final_text,
                'timestamp': self.get_current_timestamp(),
                'rule_name': '交互式脱敏',
                'mode': '交互式脱敏',
                'rule_type': '内置算法'
            }

            self.pdf_redaction_history.append(aggregated_operation)
            self.pdf_pending_redactions.append(aggregated_operation)

            self.pdf_edit.blockSignals(True)
            self.pdf_edit.setPlainText(final_text.replace('\n', '\u2029'))
            self.pdf_edit.blockSignals(False)

            QMessageBox.information(self, "脱敏完成", f"已成功脱敏 {len(operations)} 处相同内容")
        else:
            QMessageBox.warning(self, "警告", "请先选择要脱敏的文本")
    
    def undo_pdf_redaction(self):
        """撤销PDF文档编辑器的脱敏操作"""
        if not self.pdf_redaction_history:
            QMessageBox.information(self, "提示", "没有可撤销的脱敏操作")
            return
        
        # 获取用户确认
        reply = QMessageBox.question(self, "确认撤销", 
                                   f"确定要撤销最后 {len(self.pdf_redaction_history)} 个脱敏操作吗？",
                                   QMessageBox.StandardButton.Yes, QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # 逆序撤销操作（后进先出）
        restored_count = 0
        while self.pdf_redaction_history:
            operation = self.pdf_redaction_history.pop()
            if operation in self.pdf_pending_redactions:
                self.pdf_pending_redactions.remove(operation)

            # 恢复字符映射
            for backup in operation.get('char_backup', []):
                index = backup.get('index')
                original_char = backup.get('char')
                if index is not None and 0 <= index < len(self.pdf_char_map):
                    self.pdf_char_map[index]['char'] = original_char

            if operation.get('type') == 'replace_all':
                original_text = operation.get('full_original')
                if original_text is not None:
                    self.pdf_edit.blockSignals(True)
                    self.pdf_edit.setPlainText(original_text)
                    self.pdf_edit.blockSignals(False)
                restored_count += operation.get('count', 0)
            else:
                cursor = self.pdf_edit.textCursor()
                cursor.setPosition(operation.get('start', 0))
                cursor.setPosition(operation.get('end', 0), cursor.KeepAnchor)
                original_text = operation.get('original', '')
                cursor.insertText(original_text.replace('\n', '\u2029'))
                restored_count += 1
        
        if restored_count > 0:
            QMessageBox.information(self, "撤销完成", f"已成功撤销 {restored_count} 处脱敏内容")

    # 区域撤销功能已移除，仅保留单步撤销

    def undo_current_excel_redaction(self):
        """撤销当前选中区域的Excel脱敏"""
        # 获取当前选中的单元格或区域
        selected_items = self.table_widget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "警告", "请先选中需要撤销脱敏的单元格")
            return
        
        # 如果只选中一个单元格，进行单元格撤销
        if len(selected_items) == 1:
            self.undo_single_cell_redaction(selected_items[0])
        else:
            # 选中多个单元格，进行区域撤销
            self.undo_region_redaction(selected_items)

    def show_export_log_dialog(self):
        """显示导出日志对话框，让用户选择是否导出日志"""
        # 统计脱敏记录数量
        total_records = 0
        text_records = len(self.text_redaction_history)
        word_records = len(self.word_redaction_history)
        pdf_records = len(self.pdf_redaction_history)
        excel_records = sum(len(entry.get('operations', [])) for entry in self.excel_redaction_history)
        total_records = text_records + word_records + pdf_records + excel_records
        
        if total_records == 0:
            # 如果没有脱敏记录，显示提示信息
            QMessageBox.information(self, "📋 脱敏日志", 
                "当前没有脱敏操作记录。\n\n"
                "💡 提示：只有进行了脱敏操作（如选中文本/单元格右键脱敏）才会产生日志记录。")
            return False
        
        # 创建导出日志选择对话框
        dialog = QDialog(self)
        dialog.setWindowTitle("📋 导出脱敏日志")
        dialog.setModal(True)
        dialog.resize(500, 300)
        layout = QVBoxLayout()
        
        # 信息说明
        info_label = QLabel(f"""
<div style='padding: 10px; background-color: #f8f9fa; border-radius: 5px; border-left: 4px solid #007bff;'>
<h3 style='margin: 0; color: #007bff;'>📊 脱敏操作统计</h3>
<p style='margin: 5px 0;'><b>文本脱敏记录：</b>{text_records} 条</p>
<p style='margin: 5px 0;'><b>Word脱敏记录：</b>{word_records} 条</p>
<p style='margin: 5px 0;'><b>PDF脱敏记录：</b>{pdf_records} 条</p>
<p style='margin: 5px 0;'><b>Excel脱敏记录：</b>{excel_records} 条</p>
<p style='margin: 10px 0 0 0; font-weight: bold; color: #28a745;'>总计：{total_records} 条脱敏记录</p>
</div>

<div style='padding: 10px; margin-top: 10px; background-color: #fff3cd; border-radius: 5px; border-left: 4px solid #ffc107;'>
<p style='margin: 0; color: #856404;'><b>💡 导出说明：</b>日志将包含原始内容、脱敏后内容、位置信息、使用规则等详细信息，便于审核和备案。</p>
</div>
        """)
        info_label.setWordWrap(True)
        layout.addWidget(info_label)
        
        # 导出格式说明
        format_group = QGroupBox("📁 导出格式")
        format_layout = QVBoxLayout()
        
        format_info = QLabel("� PDF格式 (.pdf) - 横向页面，适合打印和存档")
        format_info.setStyleSheet("color: #007bff; font-weight: bold; padding: 5px;")
        format_layout.addWidget(format_info)
        
        format_group.setLayout(format_layout)
        layout.addWidget(format_group)
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        
        export_btn = QPushButton("导出日志")
        self.set_hollow_button(export_btn, "#28a745", font_size="14px", padding="10px 20px")
        export_btn.clicked.connect(lambda: self.export_redaction_log(dialog))
        
        skip_btn = QPushButton("跳过")
        self.set_hollow_button(skip_btn, "#6c757d", font_size="14px", padding="10px 20px")
        skip_btn.clicked.connect(dialog.reject)
        
        btn_layout.addWidget(export_btn)
        btn_layout.addWidget(skip_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        result = dialog.exec_()
        return result == QDialog.Accepted

    def export_redaction_log(self, dialog):
        """导出脱敏日志"""
        try:
            # 获取导出路径（文件名不能包含冒号，所以用下划线代替）
            import datetime
            file_timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"脱敏日志_{file_timestamp}"
            
            # 导出PDF格式
            file_path, _ = QFileDialog.getSaveFileName(
                dialog,
                "导出PDF日志",
                f"{default_filename}.pdf",
                "PDF文件 (*.pdf);;所有文件 (*)"
            )
            
            if file_path:
                self.export_to_pdf(file_path)
                dialog.accept()
            
        except Exception as e:
            QMessageBox.critical(dialog, "导出失败", f"导出日志时发生错误：{str(e)}")

    def get_current_timestamp(self):
        """获取当前时间戳字符串"""
        import datetime
        return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def export_to_pdf(self, file_path):
        """导出日志到PDF文件 - 横向页面格式"""
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from datetime import datetime
            import os

            # 注册中文字体
            try:
                # 尝试使用系统自带的中文字体
                font_paths = [
                    'C:/Windows/Fonts/simsun.ttc',  # 宋体
                    'C:/Windows/Fonts/simhei.ttf',  # 黑体
                    'C:/Windows/Fonts/simkai.ttf',  # 楷体
                    'C:/Windows/Fonts/msyh.ttc',    # 微软雅黑
                ]
                font_registered = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            font_registered = True
                            break
                        except:
                            continue
                
                if not font_registered:
                    # 如果没有找到中文字体，使用默认字体
                    font_name = 'Helvetica'
                else:
                    font_name = 'ChineseFont'
            except:
                font_name = 'Helvetica'

            # 收集所有日志记录
            all_records = []
            input_file_path = getattr(self, 'input_file_path', '')

            # 处理文本脱敏记录
            for i, record in enumerate(self.text_redaction_history):
                rule_name = record.get('rule_name', '交互式脱敏')
                mode = record.get('mode', '交互式脱敏')
                if record.get('type') == 'replace_all':
                    final_rule = '全文替换' if mode == '交互式脱敏' else rule_name
                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'TXT文本',
                        mode,
                        '整个文档',
                        self.truncate_text(record.get('original', ''), 20),
                        self.truncate_text(record.get('redacted', ''), 20),
                        f"{record.get('count', 1)} 处",
                        self.truncate_text(final_rule, 15),
                        record.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])
                else:
                    final_rule = '选中脱敏' if mode == '交互式脱敏' else rule_name
                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'TXT文本',
                        mode,
                        f"字符位置 {record.get('start', 0)}-{record.get('end', 0)}",
                        self.truncate_text(record.get('original', ''), 20),
                        self.truncate_text(record.get('redacted', ''), 20),
                        '1 处',
                        self.truncate_text(final_rule, 15),
                        record.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])

            # 处理Word脱敏记录
            for i, record in enumerate(self.word_redaction_history):
                rule_name = record.get('rule_name', '交互式脱敏')
                mode = record.get('mode', '交互式脱敏')
                if record.get('type') == 'replace_all':
                    final_rule = '全文替换' if mode == '交互式脱敏' else rule_name
                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'Word文档',
                        mode,
                        '整个文档',
                        self.truncate_text(record.get('original', ''), 20),
                        self.truncate_text(record.get('redacted', ''), 20),
                        f"{record.get('count', 1)} 处",
                        self.truncate_text(final_rule, 15),
                        record.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])
                else:
                    final_rule = '选中脱敏' if mode == '交互式脱敏' else rule_name
                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'Word文档',
                        mode,
                        f"字符位置 {record.get('start', 0)}-{record.get('end', 0)}",
                        self.truncate_text(record.get('original', ''), 20),
                        self.truncate_text(record.get('redacted', ''), 20),
                        '1 处',
                        self.truncate_text(final_rule, 15),
                        record.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])

            # 处理PDF脱敏记录
            for i, record in enumerate(self.pdf_redaction_history):
                rule_name = record.get('rule_name', '交互式脱敏')
                mode = record.get('mode', '交互式脱敏')
                original_text = record.get('original', record.get('target', ''))
                redacted_text = record.get('redacted', record.get('replacement', ''))
                if record.get('type') == 'replace_all':
                    final_rule = '全文替换' if mode == '交互式脱敏' else rule_name
                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'PDF文档',
                        mode,
                        '整个文档',
                        self.truncate_text(original_text, 20),
                        self.truncate_text(redacted_text, 20),
                        f"{record.get('count', 1)} 处",
                        self.truncate_text(final_rule, 15),
                        record.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])
                else:
                    final_rule = '选中脱敏' if mode == '交互式脱敏' else rule_name
                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'PDF文档',
                        mode,
                        f"字符位置 {record.get('start', 0)}-{record.get('end', 0)}",
                        self.truncate_text(original_text, 20),
                        self.truncate_text(redacted_text, 20),
                        '1 处',
                        self.truncate_text(final_rule, 15),
                        record.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])

            # 处理Excel脱敏记录
            for entry in self.excel_redaction_history:
                entry_type = entry.get('type', 'unknown')
                operations = entry.get('operations', [])

                type_map = {
                    'cell_redaction': '单元格脱敏',
                    'row_redaction': '行脱敏',
                    'column_redaction': '列脱敏',
                    'table_find_replace_redaction': '全表替换',
                    'auto_rule_redaction': '自动规则脱敏'
                }
                operation_type = type_map.get(entry_type, 'Excel脱敏')

                for operation in operations:
                    row = operation.get('row', 0) + 1
                    col = operation.get('col', 0) + 1
                    col_letter = self.get_excel_column_letter(col)
                    
                    rule_name = operation.get('rule_name', None)
                    mode = operation.get('mode', None)
                    
                    if entry_type == 'auto_rule_redaction':
                        if not rule_name:
                            if 'rule' in operation and hasattr(operation['rule'], 'name'):
                                final_rule = operation['rule'].name
                            else:
                                final_rule = '自动规则脱敏'
                        else:
                            final_rule = rule_name
                        if not mode:
                            mode = '自动规则脱敏'
                    else:
                        final_rule = operation_type
                        if not mode:
                            mode = '交互式脱敏'

                    all_records.append([
                        str(len(all_records) + 1),
                        self.truncate_text(input_file_path, 25),
                        'Excel表格',
                        mode,
                        f"单元格 {col_letter}{row}",
                        self.truncate_text(operation.get('original_text', ''), 20),
                        self.truncate_text(operation.get('redacted_text', ''), 20),
                        '1 个单元格',
                        self.truncate_text(final_rule, 15),
                        operation.get('timestamp', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    ])

            if all_records:
                # 创建PDF文档 - 横向页面
                doc = SimpleDocTemplate(file_path, pagesize=landscape(A4))
                elements = []

                # 设置样式
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontName=font_name,
                    fontSize=16,
                    alignment=1,  # 居中
                    spaceAfter=20
                )
                
                # 添加标题
                title = Paragraph("文件脱敏日志报告", title_style)
                elements.append(title)
                elements.append(Spacer(1, 12))
                
                # 添加基本信息
                info_style = ParagraphStyle(
                    'InfoStyle',
                    parent=styles['Normal'],
                    fontName=font_name,
                    fontSize=12,
                    alignment=1,  # 居中对齐
                    spaceAfter=15
                )
                
                # 统计各类型记录数量
                text_records = len(self.text_redaction_history)
                word_records = len(self.word_redaction_history)
                pdf_records = len(self.pdf_redaction_history)
                excel_records = sum(len(entry.get('operations', [])) for entry in self.excel_redaction_history)
                total_records = text_records + word_records + pdf_records + excel_records
                
                # 格式化导出时间
                export_time = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
                
                info_text = f"""
导出时间：{export_time}<br/>
总记录数：{total_records} 条<br/>
文本记录：{text_records} 条<br/>
Word记录：{word_records} 条<br/>
PDF记录：{pdf_records} 条<br/>
Excel记录：{excel_records} 条
                """
                
                info_para = Paragraph(info_text, info_style)
                elements.append(info_para)
                elements.append(Spacer(1, 20))

                # 创建表格数据
                table_data = [
                    ['序号', '原文件路径', '文件类型', '脱敏方式', '位置', '原始内容', '脱敏后内容', '影响数量', '脱敏规则', '操作时间']
                ]
                table_data.extend(all_records)

                # 创建表格
                table = Table(table_data, colWidths=[
                    0.5*inch,   # 序号
                    2.0*inch,   # 原文件路径  
                    0.8*inch,   # 文件类型
                    1.0*inch,   # 脱敏方式
                    1.2*inch,   # 位置
                    1.5*inch,   # 原始内容
                    1.5*inch,   # 脱敏后内容
                    0.8*inch,   # 影响数量
                    1.0*inch,   # 脱敏规则
                    1.2*inch    # 操作时间
                ])

                # 设置表格样式
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), font_name),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), font_name),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
                ]))

                elements.append(table)
                
                # 生成PDF
                doc.build(elements)

                QMessageBox.information(self, "导出成功", f"已成功导出 {len(all_records)} 条日志记录到：\n{file_path}")
            else:
                QMessageBox.warning(self, "提示", "没有可导出的日志记录")

        except ImportError:
            QMessageBox.warning(self, "依赖缺失", "导出PDF需要reportlab库\n请运行: pip install reportlab")
        except PermissionError:
            QMessageBox.critical(self, "文件保存失败", 
                "❌ 目标文件正在被其他程序占用或锁定！\n\n"
                "💡 解决方法：\n"
                "1️⃣ 关闭所有正在使用该文件的程序（如PDF阅读器等）\n"
                "2️⃣ 检查文件是否为只读状态，右键文件→属性→取消只读\n"
                "3️⃣ 如果是同步盘文件，等待同步完成后重试\n"
                "4️⃣ 尝试选择其他位置保存文件")
        except Exception as e:
            if "Permission denied" in str(e) or "errno 13" in str(e).lower():
                QMessageBox.critical(self, "文件保存失败", 
                    "❌ 目标文件正在被其他程序占用或锁定！\n\n"
                    "💡 解决方法：\n"
                    "1️⃣ 关闭所有正在使用该文件的程序（如PDF阅读器等）\n"
                    "2️⃣ 检查文件是否为只读状态，右键文件→属性→取消只读\n"
                    "3️⃣ 如果是同步盘文件，等待同步完成后重试\n"
                    "4️⃣ 尝试选择其他位置保存文件")
            else:
                QMessageBox.critical(self, "导出失败", f"导出PDF日志时发生错误：{str(e)}")

    def truncate_text(self, text, max_length):
        """截断文本，保留关键内容，处理更美观"""
        if not text:
            return ""
        
        text = str(text).strip()
        if len(text) <= max_length:
            return text
        
        # 对于文件路径，优先保留文件名
        if ('\\' in text or '/' in text) and ('.' in text):
            # 提取文件名
            filename = text.split('\\')[-1] if '\\' in text else text.split('/')[-1]
            if len(filename) <= max_length:
                return filename
            else:
                # 文件名也太长，截断文件名
                name_part, ext_part = os.path.splitext(filename)
                if len(ext_part) + 3 < max_length:  # 保留扩展名
                    return name_part[:max_length-len(ext_part)-3] + '...' + ext_part
                else:
                    return filename[:max_length-3] + '...'
        
        # 对于普通文本内容
        if max_length <= 3:
            return text[:max_length]
        
        # 优化显示：如果是中文为主，按字符截断；如果是英文/数字为主，尽量按词截断
        if len([c for c in text if '\u4e00' <= c <= '\u9fff']) > len(text) * 0.5:
            # 中文内容：保留前面大部分 + ...
            return text[:max_length-3] + '...' if len(text) > max_length else text
        else:
            # 英文/数字内容：尝试在合适位置截断
            if max_length <= 10:
                return text[:max_length-3] + '...'
            
            # 找合适的截断点（空格、标点等）
            truncate_pos = max_length - 3
            for i in range(min(truncate_pos, len(text)-1), max(truncate_pos-5, 0), -1):
                if text[i] in ' .,;:!?，。；：！？':
                    return text[:i] + '...'
            
            # 没找到合适截断点，直接截断
            return text[:max_length-3] + '...'

    def get_excel_column_letter(self, col_num):
        """将列号转换为Excel列字母（如1->A, 2->B, 27->AA）"""
        result = ""
        while col_num > 0:
            col_num -= 1
            result = chr(col_num % 26 + ord('A')) + result
            col_num //= 26
        return result
    
    def undo_single_cell_redaction(self, cell_item):
        """撤销单个单元格的脱敏"""
        row = cell_item.row()
        col = cell_item.column()
        current_text = cell_item.text()
        
        # 在历史记录中查找该单元格的最新脱敏记录
        restored_count = 0
        for entry in reversed(self.excel_redaction_history):
            # 只处理新格式
            if 'operations' in entry:
                for operation in reversed(entry['operations']):
                    if (operation.get('row') == row and 
                        operation.get('col') == col and 
                        operation.get('redacted_text') == current_text):
                        cell_item.setText(operation['original_text'])
                        cell_item.setBackground(operation.get('original_background', QColor()))
                        cell_item.setToolTip(operation.get('original_tooltip', ''))
                        entry['operations'].remove(operation)
                        if not entry['operations']:
                            self.excel_redaction_history.remove(entry)
                        restored_count = 1
                        break
                if restored_count > 0:
                    break
        
        if restored_count > 0:
            QMessageBox.information(self, "撤销完成", f"已撤销单元格({row+1}, {col+1})的脱敏内容")
        else:
            QMessageBox.information(self, "提示", f"未找到单元格({row+1}, {col+1})的脱敏历史记录")
    
    def undo_region_redaction(self, selected_items):
        """撤销选中区域的脱敏"""
        # 获取选中单元格的位置集合
        selected_positions = {(item.row(), item.column()) for item in selected_items}
        
        restored_count = 0
        operations_to_remove = []  # 记录需要移除的操作
        entries_to_remove = []     # 记录需要移除的整个entry
        
        # 遍历历史记录
        for entry in reversed(self.excel_redaction_history):
            if 'operations' in entry:
                for operation in reversed(entry['operations']):
                    pos = (operation.get('row'), operation.get('col'))
                    if pos in selected_positions:
                        cell_item = self.table_widget.item(pos[0], pos[1])
                        if (cell_item and cell_item.text() == operation.get('redacted_text')):
                            cell_item.setText(operation['original_text'])
                            cell_item.setBackground(operation.get('original_background', QColor()))
                            cell_item.setToolTip(operation.get('original_tooltip', ''))
                            operations_to_remove.append((entry, operation))
                            restored_count += 1
        for entry, operation in operations_to_remove:
            if operation in entry['operations']:
                entry['operations'].remove(operation)
                if not entry['operations']:
                    entries_to_remove.append(entry)
        for entry in entries_to_remove:
            if entry in self.excel_redaction_history:
                self.excel_redaction_history.remove(entry)
        if restored_count > 0:
            QMessageBox.information(self, "撤销完成", f"已撤销选中区域内 {restored_count} 个单元格的脱敏内容")
        else:
            QMessageBox.information(self, "提示", "未找到选中区域内的脱敏历史记录")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = UniversalRedactionTool()
    window.show()
    sys.exit(app.exec_())
